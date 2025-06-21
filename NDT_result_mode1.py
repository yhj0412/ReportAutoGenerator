#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
NDT结果通知单台账Mode1处理模块
将Excel数据填入Word文档 - Mode1模式

作者: NDT报告生成器
日期: 2025-06-21
"""

import os
import sys
import pandas as pd
from docx import Document
import argparse
import re
from datetime import datetime

def update_date_in_cell(cell, year, month, day):
    """更新单元格中的日期"""
    # 检查单元格中的所有段落
    date_found = False
    for paragraph in cell.paragraphs:
        if "年" in paragraph.text and "月" in paragraph.text and "日" in paragraph.text:
            print(f"找到日期段落: {paragraph.text}")

            # 创建新的文本，确保只有一个年月日
            new_text = paragraph.text
            # 确保年月日前没有数字
            new_text = re.sub(r'\d*年', '年', new_text)
            new_text = re.sub(r'\d*月', '月', new_text)
            new_text = re.sub(r'\d*日', '日', new_text)

            # 在年月日前插入正确的数字
            new_text = new_text.replace('年', f'{year}年')
            new_text = new_text.replace('月', f'{month}月')
            new_text = new_text.replace('日', f'{day}日')

            paragraph.text = new_text
            date_found = True
            print("已更新日期")
            break

    # 如果没有找到日期段落，尝试在现有文本后添加日期
    if not date_found:
        print("未找到日期段落，尝试添加日期...")
        if cell.paragraphs:
            # 在现有文本后添加日期
            current_text = cell.paragraphs[0].text
            if current_text and not current_text.endswith('：'):
                current_text += ' '
            cell.paragraphs[0].text = current_text + f"{year}年{month}月{day}日"
            print(f"已添加日期: {year}年{month}月{day}日")

def process_excel_to_word(excel_path, word_template_path, output_path=None, 
                         project_name=None, client_name=None, inspection_unit=None, 
                         inspection_standard=None, inspection_method=None):
    """将Excel数据填入Word文档 - Mode1模式
    
    Args:
        excel_path: Excel表格路径
        word_template_path: Word模板文档路径
        output_path: 输出Word文档路径（如果为None，将自动生成）
        project_name: 工程名称，用于替换文档中的"工程名称值"
        client_name: 委托单位，用于替换文档中的"委托单位值"
        inspection_unit: 检测单位，用于替换文档中的"检测单位值"
        inspection_standard: 检测标准，用于替换文档中的"检测标准值"
        inspection_method: 检测方法，用于替换文档中的"检测方法值"
    
    Returns:
        bool: 处理是否成功
    """
    # 检查文件是否存在
    if not os.path.exists(excel_path):
        print(f"错误: Excel文件不存在: {excel_path}")
        return False
        
    if not os.path.exists(word_template_path):
        print(f"错误: Word模板文件不存在: {word_template_path}")
        return False
    
    # 创建输出目录
    output_dir = os.path.join("生成器", "输出报告", "2_RT结果通知单台账_Mode", "2_RT结果通知单台账_Mode1")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    try:
        # 读取Excel文件
        print(f"正在读取Excel文件: {excel_path}")
        df = pd.read_excel(excel_path)
        print(f"Excel文件读取成功，共{len(df)}行数据")
        
        # 显示列名以便调试
        print("Excel文件列名:")
        for i, col in enumerate(df.columns):
            print(f"  {i}: {col}")
        
        # 建立列名映射 - 根据实际Excel列结构
        column_mapping = {}
        for col in df.columns:
            col_str = str(col).strip()
            if '完成日期' in col_str:
                column_mapping['完成日期'] = col
            elif '委托单编号' in col_str:
                column_mapping['委托单编号'] = col
            elif '检件编号' in col_str:
                column_mapping['检件编号'] = col
            elif '焊口编号' in col_str:
                column_mapping['焊口编号'] = col
            elif '规格' in col_str and '底片' not in col_str:
                column_mapping['规格'] = col
            elif '材质' in col_str:
                column_mapping['材质'] = col
            elif '合格级别' in col_str:
                column_mapping['合格级别'] = col
            elif '底片规格' in col_str or col_str == '底片规格/张数':
                column_mapping['底片规格/张数'] = col
            elif col_str == '张数':
                column_mapping['张数'] = col
            elif '合格张数' in col_str:
                column_mapping['合格张数'] = col
            elif '单元名称' in col_str:
                column_mapping['单元名称'] = col
        
        print("找到的列映射:")
        for key, value in column_mapping.items():
            print(f"  {key}: {value}")
        
        # 检查必需的列是否都找到了
        required_columns = ['委托单编号', '完成日期', '检件编号']
        missing_columns = []
        for col in required_columns:
            if col not in column_mapping:
                missing_columns.append(col)
        
        if missing_columns:
            print(f"警告: 未找到以下必需列: {', '.join(missing_columns)}")
            # 尝试使用列位置
            possible_columns = {
                '完成日期': 'B', 
                '委托单编号': 'C', 
                '检件编号': 'D', 
                '焊口编号': 'E', 
                '规格': 'G', 
                '材质': 'H', 
                '合格级别': 'I', 
                '底片规格/张数': 'L', 
                '张数': 'M', 
                '合格张数': 'N', 
                '单元名称': 'Q'
            }
            
            for key in missing_columns:
                col_letter = possible_columns.get(key)
                if col_letter:
                    col_idx = ord(col_letter) - ord('A')
                    if col_idx < len(df.columns):
                        column_mapping[key] = df.columns[col_idx]
                        print(f"使用列位置找到: '{key}' -> '{df.columns[col_idx]}'")
        
        # 检查必需的列是否都找到了
        for col in required_columns:
            if col not in column_mapping:
                print(f"错误: 未找到必需的列: '{col}'")
                return False
        
        # 按委托单编号分组
        order_column = column_mapping['委托单编号']
        grouped = df.groupby(order_column)
        
        print(f"\n按委托单编号分组，共{len(grouped)}组:")
        for order_number, group in grouped:
            print(f"  委托单编号: {order_number}, 数据行数: {len(group)}")
        
        # 处理每个委托单编号的数据
        success_count = 0
        error_count = 0
        
        for order_number, group_data in grouped:
            try:
                print(f"\n==== 处理委托单编号: {order_number} ====")
                print(f"该组数据行数: {len(group_data)}")
                
                # 加载Word模板
                doc = Document(word_template_path)
                print("Word模板加载成功")
                
                # 获取完成日期的最晚日期
                completion_date_column = column_mapping.get('完成日期')
                if completion_date_column:
                    completion_dates = group_data[completion_date_column].dropna()
                    if not completion_dates.empty:
                        # 转换为日期类型并找到最晚日期
                        try:
                            completion_dates_converted = pd.to_datetime(completion_dates, errors='coerce')
                            latest_completion_date = completion_dates_converted.max()
                            
                            if pd.notna(latest_completion_date):
                                year = latest_completion_date.year
                                month = latest_completion_date.month
                                day = latest_completion_date.day
                                print(f"最晚完成日期: {year}年{month}月{day}日")
                            else:
                                print("警告: 无法解析完成日期")
                                year, month, day = 2024, 1, 1
                        except Exception as e:
                            print(f"日期转换错误: {e}")
                            year, month, day = 2024, 1, 1
                    else:
                        print("警告: 完成日期列为空")
                        year, month, day = 2024, 1, 1
                else:
                    print("警告: 未找到完成日期列")
                    year, month, day = 2024, 1, 1
                
                # 替换文档中的参数值
                if any([project_name, client_name, inspection_unit, inspection_standard, inspection_method]):
                    print("\n==== 开始替换参数值 ====")
                    
                    # 遍历所有段落和表格中的单元格，替换参数值
                    # 1. 遍历段落
                    for paragraph in doc.paragraphs:
                        if project_name and "工程名称值" in paragraph.text:
                            paragraph.text = paragraph.text.replace("工程名称值", project_name)
                            print(f"已将段落中的'工程名称值'替换为'{project_name}'")
                        
                        if client_name and "委托单位值" in paragraph.text:
                            paragraph.text = paragraph.text.replace("委托单位值", client_name)
                            print(f"已将段落中的'委托单位值'替换为'{client_name}'")
                        
                        if inspection_unit and "检测单位值" in paragraph.text:
                            paragraph.text = paragraph.text.replace("检测单位值", inspection_unit)
                            print(f"已将段落中的'检测单位值'替换为'{inspection_unit}'")
                        
                        if inspection_standard and "检测标准值" in paragraph.text:
                            paragraph.text = paragraph.text.replace("检测标准值", inspection_standard)
                            print(f"已将段落中的'检测标准值'替换为'{inspection_standard}'")
                        
                        if inspection_method and "检测方法值" in paragraph.text:
                            paragraph.text = paragraph.text.replace("检测方法值", inspection_method)
                            print(f"已将段落中的'检测方法值'替换为'{inspection_method}'")
                
                # 处理单值替换（合格级别、单元名称、完成日期）
                print("\n==== 开始处理单值替换 ====")
                
                # 获取合格级别值
                qualification_level = ""
                if '合格级别' in column_mapping:
                    qual_values = group_data[column_mapping['合格级别']].dropna()
                    if not qual_values.empty:
                        qualification_level = str(qual_values.iloc[0])
                        print(f"合格级别值: {qualification_level}")
                
                # 获取单元名称值
                unit_name = ""
                if '单元名称' in column_mapping:
                    unit_values = group_data[column_mapping['单元名称']].dropna()
                    if not unit_values.empty:
                        unit_name = str(unit_values.iloc[0])
                        print(f"单元名称值: {unit_name}")
                
                # 在文档中替换这些值
                for paragraph in doc.paragraphs:
                    if "合格级别值" in paragraph.text and qualification_level:
                        paragraph.text = paragraph.text.replace("合格级别值", qualification_level)
                        print(f"已将'合格级别值'替换为'{qualification_level}'")

                    if "单元名称值" in paragraph.text and unit_name:
                        paragraph.text = paragraph.text.replace("单元名称值", unit_name)
                        print(f"已将'单元名称值'替换为'{unit_name}'")

                    if "完成日期值" in paragraph.text:
                        completion_date_str = f"{year}年{month}月{day}日"
                        paragraph.text = paragraph.text.replace("完成日期值", completion_date_str)
                        print(f"已将'完成日期值'替换为'{completion_date_str}'")

                # 处理表格中的单值替换
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if "合格级别值" in cell.text and qualification_level:
                                cell.text = cell.text.replace("合格级别值", qualification_level)
                                print(f"已将表格中的'合格级别值'替换为'{qualification_level}'")

                            if "单元名称值" in cell.text and unit_name:
                                cell.text = cell.text.replace("单元名称值", unit_name)
                                print(f"已将表格中的'单元名称值'替换为'{unit_name}'")

                            if "完成日期值" in cell.text:
                                completion_date_str = f"{year}年{month}月{day}日"
                                cell.text = cell.text.replace("完成日期值", completion_date_str)
                                print(f"已将表格中的'完成日期值'替换为'{completion_date_str}'")

                # 处理日期填入（施工单位、监理单位、项目部/装置、检测单位）
                print("\n==== 开始处理日期填入 ====")
                date_keywords = ["施工单位：", "监理单位：", "项目部/装置：", "检测单位："]

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            for keyword in date_keywords:
                                if keyword in cell_text:
                                    print(f"找到{keyword}单元格")
                                    # 更新单元格中的日期
                                    update_date_in_cell(cell, year, month, day)

                # 处理表格数据填入
                print("\n==== 开始处理表格数据填入 ====")

                # 准备数据
                pipe_numbers = []  # 管线/检件编号
                weld_numbers = []  # 焊口编号
                materials = []     # 材质
                specifications = [] # 规格
                film_specs = []    # 底片规格/数量（张）
                qualified_counts = [] # 合格
                unqualified_counts = [] # 不合格

                for idx, row in group_data.iterrows():
                    # 检件编号
                    if '检件编号' in column_mapping:
                        pipe_num = row[column_mapping['检件编号']]
                        pipe_numbers.append(str(pipe_num) if pd.notna(pipe_num) else "")

                    # 焊口编号
                    if '焊口编号' in column_mapping:
                        weld_num = row[column_mapping['焊口编号']]
                        weld_numbers.append(str(weld_num) if pd.notna(weld_num) else "")

                    # 材质
                    if '材质' in column_mapping:
                        material = row[column_mapping['材质']]
                        materials.append(str(material) if pd.notna(material) else "")

                    # 规格
                    if '规格' in column_mapping:
                        spec = row[column_mapping['规格']]
                        specifications.append(str(spec) if pd.notna(spec) else "")

                    # 底片规格/张数
                    if '底片规格/张数' in column_mapping:
                        film_spec = row[column_mapping['底片规格/张数']]
                        film_specs.append(str(film_spec) if pd.notna(film_spec) else "")

                    # 合格张数
                    if '合格张数' in column_mapping:
                        qualified = row[column_mapping['合格张数']]
                        qualified_counts.append(str(qualified) if pd.notna(qualified) else "0")

                    # 不合格张数（张数-合格张数）
                    total_count = 0
                    qualified_count = 0

                    if '张数' in column_mapping:
                        total = row[column_mapping['张数']]
                        if pd.notna(total):
                            try:
                                total_count = int(float(total))
                            except:
                                total_count = 0

                    if '合格张数' in column_mapping:
                        qualified = row[column_mapping['合格张数']]
                        if pd.notna(qualified):
                            try:
                                qualified_count = int(float(qualified))
                            except:
                                qualified_count = 0

                    unqualified_count = max(0, total_count - qualified_count)
                    unqualified_counts.append(str(unqualified_count))

                print(f"准备填入表格的数据行数: {len(pipe_numbers)}")

                # 查找并填入表格数据
                table_found = False
                for table_idx, table in enumerate(doc.tables):
                    # 查找表格头部，确定这是数据表格
                    header_row = None
                    column_indices = {}

                    for row_idx, row in enumerate(table.rows):
                        row_text = " ".join([cell.text.strip() for cell in row.cells])
                        if ("管线" in row_text or "检件编号" in row_text) and "焊口编号" in row_text:
                            header_row = row_idx
                            print(f"找到数据表格#{table_idx+1}，表头在第{row_idx+1}行")

                            # 确定列索引 - 分析表头行，包括多行表头结构
                            print(f"正在分析表头行 {row_idx+1}:")
                            for col_idx, cell in enumerate(row.cells):
                                cell_text = cell.text.strip()
                                print(f"  列 {col_idx}: '{cell_text}'")
                                if "管线" in cell_text or "检件编号" in cell_text:
                                    column_indices["管线/检件编号"] = col_idx
                                elif "焊口编号" in cell_text:
                                    column_indices["焊口编号"] = col_idx
                                elif "材质" in cell_text:
                                    column_indices["材质"] = col_idx
                                elif "规格" in cell_text and "底片" not in cell_text:
                                    column_indices["规格"] = col_idx
                                elif "底片规格" in cell_text or "数量" in cell_text or ("张" in cell_text and "合格" not in cell_text):
                                    column_indices["底片规格/数量（张）"] = col_idx
                                elif "合格" in cell_text and "不合格" not in cell_text:
                                    column_indices["合格"] = col_idx
                                elif "不合格" in cell_text:
                                    column_indices["不合格"] = col_idx

                            # 检查下一行是否有"合格"和"不合格"列
                            if row_idx + 1 < len(table.rows):
                                next_row = table.rows[row_idx + 1]
                                print(f"正在分析下一行 {row_idx+2}:")
                                for col_idx, cell in enumerate(next_row.cells):
                                    cell_text = cell.text.strip()
                                    print(f"  列 {col_idx}: '{cell_text}'")
                                    if "合格" in cell_text and "不合格" not in cell_text:
                                        column_indices["合格"] = col_idx
                                        print(f"  找到合格列: {col_idx}")
                                    elif "不合格" in cell_text:
                                        column_indices["不合格"] = col_idx
                                        print(f"  找到不合格列: {col_idx}")

                            print(f"找到的列索引: {column_indices}")
                            table_found = True
                            break

                    if table_found:
                        # 填入数据 - 确保从表头的下一行开始填充数据，并保护表头不被覆盖
                        data_start_row = header_row + 1
                        print(f"开始从第{data_start_row+1}行填充数据，共{len(pipe_numbers)}行数据")

                        # 寻找第一个非表头的数据行开始填充
                        actual_data_start_row = data_start_row
                        for check_row_idx in range(data_start_row, len(table.rows)):
                            check_row = table.rows[check_row_idx]
                            if check_row.cells:
                                first_cell_text = check_row.cells[0].text.strip()
                                # 如果不是表头行，则从这里开始填充数据
                                if not any(keyword in first_cell_text for keyword in ["管线", "检件编号", "焊口编号", "材质", "规格", "底片", "合格", "检测方法"]):
                                    actual_data_start_row = check_row_idx
                                    print(f"找到实际数据开始行: 第{actual_data_start_row+1}行")
                                    break

                        for i in range(len(pipe_numbers)):
                            row_idx = actual_data_start_row + i
                            if row_idx < len(table.rows):
                                row = table.rows[row_idx]
                                print(f"正在填充第{row_idx+1}行数据...")

                                # 再次检查当前行是否为表头行，如果是则跳过
                                is_header_row = False
                                if row.cells:
                                    first_cell_text = row.cells[0].text.strip()
                                    # 检查是否包含表头关键词
                                    if any(keyword in first_cell_text for keyword in ["管线", "检件编号", "焊口编号", "材质", "规格", "底片", "合格", "检测方法"]):
                                        print(f"跳过第{row_idx+1}行，这是表头行: {first_cell_text}")
                                        is_header_row = True
                                        continue  # 跳过这一行，继续下一行

                                if not is_header_row:
                                    # 填入各列数据
                                    if "管线/检件编号" in column_indices and i < len(pipe_numbers):
                                        col_idx = column_indices["管线/检件编号"]
                                        if col_idx < len(row.cells):
                                            cell = row.cells[col_idx]
                                            if cell.paragraphs:
                                                cell.paragraphs[0].text = pipe_numbers[i]
                                                print(f"已更新第{row_idx+1}行管线/检件编号: {pipe_numbers[i]}")

                                    if "焊口编号" in column_indices and i < len(weld_numbers):
                                        col_idx = column_indices["焊口编号"]
                                        if col_idx < len(row.cells):
                                            cell = row.cells[col_idx]
                                            if cell.paragraphs:
                                                cell.paragraphs[0].text = weld_numbers[i]
                                                print(f"已更新第{row_idx+1}行焊口编号: {weld_numbers[i]}")

                                    if "材质" in column_indices and i < len(materials):
                                        col_idx = column_indices["材质"]
                                        if col_idx < len(row.cells):
                                            cell = row.cells[col_idx]
                                            if cell.paragraphs:
                                                cell.paragraphs[0].text = materials[i]
                                                print(f"已更新第{row_idx+1}行材质: {materials[i]}")

                                    if "规格" in column_indices and i < len(specifications):
                                        col_idx = column_indices["规格"]
                                        if col_idx < len(row.cells):
                                            cell = row.cells[col_idx]
                                            if cell.paragraphs:
                                                cell.paragraphs[0].text = specifications[i]
                                                print(f"已更新第{row_idx+1}行规格: {specifications[i]}")

                                    if "底片规格/数量（张）" in column_indices and i < len(film_specs):
                                        col_idx = column_indices["底片规格/数量（张）"]
                                        if col_idx < len(row.cells):
                                            cell = row.cells[col_idx]
                                            if cell.paragraphs:
                                                cell.paragraphs[0].text = film_specs[i]
                                                print(f"已更新第{row_idx+1}行底片规格/数量: {film_specs[i]}")

                                    if "合格" in column_indices and i < len(qualified_counts):
                                        col_idx = column_indices["合格"]
                                        if col_idx < len(row.cells):
                                            cell = row.cells[col_idx]
                                            if cell.paragraphs:
                                                cell.paragraphs[0].text = qualified_counts[i]
                                                print(f"已更新第{row_idx+1}行合格: {qualified_counts[i]}")

                                    if "不合格" in column_indices and i < len(unqualified_counts):
                                        col_idx = column_indices["不合格"]
                                        if col_idx < len(row.cells):
                                            cell = row.cells[col_idx]
                                            if cell.paragraphs:
                                                cell.paragraphs[0].text = unqualified_counts[i]
                                                print(f"已更新第{row_idx+1}行不合格: {unqualified_counts[i]}")
                            else:
                                print(f"警告: 表格行数不足，无法填充第{i+1}条数据")
                        break

                if not table_found:
                    print("警告: 未找到合适的数据表格")

                # 保存文档
                report_output_path = os.path.join(output_dir, f"{order_number}_RT结果通知单台账_Mode1.docx")
                
                try:
                    print(f"\n正在保存文档到: {report_output_path}")
                    doc.save(report_output_path)
                    print(f"文档已成功保存至: {report_output_path}")
                    success_count += 1
                except Exception as e:
                    print(f"错误: 无法保存文档: {e}")
                    error_count += 1
                    
            except Exception as e:
                print(f"错误: 处理委托单编号 {order_number} 时出错: {e}")
                error_count += 1
        
        print(f"\n==== 处理完成 ====")
        print(f"成功处理: {success_count} 个文档")
        print(f"处理失败: {error_count} 个文档")
        
        return error_count == 0
        
    except Exception as e:
        print(f"错误: 处理过程中出现异常: {e}")
        return False

def main():
    # 创建命令行参数解析器
    parser = argparse.ArgumentParser(description='将Excel数据填入Word文档 - Mode1模式')
    parser.add_argument('-e', '--excel', default="生成器/Excel/2_生成器结果.xlsx",
                        help='Excel表格路径 (默认: 生成器/Excel/2_生成器结果.xlsx)')
    parser.add_argument('-w', '--word', default="生成器/wod/2_RT结果通知台账_Mode1.docx", 
                        help='Word模板文档路径 (默认: 生成器/wod/2_RT结果通知台账_Mode1.docx)')
    parser.add_argument('-o', '--output', 
                        help='输出目录 (可选，默认为"生成器/输出报告/2_RT结果通知单台账_Mode/2_RT结果通知单台账_Mode1"目录)')
    parser.add_argument('-p', '--project', 
                        help='工程名称，用于替换文档中的"工程名称值"')
    parser.add_argument('-c', '--client', 
                        help='委托单位，用于替换文档中的"委托单位值"')
    parser.add_argument('-u', '--unit', 
                        help='检测单位，用于替换文档中的"检测单位值"')
    parser.add_argument('-s', '--standard', 
                        help='检测标准，用于替换文档中的"检测标准值"')
    parser.add_argument('-m', '--method', 
                        help='检测方法，用于替换文档中的"检测方法值"')
    
    # 解析命令行参数
    args = parser.parse_args()
    
    # 处理Excel到Word的转换
    success = process_excel_to_word(
        args.excel, args.word, args.output, 
        args.project, args.client, args.unit, 
        args.standard, args.method
    )
    
    # 返回状态码
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
