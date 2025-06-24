import pandas as pd
import os
import sys
import argparse
from docx import Document
from datetime import datetime
import re

def find_column_with_keyword(df, keyword):
    """查找包含指定关键字的列"""
    matching_cols = [col for col in df.columns if keyword.lower() in col.lower()]
    return matching_cols[0] if matching_cols else None

def get_detection_level_by_method(detection_method):
    """根据检测方法获取对应的检测级别值

    Args:
        detection_method: 检测方法字符串

    Returns:
        str: 对应的检测级别值
    """
    if not detection_method:
        return ""

    # 转换为字符串并去除空格，转为大写进行匹配
    method = str(detection_method).strip().upper()

    # 检测方法与检测级别值的映射关系
    method_mapping = {
        # 硬度检测或YD
        '硬度检测': '力学   级',
        'YD': '力学   级',
        # 光谱检测或PMIN
        '光谱检测': '光谱分析  级',
        'PMIN': '光谱分析  级',
        # 其他检测方法
        'UT': 'UT  级',
        'PT': 'PT  级',
        'MT': 'MT  级',
        'RT': 'RT  级',
        'TOFD': 'TOFD  级',
        'PA': 'PA  级'
    }

    # 精确匹配
    if method in method_mapping:
        return method_mapping[method]

    # 模糊匹配 - 检查是否包含关键字
    for key, value in method_mapping.items():
        if key in method:
            return value

    # 如果没有匹配到，返回空字符串
    print(f"警告: 未找到检测方法 '{detection_method}' 对应的检测级别值")
    return ""

def get_output_filename(word_template_path, order_number):
    """根据Word模板路径和委托单编号生成输出文件名
    
    Args:
        word_template_path: Word模板文档路径
        order_number: 委托单编号
    
    Returns:
        str: 输出文件名
    """
    # 获取模板文件名（不含路径和扩展名）
    template_name = os.path.splitext(os.path.basename(word_template_path))[0]
    # 生成输出文件名
    return f"{template_name}_{order_number}_生成结果.docx"

def process_excel_to_word(excel_path, word_template_path, output_path=None, project_name=None, client_name=None, inspection_method=None):
    """将Excel数据填入Word文档
    
    Args:
        excel_path: Excel表格路径
        word_template_path: Word模板文档路径
        output_path: 输出Word文档路径（如果为None，将自动生成）
        project_name: 工程名称，用于替换文档中的"工程名称参数值"
        client_name: 委托单位，用于替换文档中的"委托单位参数值"
        inspection_method: 检测方法，用于替换文档中的"检测方法参数"
    
    Returns:
        bool: 处理是否成功
    """
    # 检查文件是否存在
    if not os.path.exists(excel_path):
        print(f"错误: Excel文件不存在: {excel_path}")
        return False
        
    if not os.path.exists(word_template_path):
        print(f"错误: Word模板文件不存在: {word_template_path}")
        return False
    
    # 创建输出目录 - 修改为指定的路径
    output_dir = os.path.join("生成器", "输出报告", "3_表面结果通知单台账","3_表面结果通知单台账_Mode2")
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            print(f"创建输出目录: {output_dir}")
        except Exception as e:
            print(f"错误: 无法创建输出目录: {e}")
            return False
    
    # 读取Excel数据 - 指定读取sheet3"荣信聚乙烯PT"
    print(f"正在读取Excel文件: {excel_path}")
    try:
        # 读取指定的工作表sheet3"荣信聚乙烯PT"
        df = pd.read_excel(excel_path, sheet_name="荣信聚乙烯PT")
        print(f"成功读取Excel文件sheet3'荣信聚乙烯PT'，共有{len(df)}行数据")
    except Exception as e:
        print(f"错误: 无法读取Excel文件sheet3'荣信聚乙烯PT': {e}")
        # 如果指定工作表不存在，尝试读取第一个工作表
        try:
            df = pd.read_excel(excel_path)
            print(f"警告: 未找到sheet3'荣信聚乙烯PT'，使用默认工作表，共有{len(df)}行数据")
        except Exception as e2:
            print(f"错误: 无法读取Excel文件: {e2}")
            return False
    
    # 打印所有列名，帮助调试
    print(f"Excel表格列名: {list(df.columns)}")
    
    # 定义需要查找的列关键字 - 根据新需求更新
    column_keywords = {
        '完成日期': '完成日期',           # B列
        '委托单编号': '委托单编号',       # C列
        '检件编号': '检件编号',           # D列
        '焊口编号': '焊口编号',           # E列
        '焊工号': '焊工号',               # F列
        '焊口情况': '焊口情况',           # K列 - 对应检测结果
        '返修张处数': '返修张/处数',      # L列 - 返修张/处数
        '检测方法': '检测方法',           # N列
        '单元名称': '单元名称'            # O列
    }
    
    # 查找每个关键字对应的实际列名
    column_mapping = {}
    missing_columns = []
    
    for key, keyword in column_keywords.items():
        col_name = find_column_with_keyword(df, keyword)
        if col_name:
            column_mapping[key] = col_name
            print(f"找到列: '{key}' -> '{col_name}'")
        else:
            missing_columns.append(key)
    
    if missing_columns:
        print(f"警告: 未找到以下列: {', '.join(missing_columns)}")
        # 尝试使用列位置 - 根据新需求更新
        possible_columns = {
            '完成日期': 'B',         # B列
            '委托单编号': 'C',       # C列
            '检件编号': 'D',         # D列
            '焊口编号': 'E',         # E列
            '焊工号': 'F',           # F列
            '焊口情况': 'K',         # K列
            '返修张处数': 'L',       # L列
            '检测方法': 'N',         # N列
            '单元名称': 'O'          # O列
        }
        
        for key in missing_columns:
            col_letter = possible_columns.get(key)
            if col_letter:
                col_idx = ord(col_letter) - ord('A')
                if col_idx < len(df.columns):
                    column_mapping[key] = df.columns[col_idx]
                    print(f"使用列位置找到: '{key}' -> '{df.columns[col_idx]}'")
    
    # 按委托单编号分组处理数据
    if '委托单编号' not in column_mapping:
        print("错误: 无法找到委托单编号列")
        return False
    
    # 获取所有唯一的委托单编号
    order_numbers = df[column_mapping['委托单编号']].dropna().unique().tolist()
    print(f"找到{len(order_numbers)}个不同的委托单编号: {order_numbers}")
    
    success_count = 0
    error_count = 0
    
    # 对每个委托单编号生成一份报告
    for order_number in order_numbers:
        print(f"\n{'='*50}")
        print(f"处理委托单编号: {order_number} ({success_count+error_count+1}/{len(order_numbers)})")
        print(f"{'='*50}")
        
        try:
            # 筛选该委托单编号的数据
            order_df = df[df[column_mapping['委托单编号']] == order_number]
            print(f"该委托单编号有{len(order_df)}条记录")
            
            # 为该委托单编号生成输出文件名
            output_filename = get_output_filename(word_template_path, order_number)
            report_output_path = os.path.join(output_dir, output_filename)
            print(f"输出文件路径: {report_output_path}")
            
            # 检查输出路径是否可写
            try:
                # 尝试创建一个临时文件来测试目录是否可写
                test_file = os.path.join(output_dir, "test_write.tmp")
                with open(test_file, 'w') as f:
                    f.write("test")
                os.remove(test_file)
                print(f"输出目录可写: {output_dir}")
            except Exception as e:
                print(f"警告: 输出目录可能不可写: {e}")
                # 继续尝试，但记录警告
            
            # 1) 获取该组数据中最晚的完成日期
            date_col = column_mapping.get('完成日期')
            if date_col:
                # 确保日期列是日期类型
                order_df[date_col] = pd.to_datetime(order_df[date_col], errors='coerce')
                latest_date = order_df[date_col].max()
                
                if pd.isna(latest_date):
                    print(f"警告: 委托单编号 {order_number} 没有有效的完成日期")
                    year, month, day = datetime.now().year, datetime.now().month, datetime.now().day
                else:
                    # 将日期转换为年、月、日
                    year = latest_date.year
                    month = latest_date.month
                    day = latest_date.day
                    print(f"找到最晚完成日期: {year}年{month}月{day}日")
            else:
                print("警告: 未找到完成日期列")
                year, month, day = datetime.now().year, datetime.now().month, datetime.now().day
            
            # 获取相关数据 - 根据新需求更新
            inspection_numbers = order_df[column_mapping.get('检件编号')].dropna().tolist() if '检件编号' in column_mapping else []
            weld_numbers = order_df[column_mapping.get('焊口编号')].dropna().tolist() if '焊口编号' in column_mapping else []
            welder_numbers = order_df[column_mapping.get('焊工号')].dropna().tolist() if '焊工号' in column_mapping else []
            weld_conditions = order_df[column_mapping.get('焊口情况')].tolist() if '焊口情况' in column_mapping else []  # K列焊口情况
            repair_counts = order_df[column_mapping.get('返修张处数')].tolist() if '返修张处数' in column_mapping else []  # L列返修张/处数

            # 获取单元名称（第一个非空值）- O列
            unit_name = ""
            if '单元名称' in column_mapping:
                unit_names = order_df[column_mapping['单元名称']].dropna().tolist()
                if unit_names:
                    unit_name = unit_names[0]
                    print(f"找到单元名称: {unit_name}")

            # 获取检测方法（第一个非空值）- N列
            detection_method = ""
            detection_level = ""
            if '检测方法' in column_mapping:
                detection_methods = order_df[column_mapping['检测方法']].dropna().tolist()
                if detection_methods:
                    detection_method = detection_methods[0]
                    print(f"找到检测方法: {detection_method}")

                    # 根据检测方法获取对应的检测级别值
                    detection_level = get_detection_level_by_method(detection_method)
                    if detection_level:
                        print(f"根据检测方法 '{detection_method}' 确定检测级别值: '{detection_level}'")
            
                    # 打开Word文档
            print(f"正在处理Word文档: {word_template_path}")
            
            # 检查文件扩展名，使用不同的方法处理.doc和.docx文件
            try:
                if word_template_path.lower().endswith('.doc'):
                    # 对于.doc文件，需要先转换为.docx
                    temp_docx_path = word_template_path + 'x'
                    print(f"检测到.doc文件，尝试转换为.docx: {temp_docx_path}")
                    
                    # 尝试直接打开.doc文件
                    doc = Document(word_template_path)
                    doc.save(temp_docx_path)
                    print(f"成功转换.doc为.docx")
                    doc = Document(temp_docx_path)
                else:
                    # 对于.docx文件，直接打开
                    # 每次处理新的委托单编号时，重新从模板创建文档对象
                    # 这确保了每个委托单编号都会生成一个独立的文档
                    doc = Document(word_template_path)
                    print(f"成功从模板创建新文档")
            except Exception as e:
                print(f"无法打开Word文档: {e}")
                # 跳过当前委托单编号的处理
                error_count += 1
                continue
            
            # 替换文档中的参数值
            if project_name or client_name or inspection_method:
                print("\n==== 开始替换参数值 ====")
                
                # 遍历所有段落和表格中的单元格，替换参数值
                # 1. 遍历段落
                for paragraph in doc.paragraphs:
                    if project_name and "工程名称参数值" in paragraph.text:
                        paragraph.text = paragraph.text.replace("工程名称参数值", project_name)
                        print(f"已将段落中的'工程名称参数值'替换为'{project_name}'")
                    
                    if client_name and "委托单位参数值" in paragraph.text:
                        paragraph.text = paragraph.text.replace("委托单位参数值", client_name)
                        print(f"已将段落中的'委托单位参数值'替换为'{client_name}'")
                    
                    if inspection_method and "检测方法参数" in paragraph.text:
                        paragraph.text = paragraph.text.replace("检测方法参数", inspection_method)
                        print(f"已将段落中的'检测方法参数'替换为'{inspection_method}'")
                
                # 2. 遍历表格中的单元格
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if project_name and "工程名称参数值" in paragraph.text:
                                    paragraph.text = paragraph.text.replace("工程名称参数值", project_name)
                                    print(f"已将表格单元格中的'工程名称参数值'替换为'{project_name}'")
                                
                                if client_name and "委托单位参数值" in paragraph.text:
                                    paragraph.text = paragraph.text.replace("委托单位参数值", client_name)
                                    print(f"已将表格单元格中的'委托单位参数值'替换为'{client_name}'")
                                
                                if inspection_method and "检测方法参数" in paragraph.text:
                                    paragraph.text = paragraph.text.replace("检测方法参数", inspection_method)
                                    print(f"已将表格单元格中的'检测方法参数'替换为'{inspection_method}'")
                
                print("==== 参数值替换完成 ====\n")
            
            # 替换段落中的参数值
            for paragraph in doc.paragraphs:
                if unit_name and "单元名称值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("单元名称值", unit_name)
                    print(f"已将段落中的'单元名称值'替换为'{unit_name}'")

                if detection_method and "检测方法值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("检测方法值", detection_method)
                    print(f"已将段落中的'检测方法值'替换为'{detection_method}'")

                if detection_level and "检测级别值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("检测级别值", detection_level)
                    print(f"已将段落中的'检测级别值'替换为'{detection_level}'")

                if "委托单号编号值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("委托单号编号值", str(order_number))
                    print(f"已将段落中的'委托单号编号值'替换为'{order_number}'")
            
            # 遍历表格中的单元格替换参数值
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if unit_name and "单元名称值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("单元名称值", unit_name)
                                print(f"已将表格单元格中的'单元名称值'替换为'{unit_name}'")

                            if detection_method and "检测方法值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("检测方法值", detection_method)
                                print(f"已将表格单元格中的'检测方法值'替换为'{detection_method}'")

                            if detection_level and "检测级别值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("检测级别值", detection_level)
                                print(f"已将表格单元格中的'检测级别值'替换为'{detection_level}'")

                            if "委托单号编号值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("委托单号编号值", str(order_number))
                                print(f"已将表格单元格中的'委托单号编号值'替换为'{order_number}'")
            
            # 填写通知单编号（委托单编号）
            notification_number_updated = False
            
            print("\n==== 开始查找通知单编号位置 ====")
            
            # 尝试查找特定格式的单元格(如"RX3-03-ZYLJ-DG-RT-000*")
            for table_idx, table in enumerate(doc.tables):
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        
                        # 检查单元格内容是否符合特定格式
                        if (cell_text and 
                            (("RX" in cell_text and "-DG-RT-" in cell_text) or 
                             ("RT" in cell_text and "-DG-" in cell_text and "*" in cell_text) or
                             ("RX3-03-ZYLJ-DG-RT" in cell_text))):
                            
                            print(f"找到匹配特定格式的单元格: 表格#{table_idx+1}, 第{i+1}行, 第{j+1}列")
                            print(f"单元格内容: '{cell_text}'")
                            print(f"将替换为委托单编号: {order_number}")
                            
                            # 保存原始内容以便验证
                            original_content = cell_text
                            
                            # 修改单元格内容
                            if cell.paragraphs:
                                cell.paragraphs[0].text = str(order_number)
                                print(f"已将单元格内容从 '{original_content}' 修改为 '{order_number}'")
                                notification_number_updated = True
                                break
                    if notification_number_updated:
                        break
                if notification_number_updated:
                    break
            
            # 如果未找到匹配的单元格，尝试查找表格右上角区域
            if not notification_number_updated:
                print("\n未找到完全匹配的单元格，检查表格右上角区域...")
                for table_idx, table in enumerate(doc.tables):
                    # 只检查前3行
                    for i in range(min(3, len(table.rows))):
                        if i < len(table.rows) and len(table.rows[i].cells) > 0:
                            # 检查行中的最后一个单元格
                            last_cell = table.rows[i].cells[-1]
                            cell_text = last_cell.text.strip()
                            
                            print(f"检查表格#{table_idx+1}, 第{i+1}行, 最后一列, 内容: '{cell_text}'")
                            
                            # 检查是否包含部分匹配特征
                            if (cell_text and 
                                (("RX" in cell_text) or ("RT" in cell_text) or 
                                 ("DG" in cell_text) or ("*" in cell_text))):
                                
                                print(f"找到部分匹配的单元格: 表格#{table_idx+1}, 第{i+1}行, 最后一列")
                                print(f"单元格内容: '{cell_text}'")
                                print(f"将替换为委托单编号: {order_number}")
                                
                                # 保存原始内容以便验证
                                original_content = cell_text
                                
                                # 修改单元格内容
                                if last_cell.paragraphs:
                                    last_cell.paragraphs[0].text = str(order_number)
                                    print(f"已将单元格内容从 '{original_content}' 修改为 '{order_number}'")
                                    notification_number_updated = True
                                    break
                        if notification_number_updated:
                            break
                    if notification_number_updated:
                        break
            
            # 打印查找结果总结
            print("\n==== 通知单编号位置查找结果 ====")
            if notification_number_updated:
                print(f"成功找到并替换通知单编号位置: {order_number}")
            else:
                print("警告: 未能找到合适的位置填写通知单编号")
            
            # 填写单位工程名称
            for paragraph in doc.paragraphs:
                if "单位工程名称" in paragraph.text:
                    # 找到包含"单位工程名称"的段落
                    print(f"找到单位工程名称段落: {paragraph.text}")
                    
                    # 检查是否在表格中
                    found_in_table = False
                    for table in doc.tables:
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                if "单位工程名称" in cell.text and j + 1 < len(row.cells):
                                    # 在右侧单元格填写单元名称
                                    right_cell = row.cells[j + 1]
                                    if right_cell.paragraphs and unit_name:
                                        right_cell.paragraphs[0].text = unit_name
                                        print(f"已将单元名称 {unit_name} 填入单位工程名称右侧单元格")
                                        found_in_table = True
                                        break
                            if found_in_table:
                                break
                        if found_in_table:
                            break
                    
                    if not found_in_table and unit_name:
                        # 如果不在表格中，尝试修改段落文本
                        new_text = paragraph.text
                        if "：" in new_text or ":" in new_text:
                            # 如果有冒号，在冒号后添加单元名称
                            if "：" in new_text:
                                new_text = new_text.split("：")[0] + "：" + unit_name
                            else:
                                new_text = new_text.split(":")[0] + ":" + unit_name
                        else:
                            # 否则直接在文本后添加单元名称
                            new_text = new_text + " " + unit_name
                        
                        paragraph.text = new_text
                        print(f"已将单元名称 {unit_name} 添加到单位工程名称段落")
            
            # 处理表格
            for table in doc.tables:
                # 查找日期字段
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        # 1) 处理"检测人"日期
                        if "检测人" in cell.text:
                            print(f"找到检测人单元格: 第{i+1}行, 第{j+1}列")
                            
                            # 检查单元格中的所有段落
                            date_found = False
                            for paragraph in cell.paragraphs:
                                if "年" in paragraph.text and "月" in paragraph.text and "日" in paragraph.text:
                                    print(f"找到日期段落: {paragraph.text}")
                                    
                                    # 创建新的文本，确保只有一个年月日
                                    new_text = paragraph.text
                                    # 确保年月日前没有数字
                                    new_text = re.sub(r'\d*年', '年', new_text)
                                    new_text = re.sub(r'\d*月', '月', new_text)
                                    new_text = re.sub(r'\d*日', '日', new_text)
                                    
                                    # 在年月日前插入正确的数字
                                    new_text = new_text.replace('年', f'{year}年')
                                    new_text = new_text.replace('月', f'{month}月')
                                    new_text = new_text.replace('日', f'{day}日')
                                    
                                    paragraph.text = new_text
                                    date_found = True
                                    print("已更新检测人日期")
                                    break
                            
                            # 如果没有找到日期段落，尝试创建新段落
                            if not date_found:
                                print("未在检测人单元格中找到日期段落，尝试其他方法...")
                                # 添加新段落
                                p = cell.add_paragraph(f"{year}年{month}月{day}日")
                                print("已添加检测人日期")
                        
                        # 2) 处理"审核"日期
                        if "审核" in cell.text:
                            print(f"找到审核单元格: 第{i+1}行, 第{j+1}列")
                            
                            # 检查单元格中的所有段落
                            date_found = False
                            for paragraph in cell.paragraphs:
                                if "年" in paragraph.text and "月" in paragraph.text and "日" in paragraph.text:
                                    print(f"找到日期段落: {paragraph.text}")
                                    
                                    # 创建新的文本，确保只有一个年月日
                                    new_text = paragraph.text
                                    # 确保年月日前没有数字
                                    new_text = re.sub(r'\d*年', '年', new_text)
                                    new_text = re.sub(r'\d*月', '月', new_text)
                                    new_text = re.sub(r'\d*日', '日', new_text)
                                    
                                    # 在年月日前插入正确的数字
                                    new_text = new_text.replace('年', f'{year}年')
                                    new_text = new_text.replace('月', f'{month}月')
                                    new_text = new_text.replace('日', f'{day}日')
                                    
                                    paragraph.text = new_text
                                    date_found = True
                                    print("已更新审核日期")
                                    break
                            
                            # 如果没有找到日期段落，尝试创建新段落
                            if not date_found:
                                print("未在审核单元格中找到日期段落，尝试其他方法...")
                                # 添加新段落
                                p = cell.add_paragraph(f"{year}年{month}月{day}日")
                                print("已添加审核日期")
                
                # 查找表头行，确定各列的位置
                column_indices = {}
                header_row_index = -1
                
                # 查找包含"委托单编号"、"单线号"等的行
                for i, row in enumerate(table.rows):
                    header_found = False
                    for j, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if "委托单编号" in cell_text:
                            column_indices["委托单编号"] = j
                            header_row_index = i
                            header_found = True
                        elif "单线号" in cell_text:
                            column_indices["单线号"] = j
                            header_found = True
                        elif "焊口号" in cell_text:
                            column_indices["焊口号"] = j
                            header_found = True
                        elif "焊工号" in cell_text:
                            column_indices["焊工号"] = j
                            header_found = True
                        elif "检测结果" in cell_text:
                            column_indices["检测结果"] = j
                            header_found = True
                        elif "返修张/处数" in cell_text:
                            column_indices["返修张/处数"] = j
                            header_found = True
                        elif "备注" in cell_text:
                            column_indices["备注"] = j
                            header_found = True
                    
                    if header_found and header_row_index >= 0:
                        break
                
                print(f"找到表头行: 第{header_row_index+1}行")
                print(f"列索引: {column_indices}")
                
                # 如果找到表头行，处理数据填充
                if header_row_index >= 0 and column_indices:
                    # 获取可用于填充数据的行
                    data_rows = []
                    for i in range(header_row_index + 1, len(table.rows)):
                        if i < len(table.rows):
                            # 检查是否是空行或包含特殊标记的行
                            if "以下空白" in table.rows[i].cells[0].text if len(table.rows[i].cells) > 0 else False:
                                print(f"找到'以下空白'行: 第{i+1}行")
                                break
                            # 添加可用于填充数据的行
                            data_rows.append(i)
                    
                    print(f"找到{len(data_rows)}行可用于填充数据")
                    
                    # 确定需要填充的数据行数
                    data_count = len(order_df)
                    print(f"需要填充{data_count}行数据")
                    
                    # 如果Word表格中的行数不足，需要添加新行
                    rows_needed = data_count - len(data_rows)
                    if rows_needed > 0:
                        print(f"需要添加{rows_needed}行到表格中")
                        # 找到最后一行的索引
                        last_row_idx = data_rows[-1] if data_rows else header_row_index
                        
                        # 添加新行
                        for _ in range(rows_needed):
                            # 在表格末尾添加一行
                            new_row = table.add_row()
                            data_rows.append(len(table.rows) - 1)  # 添加新行的索引
                    
                    # 处理每一行数据
                    for i in range(data_count):
                        if i < len(data_rows):
                            row_idx = data_rows[i]
                            row = table.rows[row_idx]
                            
                            # 1. 填写委托单编号
                            if "委托单编号" in column_indices:
                                col_idx = column_indices["委托单编号"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    if cell.paragraphs:
                                        cell.paragraphs[0].text = str(order_number)
                                        print(f"已更新第{row_idx+1}行委托单编号: {order_number}")
                            
                            # 2. 填写单线号（检件编号）
                            if "单线号" in column_indices and i < len(inspection_numbers):
                                col_idx = column_indices["单线号"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    if cell.paragraphs:
                                        cell.paragraphs[0].text = str(inspection_numbers[i])
                                        print(f"已更新第{row_idx+1}行单线号: {inspection_numbers[i]}")
                            
                            # 3. 填写焊口号
                            if "焊口号" in column_indices and i < len(weld_numbers):
                                col_idx = column_indices["焊口号"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    if cell.paragraphs:
                                        cell.paragraphs[0].text = str(weld_numbers[i])
                                        print(f"已更新第{row_idx+1}行焊口号: {weld_numbers[i]}")
                            
                            # 4. 填写焊工号
                            if "焊工号" in column_indices and i < len(welder_numbers):
                                col_idx = column_indices["焊工号"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    if cell.paragraphs:
                                        cell.paragraphs[0].text = str(welder_numbers[i])
                                        print(f"已更新第{row_idx+1}行焊工号: {welder_numbers[i]}")
                            
                            # 5. 填写检测结果（焊口情况）- K列对应检测结果
                            if "检测结果" in column_indices and i < len(weld_conditions):
                                col_idx = column_indices["检测结果"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    weld_condition = weld_conditions[i]
                                    if cell.paragraphs:
                                        # 检查是否为空或NaN
                                        if pd.isna(weld_condition):
                                            cell.paragraphs[0].text = ""
                                        else:
                                            cell.paragraphs[0].text = str(weld_condition)
                                        print(f"已更新第{row_idx+1}行检测结果: {weld_condition}")

                            # 6. 填写返修张/处数 - L列，空值填"0"
                            if "返修张/处数" in column_indices and i < len(repair_counts):
                                col_idx = column_indices["返修张/处数"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    repair_count = repair_counts[i]
                                    if cell.paragraphs:
                                        # 检查是否为空或NaN，空值填"0"
                                        if pd.isna(repair_count) or repair_count == "":
                                            cell.paragraphs[0].text = "0"  # 为空填写0
                                        else:
                                            cell.paragraphs[0].text = str(repair_count)
                                        print(f"已更新第{row_idx+1}行返修张/处数: {cell.paragraphs[0].text}")
            
            # 保存文档
            try:
                print(f"\n正在保存文档到: {report_output_path}")
                doc.save(report_output_path)
                print(f"文档已成功保存至: {report_output_path}")
                success_count += 1
            except Exception as e:
                print(f"错误: 无法保存文档: {e}")
                error_count += 1
        except Exception as e:
            print(f"错误: 处理委托单编号 {order_number} 时出错: {e}")
            error_count += 1
    
    print(f"\n处理完成: 共处理{len(order_numbers)}个委托单编号，成功生成{success_count}份报告，失败{error_count}份")
    if error_count > 0:
        print(f"警告: 有{error_count}个委托单编号处理失败，请检查日志")
    return success_count > 0

def main():
    # 创建命令行参数解析器
    parser = argparse.ArgumentParser(description='将Excel数据填入Word文档')
    parser.add_argument('-e', '--excel', default="生成器/Excel/3_生成器表面结果.xlsx", 
                        help='Excel表格路径 (默认: 生成器/Excel/3_生成器表面结果.xlsx)')
    parser.add_argument('-w', '--word', default="生成器/word/3_表面结果通知单台账_Mode2.docx", 
                        help='Word模板文档路径 (默认: 生成器/word/3_表面结果通知单台账_Mode2.docx)')
    parser.add_argument('-o', '--output', 
                        help='输出目录 (可选，默认为"生成器/输出报告"目录)')
    parser.add_argument('-p', '--project', 
                        help='工程名称，用于替换文档中的"工程名称参数值"')
    parser.add_argument('-c', '--client', 
                        help='委托单位，用于替换文档中的"委托单位参数值"')
    parser.add_argument('-m', '--method', 
                        help='检测方法，用于替换文档中的"检测方法参数"')
    
    # 解析命令行参数
    args = parser.parse_args()
    
    # 处理Excel到Word的转换
    success = process_excel_to_word(args.excel, args.word, args.output, args.project, args.client, args.method)
    
    # 返回状态码
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main() 