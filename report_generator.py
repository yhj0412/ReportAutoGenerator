import os
import pandas as pd
from docx import Document
from datetime import datetime
import logging
import tkinter as tk
from tkinter import filedialog, messagebox
import re

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

class ReportGenerator:
    def __init__(self):
        self.excel_file = None
        self.word_template = None
        self.output_dir = None
        self.df = None
        
    def set_excel_file(self, excel_file):
        """设置Excel文件路径"""
        self.excel_file = excel_file
        logger.info(f"Excel文件已设置: {excel_file}")
        
    def set_word_template(self, word_template):
        """设置Word模板文件路径"""
        self.word_template = word_template
        logger.info(f"Word模板已设置: {word_template}")
        
    def set_output_dir(self, output_dir):
        """设置输出目录"""
        self.output_dir = output_dir
        # 确保输出目录存在
        if not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
                logger.info(f"创建输出目录: {output_dir}")
            except Exception as e:
                logger.error(f"创建输出目录失败: {str(e)}")
        logger.info(f"输出目录已设置: {output_dir}")
        
    def load_excel_data(self):
        """加载Excel数据"""
        try:
            self.df = pd.read_excel(self.excel_file)
            logger.info(f"Excel数据加载成功，共{len(self.df)}条记录")
            
            # 尝试将列名标准化
            rename_dict = {}
            for col in self.df.columns:
                # 移除任何非中文字符来匹配列名
                clean_col = re.sub(r'[^\u4e00-\u9fa5]', '', str(col))
                if '委托日期' in clean_col or '日期' in clean_col:
                    rename_dict[col] = '委托日期'
                elif '检件编号' in clean_col or '管道编号' in clean_col:
                    rename_dict[col] = '检件编号'
                elif '焊口编号' in clean_col or '焊口号' in clean_col:
                    rename_dict[col] = '焊口编号'
                elif '焊工号' in clean_col or '焊工' in clean_col:
                    rename_dict[col] = '焊工号'
                elif '规格' in clean_col:
                    rename_dict[col] = '规格'
                elif '材质' in clean_col:
                    rename_dict[col] = '材质'
            
            if rename_dict:
                self.df = self.df.rename(columns=rename_dict)
                logger.info(f"已将列名标准化: {rename_dict}")
            
            # 检查必要的列是否存在
            required_cols = ['委托日期', '检件编号', '焊口编号', '焊工号', '规格', '材质']
            missing_cols = [col for col in required_cols if col not in self.df.columns]
            
            if missing_cols:
                logger.warning(f"Excel中缺少以下列: {missing_cols}")
                logger.info(f"可用的列: {list(self.df.columns)}")
                
                # 尝试根据位置猜测列
                if len(self.df.columns) >= 8:  # 假设至少有8列
                    col_map = {}
                    if '委托日期' in missing_cols and 0 < len(self.df.columns):
                        col_map[self.df.columns[0]] = '委托日期'
                    if '检件编号' in missing_cols and 3 < len(self.df.columns):
                        col_map[self.df.columns[3]] = '检件编号'
                    if '焊口编号' in missing_cols and 4 < len(self.df.columns):
                        col_map[self.df.columns[4]] = '焊口编号'
                    if '焊工号' in missing_cols and 5 < len(self.df.columns):
                        col_map[self.df.columns[5]] = '焊工号'
                    if '规格' in missing_cols and 6 < len(self.df.columns):
                        col_map[self.df.columns[6]] = '规格'
                    if '材质' in missing_cols and 7 < len(self.df.columns):
                        col_map[self.df.columns[7]] = '材质'
                    
                    if col_map:
                        self.df = self.df.rename(columns=col_map)
                        logger.info(f"根据位置推断列名: {col_map}")
            
            return True
        except Exception as e:
            logger.error(f"Excel数据加载失败: {str(e)}")
            return False
    
    def get_latest_date(self, dates_list):
        """获取最晚日期"""
        try:
            if not dates_list:
                logger.warning("日期列表为空，使用当前日期")
                return datetime.now()
                
            valid_dates = []
            for date in dates_list:
                try:
                    if pd.notna(date):
                        if isinstance(date, str):
                            try:
                                # 尝试解析各种日期格式
                                if '.' in date:  # 处理 2024.06.05 格式
                                    date = datetime.strptime(date, '%Y.%m.%d')
                                elif '-' in date:  # 处理 2024-06-05 格式
                                    date = datetime.strptime(date, '%Y-%m-%d')
                                elif '/' in date:  # 处理 2024/06/05 格式
                                    date = datetime.strptime(date, '%Y/%m/%d')
                                else:
                                    try:
                                        # 尝试通用解析
                                        date = pd.to_datetime(date).to_pydatetime()
                                    except:
                                        logger.warning(f"无法解析日期格式: {date}")
                                        continue
                            except ValueError as e:
                                try:
                                    # 尝试更宽松的解析
                                    date = pd.to_datetime(date).to_pydatetime()
                                except:
                                    logger.warning(f"无法解析日期格式: {date}")
                                    continue
                        elif isinstance(date, pd.Timestamp):
                            date = date.to_pydatetime()
                        elif isinstance(date, datetime):
                            pass  # 已经是datetime格式
                        else:
                            try:
                                date = pd.to_datetime(date).to_pydatetime()
                            except:
                                logger.warning(f"无法解析日期类型: {type(date)}")
                                continue
                                
                        valid_dates.append(date)
                except Exception as inner_e:
                    logger.warning(f"处理日期 {date} 时出错: {str(inner_e)}")
                    continue
            
            if valid_dates:
                return max(valid_dates)
            
            # 如果没有有效日期，使用当前日期
            logger.warning("未找到有效日期，使用当前日期")
            return datetime.now()
        except Exception as e:
            logger.error(f"获取最晚日期失败: {str(e)}")
            # 返回当前日期作为后备
            return datetime.now()
    
    def find_cell_with_text(self, table, search_text):
        """查找包含指定文本的单元格"""
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if search_text in cell.text:
                    return (i, j)
        return None
    
    def find_column_index_by_header(self, table, header_text):
        """根据表头文本查找列索引，支持模糊匹配"""
        if not table.rows or len(table.rows) == 0:
            return None
            
        # 先尝试精确匹配
        for i, row in enumerate(table.rows):
            if i == 0:  # 只在第一行(表头行)查找
                for j, cell in enumerate(row.cells):
                    if header_text in cell.text:
                        return j
        
        # 如果精确匹配失败，尝试部分匹配
        header_key = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', header_text)  # 移除特殊字符
        for i, row in enumerate(table.rows):
            if i == 0:  # 只在第一行(表头行)查找
                for j, cell in enumerate(row.cells):
                    cell_text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', cell.text)  # 移除特殊字符
                    if header_key in cell_text or cell_text in header_key:
                        logger.info(f"通过模糊匹配找到列: '{header_text}' -> '{cell.text}' (列索引 {j})")
                        return j
        
        return None
    
    def generate_reports(self):
        """生成所有报告"""
        if not self.excel_file or not self.word_template or not self.output_dir:
            logger.error("缺少必要的文件路径")
            return False
        
        if self.df is None:
            if not self.load_excel_data():
                return False
        
        # 获取委托日期列的最晚日期
        try:
            if '委托日期' in self.df.columns:
                # 安全地获取列数据
                date_values = []
                for value in self.df['委托日期'].values:
                    # 确保每个值都经过转换处理
                    try:
                        if pd.notna(value):
                            if isinstance(value, (str, datetime, pd.Timestamp)):
                                date_values.append(value)
                    except Exception:
                        pass
                
                latest_date = self.get_latest_date(date_values)
                logger.info(f"最晚委托日期: {latest_date}")
            else:
                latest_date = datetime.now()
                logger.warning(f"未找到委托日期列，使用当前日期: {latest_date}")
        except Exception as e:
            logger.error(f"获取最晚日期出错: {str(e)}")
            latest_date = datetime.now()
            logger.warning(f"使用当前日期作为备选: {latest_date}")
        
        # 创建一个新的Word文档副本
        try:
            doc = Document(self.word_template)
            logger.info(f"已加载Word模板: {self.word_template}")
            
            # 处理委托人日期 - 在表格外查找
            date_replaced = False
            for paragraph in doc.paragraphs:
                if "委托人" in paragraph.text:
                    # 将日期格式化为yyyy年mm月dd日
                    if isinstance(latest_date, datetime):
                        date_str = latest_date.strftime('%Y年%m月%d日')
                    else:
                        try:
                            if isinstance(latest_date, str) and '.' in latest_date:
                                date_obj = datetime.strptime(latest_date, '%Y.%m.%d')
                            else:
                                date_obj = pd.to_datetime(latest_date).to_pydatetime()
                            date_str = date_obj.strftime('%Y年%m月%d日')
                        except:
                            date_str = str(datetime.now().strftime('%Y年%m月%d日'))
                    
                    # 替换日期文本
                    paragraph.text = paragraph.text.replace("委托人", f"委托人 {date_str}")
                    date_replaced = True
                    break
            
            # 查找表格中的委托人字段
            if not date_replaced:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if "委托人" in cell.text:
                                for paragraph in cell.paragraphs:
                                    if "委托人" in paragraph.text:
                                        # 将日期格式化为yyyy年mm月dd日
                                        if isinstance(latest_date, datetime):
                                            date_str = latest_date.strftime('%Y年%m月%d日')
                                        else:
                                            try:
                                                if isinstance(latest_date, str) and '.' in latest_date:
                                                    date_obj = datetime.strptime(latest_date, '%Y.%m.%d')
                                                else:
                                                    date_obj = pd.to_datetime(latest_date).to_pydatetime()
                                                date_str = date_obj.strftime('%Y年%m月%d日')
                                            except:
                                                date_str = str(datetime.now().strftime('%Y年%m月%d日'))
                                        
                                        # 替换日期文本
                                        paragraph.text = paragraph.text.replace("委托人", f"委托人 {date_str}")
                                        date_replaced = True
                                        break
            
            # 处理每一行数据
            successful_count = 0
            total_count = len(self.df)
            
            for index, row in self.df.iterrows():
                row_data = row.to_dict()
                # 使用最晚日期替换行数据中的日期
                if latest_date:
                    row_data['委托日期'] = latest_date
                    
                logger.info(f"正在处理第{index+1}条记录...")
                
                # 提取数据
                try:
                    委托日期 = row_data.get('委托日期', None)
                    检件编号 = row_data.get('检件编号', '')
                    焊口编号 = row_data.get('焊口编号', '')
                    焊工号 = row_data.get('焊工号', '')
                    规格 = row_data.get('规格', '')
                    材质 = row_data.get('材质', '')
                    
                    logger.info(f"行 {index+1} 数据: 委托日期={委托日期}, 检件编号={检件编号}, 焊口编号={焊口编号}, 焊工号={焊工号}, 规格={规格}, 材质={材质}")
                except Exception as e:
                    logger.error(f"数据提取失败: {str(e)}")
                    continue
                
                # 查找并替换表格中的值
                for table in doc.tables:
                    # 查找各字段所在列索引（增加模糊匹配）
                    管道编号_col = self.find_column_index_by_header(table, "管道编号") or self.find_column_index_by_header(table, "检件编号")
                    焊口号_col = self.find_column_index_by_header(table, "焊口号") or self.find_column_index_by_header(table, "焊口编号")
                    焊工号_col = self.find_column_index_by_header(table, "焊工号") or self.find_column_index_by_header(table, "焊工")
                    焊口规格_col = self.find_column_index_by_header(table, "焊口规格") or self.find_column_index_by_header(table, "规格")
                    焊口材质_col = self.find_column_index_by_header(table, "焊口材质") or self.find_column_index_by_header(table, "材质")
                    
                    # 如果没有找到列，尝试通过位置猜测
                    if 管道编号_col is None and len(table.rows) > 0 and len(table.rows[0].cells) > 1:
                        管道编号_col = 1  # 假设第2列是管道编号
                        logger.info(f"未找到管道编号列，使用第2列")
                    if 焊口号_col is None and len(table.rows) > 0 and len(table.rows[0].cells) > 2:
                        焊口号_col = 2  # 假设第3列是焊口号
                        logger.info(f"未找到焊口号列，使用第3列")
                    if 焊工号_col is None and len(table.rows) > 0 and len(table.rows[0].cells) > 3:
                        焊工号_col = 3  # 假设第4列是焊工号
                        logger.info(f"未找到焊工号列，使用第4列")
                    if 焊口规格_col is None and len(table.rows) > 0 and len(table.rows[0].cells) > 4:
                        焊口规格_col = 4  # 假设第5列是焊口规格
                        logger.info(f"未找到焊口规格列，使用第5列")
                    if 焊口材质_col is None and len(table.rows) > 0 and len(table.rows[0].cells) > 5:
                        焊口材质_col = 5  # 假设第6列是焊口材质
                        logger.info(f"未找到焊口材质列，使用第6列")
                        
                    logger.info(f"表格列索引: 管道编号={管道编号_col}, 焊口号={焊口号_col}, 焊工号={焊工号_col}, 焊口规格={焊口规格_col}, 焊口材质={焊口材质_col}")
                    
                    # 查找表格中的数据行
                    # 首先检查表格是否有足够的行
                    header_row = -1
                    
                    # 尝试找到表头行
                    for i, row in enumerate(table.rows):
                        for cell in row.cells:
                            if "管道编号" in cell.text or "检件编号" in cell.text or "焊口号" in cell.text or "焊口编号" in cell.text:
                                header_row = i
                                break
                        if header_row >= 0:
                            break
                    
                    if header_row < 0:
                        header_row = 0  # 如果找不到表头行，假设第一行是表头
                        logger.warning("未找到表头行，假设第一行是表头")
                    
                    # 数据行从表头行之后开始
                    data_row = header_row + index + 1
                    
                    # 如果行数不够，添加新行
                    while data_row >= len(table.rows):
                        new_row = table.add_row()
                        # 确保新行有足够的单元格
                        if len(table.rows[0].cells) > len(new_row.cells):
                            for _ in range(len(table.rows[0].cells) - len(new_row.cells)):
                                new_row.add_cell()
                    
                    logger.info(f"数据将填入第{data_row+1}行")
                    
                    # 填充相应字段到找到的行
                    values_replaced = False
                    
                    # 清空行中的所有单元格，确保没有残留数据
                    if data_row < len(table.rows):
                        for j in range(len(table.rows[data_row].cells)):
                            # 保留第一列的序号（如果有）
                            if j > 0 or not table.rows[data_row].cells[j].text.strip().isdigit():
                                table.rows[data_row].cells[j].text = ""
                    
                    # 填充检件编号/管道编号
                    if 管道编号_col is not None and data_row < len(table.rows):
                        j = 管道编号_col
                        if j < len(table.rows[data_row].cells):
                            # 使用段落而不是直接设置文本，以保持格式
                            cell = table.rows[data_row].cells[j]
                            if not cell.paragraphs:
                                cell.add_paragraph()
                            cell.paragraphs[0].text = str(检件编号)
                            values_replaced = True
                            logger.info(f"已填充管道编号: {检件编号} 到 行{data_row+1}列{j+1}")
                    
                    # 填充焊口编号/焊口号
                    if 焊口号_col is not None and data_row < len(table.rows):
                        j = 焊口号_col
                        if j < len(table.rows[data_row].cells):
                            cell = table.rows[data_row].cells[j]
                            if not cell.paragraphs:
                                cell.add_paragraph()
                            cell.paragraphs[0].text = str(焊口编号)
                            values_replaced = True
                            logger.info(f"已填充焊口号: {焊口编号} 到 行{data_row+1}列{j+1}")
                    
                    # 填充焊工号
                    if 焊工号_col is not None and data_row < len(table.rows):
                        j = 焊工号_col
                        if j < len(table.rows[data_row].cells):
                            cell = table.rows[data_row].cells[j]
                            if not cell.paragraphs:
                                cell.add_paragraph()
                            cell.paragraphs[0].text = str(焊工号)
                            values_replaced = True
                            logger.info(f"已填充焊工号: {焊工号} 到 行{data_row+1}列{j+1}")
                    
                    # 填充焊口规格
                    if 焊口规格_col is not None and data_row < len(table.rows):
                        j = 焊口规格_col
                        if j < len(table.rows[data_row].cells):
                            cell = table.rows[data_row].cells[j]
                            if not cell.paragraphs:
                                cell.add_paragraph()
                            cell.paragraphs[0].text = str(规格)
                            values_replaced = True
                            logger.info(f"已填充规格: {规格} 到 行{data_row+1}列{j+1}")
                    
                    # 填充焊口材质
                    if 焊口材质_col is not None and data_row < len(table.rows):
                        j = 焊口材质_col
                        if j < len(table.rows[data_row].cells):
                            cell = table.rows[data_row].cells[j]
                            if not cell.paragraphs:
                                cell.add_paragraph()
                            cell.paragraphs[0].text = str(材质)
                            values_replaced = True
                            logger.info(f"已填充材质: {材质} 到 行{data_row+1}列{j+1}")
                    
                    # 无论是否找到匹配的列，都认为这条记录已处理
                    successful_count += 1
                    break  # 只处理第一个找到的表格
            
            # 检查表格内容是否已填充
            table_content_filled = False
            for table in doc.tables:
                if len(table.rows) > 1:  # 至少有表头和一行数据
                    for i in range(1, len(table.rows)):
                        row_has_data = False
                        for cell in table.rows[i].cells:
                            if cell.text.strip():
                                row_has_data = True
                                break
                        if row_has_data:
                            table_content_filled = True
                            break
                if table_content_filled:
                    break
            
            if not table_content_filled:
                logger.warning("警告：表格似乎没有填充任何数据！")
                # 打印表格结构以进行调试
                for t_idx, table in enumerate(doc.tables):
                    logger.info(f"表格 {t_idx+1} 结构:")
                    for r_idx, row in enumerate(table.rows):
                        row_text = []
                        for c_idx, cell in enumerate(row.cells):
                            row_text.append(f"[{c_idx}]:{cell.text}")
                        logger.info(f"  行 {r_idx+1}: {' | '.join(row_text)}")
            
            # 生成输出文件名
            output_filename = f"委托台账汇总.docx"
            output_path = os.path.join(self.output_dir, output_filename)
            
            # 保存文档
            doc.save(output_path)
            logger.info(f"成功生成文档: {output_filename}")
            
            # 判断是否真正成功
            if table_content_filled:
                logger.info(f"报告生成完成！共处理{total_count}条记录，成功填充数据到表格中")
                return True
            else:
                logger.error(f"报告生成失败！处理了{total_count}条记录，但未能成功填充到表格中")
                return False
            
        except Exception as e:
            logger.error(f"生成报告失败: {str(e)}")
            return False

# 与界面集成的函数
def generate_delegation_reports(excel_file, word_template, output_dir, log_callback=None):
    """生成射线检测委托台账报告"""
    # 创建日志处理程序，将日志输出到回调函数
    if log_callback:
        handler = logging.StreamHandler()
        handler.setLevel(logging.INFO)
        handler.setFormatter(logging.Formatter('%(message)s'))
        logger.addHandler(handler)
        
        # 覆盖默认的logger输出
        def log_output(message, level='INFO'):
            if level == 'INFO':
                logger.info(message)
            elif level == 'ERROR':
                logger.error(message)
            elif level == 'WARNING':
                logger.warning(message)
            log_callback(message)
    else:
        log_output = lambda msg, level='INFO': None
    
    try:
        # 创建报告生成器
        generator = ReportGenerator()
        
        # 设置文件路径
        generator.set_excel_file(excel_file)
        generator.set_word_template(word_template)
        generator.set_output_dir(output_dir)
        
        # 生成报告
        log_output("开始生成射线检测委托台账报告...", 'INFO')
        success = generator.generate_reports()
        
        if success:
            log_output("射线检测委托台账报告生成成功！", 'INFO')
            return True
        else:
            log_output("射线检测委托台账报告生成失败！", 'ERROR')
            return False
            
    except Exception as e:
        log_output(f"生成射线检测委托台账报告时发生错误: {str(e)}", 'ERROR')
        return False

# 如果直接运行脚本，则提供命令行界面
if __name__ == "__main__":
    # 创建一个简单的图形界面
    root = tk.Tk()
    root.withdraw()  # 隐藏主窗口
    
    # 选择Excel文件
    excel_file = filedialog.askopenfilename(title="选择Excel文件", filetypes=[("Excel文件", "*.xlsx;*.xls")])
    if not excel_file:
        messagebox.showerror("错误", "未选择Excel文件")
        exit()
    
    # 选择Word模板
    word_template = filedialog.askopenfilename(title="选择Word模板", filetypes=[("Word文件", "*.docx")])
    if not word_template:
        messagebox.showerror("错误", "未选择Word模板")
        exit()
    
    # 选择输出目录
    output_dir = filedialog.askdirectory(title="选择输出目录")
    if not output_dir:
        messagebox.showerror("错误", "未选择输出目录")
        exit()
    
    # 生成报告
    success = generate_delegation_reports(excel_file, word_template, output_dir)
    
    if success:
        messagebox.showinfo("成功", "报告生成成功！")
    else:
        messagebox.showerror("失败", "报告生成失败！") 