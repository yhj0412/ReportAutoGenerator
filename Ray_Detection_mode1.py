import pandas as pd
import os
import sys
import argparse
from docx import Document
from datetime import datetime
import re

def find_column_with_keyword(df, keyword):
    """查找包含指定关键字的列"""
    matching_cols = [col for col in df.columns if keyword.lower() in col.lower()]
    return matching_cols[0] if matching_cols else None

def get_output_filename(word_template_path, order_number):
    """根据Word模板路径和委托单编号生成输出文件名
    
    Args:
        word_template_path: Word模板文档路径
        order_number: 委托单编号
    
    Returns:
        str: 输出文件名
    """
    # 获取模板文件名（不含路径和扩展名）
    template_name = os.path.splitext(os.path.basename(word_template_path))[0]
    # 生成输出文件名
    return f"{template_name}_{order_number}_生成结果.docx"

def process_excel_to_word(excel_path, word_template_path, output_path=None,
                         project_name=None, client_name=None,
                         inspection_standard=None, acceptance_specification=None,
                         inspection_method=None, inspection_tech_level=None,
                         appearance_check=None, groove_type=None):
    """将Excel数据填入Word文档

    Args:
        excel_path: Excel表格路径
        word_template_path: Word模板文档路径
        output_path: 输出目录路径（如果为None，将自动生成）
        project_name: 工程名称，用于替换文档中的"工程名称值"
        client_name: 委托单位，用于替换文档中的"委托单位值"
        inspection_standard: 检测标准，用于替换文档中的"检测标准值"
        acceptance_specification: 验收规范，用于替换文档中的"验收规范值"
        inspection_method: 检测方法，用于替换文档中的"检测方法值"
        inspection_tech_level: 检测技术等级，用于替换文档中的"检测技术等级值"
        appearance_check: 外观检查，用于替换文档中的"外观检查值"
        groove_type: 坡口形式，用于替换文档中的"坡口形式值"

    Returns:
        bool: 处理是否成功
    """
    # 检查文件是否存在
    if not os.path.exists(excel_path):
        print(f"错误: Excel文件不存在: {excel_path}")
        return False
        
    if not os.path.exists(word_template_path):
        print(f"错误: Word模板文件不存在: {word_template_path}")
        return False
    
    # 获取模板文件名（不含路径和扩展名）作为子目录名
    template_name = os.path.splitext(os.path.basename(word_template_path))[0]
    
    # 创建输出目录
    output_dir = os.path.join("生成器", "输出报告","1_射线检测委托台账", template_name) if output_path is None else output_path
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 读取Excel数据
    print(f"正在读取Excel文件: {excel_path}")
    df = pd.read_excel(excel_path)
    
    # 打印所有列名，帮助调试
    print(f"Excel表格列名: {list(df.columns)}")
    
    # 定义需要查找的列关键字
    column_keywords = {
        '委托日期': '委托日期',           # A列
        '委托单编号': '委托单编号',       # C列
        '检件编号': '检件编号',           # D列
        '焊口编号': '焊口编号',           # E列
        '焊工号': '焊工号',              # F列
        '规格': '规格',                  # G列
        '材质': '材质',                  # H列
        '合格级别': '合格级别',           # I列
        '检测比例': '检测比例',           # J列
        '备注': '备注',                  # O列
        '单元名称': '单元名称',           # Q列
        '焊接方法': '焊接方法',           # R列
        '区号': '区号',                  # S列
        '单线号': '单线号',              # T列
        '检测时机': '检测时机'            # V列
    }
    
    # 查找每个关键字对应的实际列名
    column_mapping = {}
    missing_columns = []
    
    for key, keyword in column_keywords.items():
        col_name = find_column_with_keyword(df, keyword)
        if col_name:
            column_mapping[key] = col_name
            print(f"找到列: '{key}' -> '{col_name}'")
        else:
            missing_columns.append(key)
    
    if missing_columns:
        print(f"警告: 未找到以下列: {', '.join(missing_columns)}")
        # 尝试使用列位置
        possible_columns = {
            '委托日期': 'A', 
            '委托单编号': 'C', 
            '检件编号': 'D', 
            '焊口编号': 'E', 
            '焊工号': 'F', 
            '规格': 'G', 
            '材质': 'H', 
            '合格级别': 'I', 
            '检测比例': 'J', 
            '备注': 'O', 
            '单元名称': 'Q', 
            '焊接方法': 'R', 
            '区号': 'S', 
            '单线号': 'T', 
            '检测时机': 'V'
        }
        
        for key in missing_columns:
            col_letter = possible_columns.get(key)
            if col_letter:
                col_idx = ord(col_letter) - ord('A')
                if col_idx < len(df.columns):
                    column_mapping[key] = df.columns[col_idx]
                    print(f"使用列位置找到: '{key}' -> '{df.columns[col_idx]}'")
    
    # 检查必需的列是否都找到了
    required_columns = ['委托单编号', '委托日期', '检件编号']
    for col in required_columns:
        if col not in column_mapping:
            print(f"错误: 未找到必需的列: '{col}'")
            return False
    
    # 按委托单编号分组处理数据
    order_numbers = df[column_mapping['委托单编号']].dropna().unique().tolist()
    print(f"找到{len(order_numbers)}个不同的委托单编号")
    
    success_count = 0
    
    # 对每个委托单编号生成一份报告
    for order_number in order_numbers:
        print(f"\n处理委托单编号: {order_number}")
        
        # 筛选该委托单编号的数据
        order_df = df[df[column_mapping['委托单编号']] == order_number]
        print(f"该委托单编号有{len(order_df)}条记录")
        
        # 为该委托单编号生成输出文件名
        output_filename = get_output_filename(word_template_path, order_number)
        report_output_path = os.path.join(output_dir, output_filename)
        
        # 1) 获取该组数据中最晚的委托日期
        date_col = column_mapping.get('委托日期')
        if date_col:
            # 确保日期列是日期类型
            order_df[date_col] = pd.to_datetime(order_df[date_col], errors='coerce')
            latest_date = order_df[date_col].max()
            
            if pd.isna(latest_date):
                print(f"警告: 委托单编号 {order_number} 没有有效的委托日期")
                year, month, day = datetime.now().year, datetime.now().month, datetime.now().day
            else:
                # 将日期转换为年、月、日
                year = latest_date.year
                month = latest_date.month
                day = latest_date.day
                print(f"找到最晚委托日期: {year}年{month}月{day}日")
        else:
            print("警告: 未找到委托日期列")
            year, month, day = datetime.now().year, datetime.now().month, datetime.now().day
        
        # 获取该委托单编号下的相关数据
        pipe_codes = order_df[column_mapping.get('检件编号')].dropna().tolist() if '检件编号' in column_mapping else []
        weld_numbers = order_df[column_mapping.get('焊口编号')].dropna().tolist() if '焊口编号' in column_mapping else []
        welder_numbers = order_df[column_mapping.get('焊工号')].dropna().tolist() if '焊工号' in column_mapping else []
        specifications = order_df[column_mapping.get('规格')].dropna().tolist() if '规格' in column_mapping else []
        materials = order_df[column_mapping.get('材质')].dropna().tolist() if '材质' in column_mapping else []
        notes = order_df[column_mapping.get('备注')].tolist() if '备注' in column_mapping else []
        line_numbers = order_df[column_mapping.get('单线号')].dropna().tolist() if '单线号' in column_mapping else []
        
        # 获取单个值字段（每个委托单编号只需一个值）
        unit_name = ""
        if '单元名称' in column_mapping:
            unit_names = order_df[column_mapping['单元名称']].dropna().tolist()
            if unit_names:
                unit_name = unit_names[0]
                print(f"找到单元名称: {unit_name}")
        
        welding_method = ""
        if '焊接方法' in column_mapping:
            methods = order_df[column_mapping['焊接方法']].dropna().tolist()
            if methods:
                welding_method = methods[0]
                print(f"找到焊接方法: {welding_method}")
        
        area_number = ""
        if '区号' in column_mapping:
            areas = order_df[column_mapping['区号']].dropna().tolist()
            if areas:
                area_number = areas[0]
                print(f"找到区号: {area_number}")
        
        inspection_timing = ""
        if '检测时机' in column_mapping:
            timings = order_df[column_mapping['检测时机']].dropna().tolist()
            if timings:
                inspection_timing = timings[0]
                print(f"找到检测时机: {inspection_timing}")
        
        qualification_level = ""
        if '合格级别' in column_mapping:
            levels = order_df[column_mapping['合格级别']].dropna().tolist()
            if levels:
                qualification_level = levels[0]
                print(f"找到合格级别: {qualification_level}")
        
        inspection_ratio = ""
        if '检测比例' in column_mapping:
            ratios = order_df[column_mapping['检测比例']].dropna().tolist()
            if ratios:
                # 处理检测比例，确保以百分数格式显示
                ratio_value = ratios[0]
                # 检查是否已经是字符串形式的百分数
                if isinstance(ratio_value, str) and '%' in ratio_value:
                    inspection_ratio = ratio_value
                else:
                    # 尝试将值转换为浮点数，然后转为百分数格式
                    try:
                        ratio_float = float(ratio_value)
                        # 判断是否已经是0到100的值
                        if ratio_float > 1:
                            # 假设已经是百分数值，如50表示50%
                            inspection_ratio = f"{ratio_float:.0f}%"
                        else:
                            # 假设是小数，如0.5表示50%
                            inspection_ratio = f"{ratio_float*100:.0f}%"
                    except (ValueError, TypeError):
                        # 无法转换，直接使用原始值
                        inspection_ratio = str(ratio_value)
                print(f"找到检测比例: {inspection_ratio}")
        
        # 打开Word文档
        print(f"正在处理Word文档: {word_template_path}")
        doc = Document(word_template_path)
        
        # 替换文档中的参数值
        print("\n==== 开始替换参数值 ====")
        replacement_dict = {
            "工程名称值": project_name if project_name else "",
            "委托单位值": client_name if client_name else "",
            "检测标准值": inspection_standard if inspection_standard else "",
            "验收规范值": acceptance_specification if acceptance_specification else "",
            "检测方法值": inspection_method if inspection_method else "",
            "检测技术等级值": inspection_tech_level if inspection_tech_level else "",
            "外观检查值": appearance_check if appearance_check else "",
            "坡口形式值": groove_type if groove_type else "",
            "委托单编号值": str(order_number),
            "单元名称值": unit_name,
            "焊接方法值": welding_method,
            "区号值": area_number,
            "检测时机值": inspection_timing,
            "合格级别值": qualification_level,
            "检测比例值": inspection_ratio
        }
        
        # 遍历所有段落和表格中的单元格，替换参数值
        # 1. 遍历段落
        for paragraph in doc.paragraphs:
            for key, value in replacement_dict.items():
                if key in paragraph.text and value:
                    paragraph.text = paragraph.text.replace(key, value)
                    print(f"已将段落中的'{key}'替换为'{value}'")
        
        # 2. 遍历表格中的单元格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacement_dict.items():
                            if key in paragraph.text and value:
                                paragraph.text = paragraph.text.replace(key, value)
                                print(f"已将表格单元格中的'{key}'替换为'{value}'")
        
        print("==== 参数值替换完成 ====\n")
        
        # 处理表格 - 根据需求2，将委托日期的最晚日期填入到指定位置
        print("\n==== 开始处理日期填入 ====")
        date_keywords = ["施工单位", "监理单位", "项目部/装置", "检测单位"]

        for table in doc.tables:
            # 查找施工单位、监理单位、项目部/装置、检测单位的日期
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    for keyword in date_keywords:
                        if keyword in cell_text:
                            print(f"找到{keyword}单元格: 第{i+1}行, 第{j+1}列")
                            update_date_in_cell(cell, year, month, day)
            
            # 查找表头行，确定各列的位置
            column_indices = {}
            header_row_index = -1
            
            # 查找包含"管道编号"、"焊口号"、"焊工号"等的行
            for i, row in enumerate(table.rows):
                header_found = False
                for j, cell in enumerate(row.cells):
                    if "管道编号" in cell.text:
                        column_indices["管道编号"] = j
                        header_row_index = i
                        header_found = True
                    elif "焊口号" in cell.text:
                        column_indices["焊口号"] = j
                        header_found = True
                    elif "焊工号" in cell.text:
                        column_indices["焊工号"] = j
                        header_found = True
                    elif "焊口规格" in cell.text:
                        column_indices["焊口规格"] = j
                        header_found = True
                    elif "焊口材质" in cell.text:
                        column_indices["焊口材质"] = j
                        header_found = True
                    elif "备注" in cell.text:
                        column_indices["备注"] = j
                        header_found = True
                    elif "单线号" in cell.text:
                        column_indices["单线号"] = j
                        header_found = True
                
                if header_found and header_row_index >= 0:
                    break
            
            if header_row_index >= 0:
                print(f"找到表头行: 第{header_row_index+1}行")
                print(f"列索引: {column_indices}")
                
                # 如果找到表头行，处理数据填充
                # 获取可用于填充数据的行
                data_rows = []
                for i in range(header_row_index + 1, len(table.rows)):
                    if i < len(table.rows):
                        # 检查是否是空行或包含特殊标记的行
                        if len(table.rows[i].cells) > 0 and "以下空白" in table.rows[i].cells[0].text:
                            print(f"找到'以下空白'行: 第{i+1}行")
                            break
                        # 添加可用于填充数据的行
                        data_rows.append(i)
                
                print(f"找到{len(data_rows)}行可用于填充数据")
                
                # 确定需要填充的数据行数
                data_fields = [pipe_codes, weld_numbers, welder_numbers, specifications, materials, line_numbers]
                non_empty_fields = [field for field in data_fields if field]
                data_count = min([len(field) for field in non_empty_fields]) if non_empty_fields else 0
                
                print(f"需要填充{data_count}行数据")
                
                # 如果Word表格中的行数不足，需要添加新行
                rows_needed = data_count - len(data_rows)
                if rows_needed > 0:
                    print(f"需要添加{rows_needed}行到表格中")
                    # 找到最后一行的索引
                    last_row_idx = data_rows[-1] if data_rows else header_row_index
                    
                    # 添加新行
                    for _ in range(rows_needed):
                        # 在表格末尾添加一行
                        new_row = table.add_row()
                        data_rows.append(len(table.rows) - 1)  # 添加新行的索引
                
                # 处理每一行数据
                for i in range(data_count):
                    if i < len(data_rows):
                        row_idx = data_rows[i]
                        row = table.rows[row_idx]
                        
                        # 1. 填写管道编号（检件编号）
                        if "管道编号" in column_indices and i < len(pipe_codes):
                            col_idx = column_indices["管道编号"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                if cell.paragraphs:
                                    cell.paragraphs[0].text = str(pipe_codes[i])
                                    print(f"已更新第{row_idx+1}行管道编号: {pipe_codes[i]}")
                        
                        # 2. 填写焊口号
                        if "焊口号" in column_indices and i < len(weld_numbers):
                            col_idx = column_indices["焊口号"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                if cell.paragraphs:
                                    cell.paragraphs[0].text = str(weld_numbers[i])
                                    print(f"已更新第{row_idx+1}行焊口号: {weld_numbers[i]}")
                        
                        # 3. 填写焊工号
                        if "焊工号" in column_indices and i < len(welder_numbers):
                            col_idx = column_indices["焊工号"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                if cell.paragraphs:
                                    cell.paragraphs[0].text = str(welder_numbers[i])
                                    print(f"已更新第{row_idx+1}行焊工号: {welder_numbers[i]}")
                        
                        # 4. 填写焊口规格
                        if "焊口规格" in column_indices and i < len(specifications):
                            col_idx = column_indices["焊口规格"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                if cell.paragraphs:
                                    cell.paragraphs[0].text = str(specifications[i])
                                    print(f"已更新第{row_idx+1}行焊口规格: {specifications[i]}")
                        
                        # 5. 填写焊口材质
                        if "焊口材质" in column_indices and i < len(materials):
                            col_idx = column_indices["焊口材质"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                if cell.paragraphs:
                                    cell.paragraphs[0].text = str(materials[i])
                                    print(f"已更新第{row_idx+1}行焊口材质: {materials[i]}")
                        
                        # 6. 填写备注
                        if "备注" in column_indices and i < len(notes):
                            col_idx = column_indices["备注"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                if cell.paragraphs:
                                    note = notes[i]
                                    if pd.notna(note):  # 检查是否为NaN
                                        cell.paragraphs[0].text = str(note)
                                        print(f"已更新第{row_idx+1}行备注")
                        
                        # 7. 填写单线号
                        if "单线号" in column_indices and i < len(line_numbers):
                            col_idx = column_indices["单线号"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                if cell.paragraphs:
                                    cell.paragraphs[0].text = str(line_numbers[i])
                                    print(f"已更新第{row_idx+1}行单线号: {line_numbers[i]}")
        
        # 保存文档
        doc.save(report_output_path)
        print(f"文档已保存至: {report_output_path}")
        success_count += 1
    
    print(f"\n处理完成: 共处理{len(order_numbers)}个委托单编号，成功生成{success_count}份报告")
    return success_count > 0

def update_date_in_cell(cell, year, month, day):
    """更新单元格中的日期"""
    # 检查单元格中的所有段落
    date_found = False
    for paragraph in cell.paragraphs:
        if "年" in paragraph.text and "月" in paragraph.text and "日" in paragraph.text:
            print(f"找到日期段落: {paragraph.text}")
            
            # 创建新的文本，确保只有一个年月日
            new_text = paragraph.text
            # 确保年月日前没有数字
            new_text = re.sub(r'\d*年', '年', new_text)
            new_text = re.sub(r'\d*月', '月', new_text)
            new_text = re.sub(r'\d*日', '日', new_text)
            
            # 在年月日前插入正确的数字
            new_text = new_text.replace('年', f'{year}年')
            new_text = new_text.replace('月', f'{month}月')
            new_text = new_text.replace('日', f'{day}日')
            
            paragraph.text = new_text
            date_found = True
            print("已更新日期")
            break
    
    # 如果没有找到日期段落，尝试创建新段落
    if not date_found:
        print("未找到日期段落，尝试添加新日期...")
        p = cell.add_paragraph(f"{year}年{month}月{day}日")
        print("已添加日期")

def main():
    # 创建命令行参数解析器
    parser = argparse.ArgumentParser(description='将Excel数据填入Word文档')
    parser.add_argument('-e', '--excel', default="生成器/Excel/1_生成器委托.xlsx", 
                       help='Excel表格路径 (默认: 生成器/Excel/1_生成器委托.xlsx)')
    parser.add_argument('-w', '--word', default="生成器/wod/1_射线检测委托台账_Mode1.docx", 
                       help='Word模板文档路径 (默认: 生成器/wod/1_射线检测委托台账_Mode1.docx)')
    parser.add_argument('-o', '--output', 
                       help='输出目录路径 (可选，默认为"生成器/输出报告/1_射线检测委托台账/1_射线检测委托台账_Mode1"目录)')
    parser.add_argument('-p', '--project',
                       help='工程名称，用于替换文档中的"工程名称值"')
    parser.add_argument('-c', '--client',
                       help='委托单位，用于替换文档中的"委托单位值"')
    parser.add_argument('-s', '--standard',
                       help='检测标准，用于替换文档中的"检测标准值"')
    parser.add_argument('-a', '--acceptance',
                       help='验收规范，用于替换文档中的"验收规范值"')
    parser.add_argument('-m', '--method',
                       help='检测方法，用于替换文档中的"检测方法值"')
    parser.add_argument('-t', '--tech-level',
                       help='检测技术等级，用于替换文档中的"检测技术等级值"')
    parser.add_argument('-v', '--appearance',
                       help='外观检查，用于替换文档中的"外观检查值"')
    parser.add_argument('-g', '--groove',
                       help='坡口形式，用于替换文档中的"坡口形式值"')

    # 解析命令行参数
    args = parser.parse_args()

    # 处理Excel到Word的转换
    success = process_excel_to_word(
        args.excel, args.word, args.output,
        args.project, args.client, args.standard, args.acceptance,
        args.method, getattr(args, 'tech_level'), args.appearance, args.groove
    )
    
    # 返回状态码
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()