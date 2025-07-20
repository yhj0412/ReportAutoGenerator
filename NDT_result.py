import pandas as pd
import os
import sys
import argparse
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re

def find_column_with_keyword(df, keyword):
    """查找包含指定关键字的列"""
    matching_cols = [col for col in df.columns if keyword.lower() in col.lower()]
    return matching_cols[0] if matching_cols else None

def replace_text_in_paragraph(paragraph, old_text, new_text):
    """在段落中精确替换文本，只对替换的部分设置楷体五号字体

    Args:
        paragraph: Word段落对象
        old_text: 要替换的文本
        new_text: 新文本
    """
    # 获取段落的完整文本
    full_text = paragraph.text

    if old_text in full_text:
        # 清空段落
        paragraph.clear()

        # 分割文本
        parts = full_text.split(old_text)

        # 重新构建段落
        for i, part in enumerate(parts):
            if i > 0:
                # 添加替换的文本（设置楷体五号字体）
                run = paragraph.add_run(new_text)
                run.font.name = "楷体"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                run.font.size = Pt(10.5)

            if part:
                # 添加原始文本部分（保持原有格式）
                paragraph.add_run(part)

def set_cell_center_alignment(cell):
    """设置单元格文本居中对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def get_detection_level_by_method(detection_method):
    """根据检测方法获取对应的检测级别值

    Args:
        detection_method: 检测方法字符串

    Returns:
        str: 对应的检测级别值
    """
    if not detection_method:
        return ""

    # 转换为字符串并去除空格，转为大写进行匹配
    method = str(detection_method).strip().upper()

    # 检测方法与检测级别值的映射关系
    method_mapping = {
        # 硬度检测或YD
        '硬度检测': '力学   级',
        'YD': '力学   级',
        # 光谱检测或PMIN
        '光谱检测': '光谱分析  级',
        'PMIN': '光谱分析  级',
        # 其他检测方法
        'UT': 'UT  级',
        'PT': 'PT  级',
        'MT': 'MT  级',
        'RT': 'RT  级',
        'TOFD': 'TOFD  级',
        'PA': 'PA  级'
    }

    # 精确匹配
    if method in method_mapping:
        return method_mapping[method]

    # 模糊匹配 - 检查是否包含关键字
    for key, value in method_mapping.items():
        if key in method:
            return value

    # 如果没有匹配到，返回空字符串
    print(f"警告: 未找到检测方法 '{detection_method}' 对应的检测级别值")
    return ""

def get_output_filename(word_template_path, order_number):
    """根据Word模板路径和委托单编号生成输出文件名
    
    Args:
        word_template_path: Word模板文档路径
        order_number: 委托单编号
    
    Returns:
        str: 输出文件名
    """
    # 获取模板文件名（不含路径和扩展名）
    template_name = os.path.splitext(os.path.basename(word_template_path))[0]
    # 生成输出文件名
    return f"{template_name}_{order_number}_生成结果.docx"

def process_excel_to_word(excel_path, word_template_path, output_path=None, project_name=None, client_name=None, inspection_method=None):
    """将Excel数据填入Word文档
    
    Args:
        excel_path: Excel表格路径
        word_template_path: Word模板文档路径
        output_path: 输出Word文档路径（如果为None，将自动生成）
        project_name: 工程名称，用于替换文档中的"工程名称参数值"
        client_name: 委托单位，用于替换文档中的"委托单位参数值"
        inspection_method: 检测方法，用于替换文档中的"检测方法参数"
    
    Returns:
        bool: 处理是否成功
    """
    # 检查文件是否存在
    if not os.path.exists(excel_path):
        print(f"错误: Excel文件不存在: {excel_path}")
        return False
        
    if not os.path.exists(word_template_path):
        print(f"错误: Word模板文件不存在: {word_template_path}")
        return False
    
    # 创建输出目录
    if output_path:
        # 使用GUI传递的输出路径
        output_dir = output_path
        print(f"使用GUI指定的输出路径: {output_dir}")
    else:
        # 使用默认输出路径
        output_dir = os.path.join("生成器", "输出报告", "2_RT结果通知单台账", "2_RT结果通知单台账_Mode2")
        print(f"使用默认输出路径: {output_dir}")

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 读取Excel数据
    print(f"正在读取Excel文件: {excel_path}")
    df = pd.read_excel(excel_path)
    
    # 打印所有列名，帮助调试
    print(f"Excel表格列名: {list(df.columns)}")
    
    # 定义需要查找的列关键字
    column_keywords = {
        '完成日期': '完成日期',
        '委托单编号': '委托单编号',
        '检件编号': '检件编号',
        '焊口编号': '焊口编号',
        '焊工号': '焊工号',
        '返修补片': '返修补片',
        '实际不合格': '实际不合格',
        '备注': '备注',
        '单元名称': '单元名称'
    }
    
    # 查找每个关键字对应的实际列名
    column_mapping = {}
    missing_columns = []
    
    for key, keyword in column_keywords.items():
        col_name = find_column_with_keyword(df, keyword)
        if col_name:
            column_mapping[key] = col_name
            print(f"找到列: '{key}' -> '{col_name}'")
        else:
            missing_columns.append(key)
    
    if missing_columns:
        print(f"警告: 未找到以下列: {', '.join(missing_columns)}")
        # 尝试使用列位置
        possible_columns = {
            '完成日期': 'B', 
            '委托单编号': 'C', 
            '检件编号': 'D', 
            '焊口编号': 'E', 
            '焊工号': 'F', 
            '返修补片': 'K', 
            '实际不合格': 'W', 
            '备注': 'O', 
            '单元名称': 'Q'
        }
        
        for key in missing_columns:
            col_letter = possible_columns.get(key)
            if col_letter:
                col_idx = ord(col_letter) - ord('A')
                if col_idx < len(df.columns):
                    column_mapping[key] = df.columns[col_idx]
                    print(f"使用列位置找到: '{key}' -> '{df.columns[col_idx]}'")
    
    # 按委托单编号分组处理数据
    if '委托单编号' not in column_mapping:
        print("错误: 无法找到委托单编号列")
        return False
    
    # 获取所有唯一的委托单编号
    order_numbers = df[column_mapping['委托单编号']].dropna().unique().tolist()
    print(f"找到{len(order_numbers)}个不同的委托单编号")
    
    success_count = 0
    
    # 对每个委托单编号生成一份报告
    for order_number in order_numbers:
        print(f"\n处理委托单编号: {order_number}")
        
        # 筛选该委托单编号的数据
        order_df = df[df[column_mapping['委托单编号']] == order_number]
        print(f"该委托单编号有{len(order_df)}条记录")
        
        # 为该委托单编号生成输出文件名
        output_filename = get_output_filename(word_template_path, order_number)
        report_output_path = os.path.join(output_dir, output_filename)
        
        # 1) 获取该组数据中最晚的完成日期
        date_col = column_mapping.get('完成日期')
        if date_col:
            # 确保日期列是日期类型
            order_df[date_col] = pd.to_datetime(order_df[date_col], errors='coerce')
            latest_date = order_df[date_col].max()
            
            if pd.isna(latest_date):
                print(f"警告: 委托单编号 {order_number} 没有有效的完成日期")
                year, month, day = datetime.now().year, datetime.now().month, datetime.now().day
            else:
                # 将日期转换为年、月、日
                year = latest_date.year
                month = latest_date.month
                day = latest_date.day
                print(f"找到最晚完成日期: {year}年{month}月{day}日")
        else:
            print("警告: 未找到完成日期列")
            year, month, day = datetime.now().year, datetime.now().month, datetime.now().day
        
        # 获取相关数据
        inspection_numbers = order_df[column_mapping.get('检件编号')].dropna().tolist() if '检件编号' in column_mapping else []
        weld_numbers = order_df[column_mapping.get('焊口编号')].dropna().tolist() if '焊口编号' in column_mapping else []
        welder_numbers = order_df[column_mapping.get('焊工号')].dropna().tolist() if '焊工号' in column_mapping else []
        repair_results = order_df[column_mapping.get('返修补片')].tolist() if '返修补片' in column_mapping else []
        failure_counts = order_df[column_mapping.get('实际不合格')].tolist() if '实际不合格' in column_mapping else []
        notes = order_df[column_mapping.get('备注')].tolist() if '备注' in column_mapping else []
        
        # 获取单元名称（第一个非空值）
        unit_name = ""
        if '单元名称' in column_mapping:
            unit_names = order_df[column_mapping['单元名称']].dropna().tolist()
            if unit_names:
                unit_name = unit_names[0]
                print(f"找到单元名称: {unit_name}")
        
        # 打开Word文档
        print(f"正在处理Word文档: {word_template_path}")
        
        # 检查文件扩展名，使用不同的方法处理.doc和.docx文件
        if word_template_path.lower().endswith('.doc'):
            # 对于.doc文件，需要先转换为.docx
            temp_docx_path = word_template_path + 'x'
            print(f"检测到.doc文件，尝试转换为.docx: {temp_docx_path}")
            
            try:
                # 尝试直接打开.doc文件
                doc = Document(word_template_path)
                doc.save(temp_docx_path)
                print(f"成功转换.doc为.docx")
                doc = Document(temp_docx_path)
            except Exception as e:
                print(f"无法直接打开.doc文件: {e}")
                print("请将.doc文件转换为.docx格式后重试")
                continue
        else:
            # 对于.docx文件，直接打开
            doc = Document(word_template_path)
        
        # 替换文档中的参数值
        if project_name or client_name or inspection_method:
            print("\n==== 开始替换参数值 ====")

            # 根据传入的检测方法参数计算检测级别值
            detection_level = ""
            if inspection_method:
                detection_level = get_detection_level_by_method(inspection_method)
                if detection_level:
                    print(f"根据传参检测方法 '{inspection_method}' 确定检测级别值: '{detection_level}'")

            # 遍历所有段落和表格中的单元格，替换参数值
            # 1. 遍历段落
            for paragraph in doc.paragraphs:
                if project_name and "工程名称参数值" in paragraph.text:
                    # 保存原始文本
                    original_text = paragraph.text
                    # 只有当段落文本完全等于占位符时，才替换整个段落
                    if original_text.strip() == "工程名称参数值":
                        # 清空段落内容并重新添加
                        paragraph.clear()
                        run = paragraph.add_run(project_name)
                        # 设置楷体五号字体
                        run.font.name = "楷体"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                        run.font.size = Pt(10.5)
                        print(f"已将段落中的'工程名称参数值'替换为'{project_name}'并设置为楷体五号字体")
                    else:
                        # 如果段落包含其他内容，需要精确替换
                        replace_text_in_paragraph(paragraph, "工程名称参数值", project_name)
                        print(f"已将段落中的'工程名称参数值'替换为'{project_name}'并设置为楷体五号字体")

                if client_name and "委托单位参数值" in paragraph.text:
                    # 保存原始文本
                    original_text = paragraph.text
                    # 只有当段落文本完全等于占位符时，才替换整个段落
                    if original_text.strip() == "委托单位参数值":
                        # 清空段落内容并重新添加
                        paragraph.clear()
                        run = paragraph.add_run(client_name)
                        # 设置楷体五号字体
                        run.font.name = "楷体"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                        run.font.size = Pt(10.5)
                        print(f"已将段落中的'委托单位参数值'替换为'{client_name}'并设置为楷体五号字体")
                    else:
                        # 如果段落包含其他内容，需要精确替换
                        replace_text_in_paragraph(paragraph, "委托单位参数值", client_name)
                        print(f"已将段落中的'委托单位参数值'替换为'{client_name}'并设置为楷体五号字体")

                if inspection_method and "检测方法参数" in paragraph.text:
                    # 保存原始文本
                    original_text = paragraph.text
                    # 只有当段落文本完全等于占位符时，才替换整个段落
                    if original_text.strip() == "检测方法参数":
                        # 清空段落内容并重新添加
                        paragraph.clear()
                        run = paragraph.add_run(inspection_method)
                        # 设置楷体五号字体
                        run.font.name = "楷体"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                        run.font.size = Pt(10.5)
                        print(f"已将段落中的'检测方法参数'替换为'{inspection_method}'并设置为楷体五号字体")
                    else:
                        # 如果段落包含其他内容，需要精确替换
                        replace_text_in_paragraph(paragraph, "检测方法参数", inspection_method)
                        print(f"已将段落中的'检测方法参数'替换为'{inspection_method}'并设置为楷体五号字体")

                if detection_level and "检测级别值" in paragraph.text:
                    # 保存原始文本
                    original_text = paragraph.text
                    # 只有当段落文本完全等于占位符时，才替换整个段落
                    if original_text.strip() == "检测级别值":
                        # 清空段落内容并重新添加
                        paragraph.clear()
                        run = paragraph.add_run(detection_level)
                        # 设置楷体五号字体
                        run.font.name = "楷体"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                        run.font.size = Pt(10.5)
                        print(f"已将段落中的'检测级别值'替换为'{detection_level}'并设置为楷体五号字体")
                    else:
                        # 如果段落包含其他内容，需要精确替换
                        replace_text_in_paragraph(paragraph, "检测级别值", detection_level)
                        print(f"已将段落中的'检测级别值'替换为'{detection_level}'并设置为楷体五号字体")
            
            # 2. 遍历表格中的单元格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if project_name and "工程名称参数值" in paragraph.text:
                                # 保存原始文本
                                original_text = paragraph.text
                                # 只有当段落文本完全等于占位符时，才替换整个段落
                                if original_text.strip() == "工程名称参数值":
                                    # 清空段落内容并重新添加
                                    paragraph.clear()
                                    run = paragraph.add_run(project_name)
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已将表格单元格中的'工程名称参数值'替换为'{project_name}'并设置为楷体五号字体")
                                else:
                                    # 如果段落包含其他内容，需要精确替换
                                    replace_text_in_paragraph(paragraph, "工程名称参数值", project_name)
                                    print(f"已将表格单元格中的'工程名称参数值'替换为'{project_name}'并设置为楷体五号字体")

                            if client_name and "委托单位参数值" in paragraph.text:
                                # 保存原始文本
                                original_text = paragraph.text
                                # 只有当段落文本完全等于占位符时，才替换整个段落
                                if original_text.strip() == "委托单位参数值":
                                    # 清空段落内容并重新添加
                                    paragraph.clear()
                                    run = paragraph.add_run(client_name)
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已将表格单元格中的'委托单位参数值'替换为'{client_name}'并设置为楷体五号字体")
                                else:
                                    # 如果段落包含其他内容，需要精确替换
                                    replace_text_in_paragraph(paragraph, "委托单位参数值", client_name)
                                    print(f"已将表格单元格中的'委托单位参数值'替换为'{client_name}'并设置为楷体五号字体")

                            if inspection_method and "检测方法参数" in paragraph.text:
                                # 保存原始文本
                                original_text = paragraph.text
                                # 只有当段落文本完全等于占位符时，才替换整个段落
                                if original_text.strip() == "检测方法参数":
                                    # 清空段落内容并重新添加
                                    paragraph.clear()
                                    run = paragraph.add_run(inspection_method)
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已将表格单元格中的'检测方法参数'替换为'{inspection_method}'并设置为楷体五号字体")
                                else:
                                    # 如果段落包含其他内容，需要精确替换
                                    replace_text_in_paragraph(paragraph, "检测方法参数", inspection_method)
                                    print(f"已将表格单元格中的'检测方法参数'替换为'{inspection_method}'并设置为楷体五号字体")

                            if detection_level and "检测级别值" in paragraph.text:
                                # 保存原始文本
                                original_text = paragraph.text
                                # 只有当段落文本完全等于占位符时，才替换整个段落
                                if original_text.strip() == "检测级别值":
                                    # 清空段落内容并重新添加
                                    paragraph.clear()
                                    run = paragraph.add_run(detection_level)
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已将表格单元格中的'检测级别值'替换为'{detection_level}'并设置为楷体五号字体")
                                else:
                                    # 如果段落包含其他内容，需要精确替换
                                    replace_text_in_paragraph(paragraph, "检测级别值", detection_level)
                                    print(f"已将表格单元格中的'检测级别值'替换为'{detection_level}'并设置为楷体五号字体")
            
            print("==== 参数值替换完成 ====\n")
        
        # 填写通知单编号（委托单编号）
        notification_number_updated = False
        
        print("\n==== 开始查找通知单编号位置 ====")
        
        # 打印表格结构以便调试
        for table_idx, table in enumerate(doc.tables):
            print(f"检查表格 #{table_idx+1}，共有 {len(table.rows)} 行")
            
            # 打印前5行的内容，帮助理解表格结构
            for i in range(min(5, len(table.rows))):
                if i < len(table.rows):
                    row = table.rows[i]
                    row_content = []
                    for cell in row.cells:
                        # 截断长文本，只显示前20个字符
                        cell_text = cell.text.strip()
                        if len(cell_text) > 20:
                            cell_text = cell_text[:20] + "..."
                        row_content.append(cell_text)
                    print(f"  第{i+1}行内容: {row_content}")
        
        # 特别查找可能包含通知单编号的单元格
        print("\n查找可能包含通知单编号的单元格:")
        for table_idx, table in enumerate(doc.tables):
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    # 检查是否包含特定关键字
                    if (("RX" in cell_text) or ("RT" in cell_text) or 
                        ("-DG-" in cell_text) or ("*" in cell_text) or
                        ("通知单" in cell_text)):
                        print(f"  表格#{table_idx+1}, 第{i+1}行, 第{j+1}列: '{cell_text}'")
        
        # 直接查找包含特定格式的单元格(如"RX3-03-ZYLJ-DG-RT-000*")
        for table_idx, table in enumerate(doc.tables):
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    
                    # 检查单元格内容是否符合特定格式
                    if (cell_text and 
                        (("RX" in cell_text and "-DG-RT-" in cell_text) or 
                         ("RT" in cell_text and "-DG-" in cell_text and "*" in cell_text) or
                         ("RX3-03-ZYLJ-DG-RT" in cell_text))):
                        
                        print(f"找到匹配特定格式的单元格: 表格#{table_idx+1}, 第{i+1}行, 第{j+1}列")
                        print(f"单元格内容: '{cell_text}'")
                        print(f"将替换为委托单编号: {order_number}")
                        
                        # 保存原始内容以便验证
                        original_content = cell_text
                        
                        # 修改单元格内容
                        if cell.paragraphs:
                            paragraph = cell.paragraphs[0]
                            paragraph.clear()
                            run = paragraph.add_run(str(order_number))
                            # 设置楷体五号字体
                            run.font.name = "楷体"
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                            run.font.size = Pt(10.5)
                            print(f"已将单元格内容从 '{original_content}' 修改为 '{order_number}'并设置为楷体五号字体")
                            notification_number_updated = True
                            break
                if notification_number_updated:
                    break
            if notification_number_updated:
                break
        
        # 如果未找到匹配的单元格，尝试查找表格右上角区域
        if not notification_number_updated:
            print("\n未找到完全匹配的单元格，检查表格右上角区域...")
            for table_idx, table in enumerate(doc.tables):
                # 只检查前3行
                for i in range(min(3, len(table.rows))):
                    if i < len(table.rows) and len(table.rows[i].cells) > 0:
                        # 检查行中的最后一个单元格
                        last_cell = table.rows[i].cells[-1]
                        cell_text = last_cell.text.strip()
                        
                        print(f"检查表格#{table_idx+1}, 第{i+1}行, 最后一列, 内容: '{cell_text}'")
                        
                        # 检查是否包含部分匹配特征
                        if (cell_text and 
                            (("RX" in cell_text) or ("RT" in cell_text) or 
                             ("DG" in cell_text) or ("*" in cell_text))):
                            
                            print(f"找到部分匹配的单元格: 表格#{table_idx+1}, 第{i+1}行, 最后一列")
                            print(f"单元格内容: '{cell_text}'")
                            print(f"将替换为委托单编号: {order_number}")
                            
                            # 保存原始内容以便验证
                            original_content = cell_text
                            
                            # 修改单元格内容
                            if last_cell.paragraphs:
                                paragraph = last_cell.paragraphs[0]
                                paragraph.clear()
                                run = paragraph.add_run(str(order_number))
                                # 设置楷体五号字体
                                run.font.name = "楷体"
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                run.font.size = Pt(10.5)
                                print(f"已将单元格内容从 '{original_content}' 修改为 '{order_number}'并设置为楷体五号字体")
                                notification_number_updated = True
                                break
                    if notification_number_updated:
                        break
                if notification_number_updated:
                    break
        
        # 打印查找结果总结
        print("\n==== 通知单编号位置查找结果 ====")
        if notification_number_updated:
            print(f"成功找到并替换通知单编号位置: {order_number}")
        else:
            print("警告: 未能找到合适的位置填写通知单编号")
        
        # 填写单位工程名称
        for paragraph in doc.paragraphs:
            if "单位工程名称" in paragraph.text:
                # 找到包含"单位工程名称"的段落
                print(f"找到单位工程名称段落: {paragraph.text}")
                
                # 检查是否在表格中
                found_in_table = False
                for table in doc.tables:
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if "单位工程名称" in cell.text and j + 1 < len(row.cells):
                                # 在右侧单元格填写单元名称
                                right_cell = row.cells[j + 1]
                                if right_cell.paragraphs and unit_name:
                                    paragraph = right_cell.paragraphs[0]
                                    paragraph.clear()
                                    run = paragraph.add_run(unit_name)
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已将单元名称 {unit_name} 填入单位工程名称右侧单元格并设置为楷体五号字体")
                                    found_in_table = True
                                    break
                        if found_in_table:
                            break
                    if found_in_table:
                        break
                
                if not found_in_table and unit_name:
                    # 如果不在表格中，尝试修改段落文本
                    new_text = paragraph.text
                    if "：" in new_text or ":" in new_text:
                        # 如果有冒号，在冒号后添加单元名称
                        if "：" in new_text:
                            new_text = new_text.split("：")[0] + "：" + unit_name
                        else:
                            new_text = new_text.split(":")[0] + ":" + unit_name
                    else:
                        # 否则直接在文本后添加单元名称
                        new_text = new_text + " " + unit_name
                    
                    # 使用精确替换来保持格式
                    if "：" in paragraph.text:
                        replace_text_in_paragraph(paragraph, "：", f"：{unit_name}")
                    elif ":" in paragraph.text:
                        replace_text_in_paragraph(paragraph, ":", f":{unit_name}")
                    else:
                        # 如果没有冒号，直接添加
                        paragraph.clear()
                        run = paragraph.add_run(new_text)
                        run.font.name = "楷体"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                        run.font.size = Pt(10.5)
                    print(f"已将单元名称 {unit_name} 添加到单位工程名称段落并设置为楷体五号字体")
        
        # 处理表格
        for table in doc.tables:
            # 查找日期字段
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    # 1) 处理"检测人"日期
                    if "检测人" in cell.text:
                        print(f"找到检测人单元格: 第{i+1}行, 第{j+1}列")
                        
                        # 检查单元格中的所有段落
                        date_found = False
                        for paragraph in cell.paragraphs:
                            if "年" in paragraph.text and "月" in paragraph.text and "日" in paragraph.text:
                                print(f"找到日期段落: {paragraph.text}")

                                original_text = paragraph.text
                                new_date = f'{year}年{month}月{day}日'

                                # 查找日期部分的模式 - 匹配各种日期格式
                                date_pattern = r'(\d+年\d+月\d+日|年\s*月\s*日\.?|年月日\.?)'

                                if re.search(date_pattern, original_text):
                                    # 清空段落
                                    paragraph.clear()

                                    # 使用正则表达式替换，同时保持格式
                                    current_pos = 0

                                    for match in re.finditer(date_pattern, original_text):
                                        # 添加匹配前的文本（标签部分），保持原有格式
                                        before_text = original_text[current_pos:match.start()]
                                        if before_text:
                                            paragraph.add_run(before_text)

                                        # 添加日期部分，设置楷体五号字体
                                        date_run = paragraph.add_run(new_date)
                                        date_run.font.name = "楷体"
                                        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                        date_run.font.size = Pt(10.5)

                                        current_pos = match.end()

                                    # 添加剩余的文本（如果有的话）
                                    remaining_text = original_text[current_pos:]
                                    if remaining_text:
                                        paragraph.add_run(remaining_text)

                                    date_found = True
                                    print("已更新检测人日期并设置为楷体五号字体")
                                    break

                        # 如果没有找到日期段落，尝试创建新段落
                        if not date_found:
                            print("未在检测人单元格中找到日期段落，尝试其他方法...")
                            # 添加新段落
                            p = cell.add_paragraph()
                            run = p.add_run(f"{year}年{month}月{day}日")
                            # 设置楷体五号字体
                            run.font.name = "楷体"
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                            run.font.size = Pt(10.5)
                            print("已添加检测人日期并设置为楷体五号字体")
                    
                    # 2) 处理"审核"日期
                    if "审核" in cell.text:
                        print(f"找到审核单元格: 第{i+1}行, 第{j+1}列")
                        
                        # 检查单元格中的所有段落
                        date_found = False
                        for paragraph in cell.paragraphs:
                            if "年" in paragraph.text and "月" in paragraph.text and "日" in paragraph.text:
                                print(f"找到日期段落: {paragraph.text}")

                                original_text = paragraph.text
                                new_date = f'{year}年{month}月{day}日'

                                # 查找日期部分的模式 - 匹配各种日期格式
                                date_pattern = r'(\d+年\d+月\d+日|年\s*月\s*日\.?|年月日\.?)'

                                if re.search(date_pattern, original_text):
                                    # 清空段落
                                    paragraph.clear()

                                    # 使用正则表达式替换，同时保持格式
                                    current_pos = 0

                                    for match in re.finditer(date_pattern, original_text):
                                        # 添加匹配前的文本（标签部分），保持原有格式
                                        before_text = original_text[current_pos:match.start()]
                                        if before_text:
                                            paragraph.add_run(before_text)

                                        # 添加日期部分，设置楷体五号字体
                                        date_run = paragraph.add_run(new_date)
                                        date_run.font.name = "楷体"
                                        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                        date_run.font.size = Pt(10.5)

                                        current_pos = match.end()

                                    # 添加剩余的文本（如果有的话）
                                    remaining_text = original_text[current_pos:]
                                    if remaining_text:
                                        paragraph.add_run(remaining_text)

                                    date_found = True
                                    print("已更新审核日期并设置为楷体五号字体")
                                    break

                        # 如果没有找到日期段落，尝试创建新段落
                        if not date_found:
                            print("未在审核单元格中找到日期段落，尝试其他方法...")
                            # 添加新段落
                            p = cell.add_paragraph()
                            run = p.add_run(f"{year}年{month}月{day}日")
                            # 设置楷体五号字体
                            run.font.name = "楷体"
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                            run.font.size = Pt(10.5)
                            print("已添加审核日期并设置为楷体五号字体")
            
            # 查找表头行，确定各列的位置
            column_indices = {}
            header_row_index = -1
            
            # 查找包含"委托单编号"、"单线号"等的行
            for i, row in enumerate(table.rows):
                header_found = False
                for j, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if "委托单编号" in cell_text:
                        column_indices["委托单编号"] = j
                        header_row_index = i
                        header_found = True
                    elif "检测批号" in cell_text:
                        column_indices["检测批号"] = j
                        header_found = True
                    elif "单线号" in cell_text:
                        column_indices["单线号"] = j
                        header_found = True
                    elif "焊口号" in cell_text:
                        column_indices["焊口号"] = j
                        header_found = True
                    elif "焊工号" in cell_text:
                        column_indices["焊工号"] = j
                        header_found = True
                    elif "检测结果" in cell_text:
                        column_indices["检测结果"] = j
                        header_found = True
                    elif "返修张/处数" in cell_text:
                        column_indices["返修张/处数"] = j
                        header_found = True
                    elif "备注" in cell_text:
                        column_indices["备注"] = j
                        header_found = True
                
                if header_found and header_row_index >= 0:
                    break
            
            print(f"找到表头行: 第{header_row_index+1}行")
            print(f"列索引: {column_indices}")
            
            # 如果找到表头行，处理数据填充
            if header_row_index >= 0 and column_indices:
                # 获取可用于填充数据的行
                data_rows = []
                for i in range(header_row_index + 1, len(table.rows)):
                    if i < len(table.rows):
                        # 检查是否是空行或包含特殊标记的行
                        if "以下空白" in table.rows[i].cells[0].text if len(table.rows[i].cells) > 0 else False:
                            print(f"找到'以下空白'行: 第{i+1}行")
                            break
                        # 添加可用于填充数据的行
                        data_rows.append(i)
                
                print(f"找到{len(data_rows)}行可用于填充数据")
                
                # 确定需要填充的数据行数
                data_count = len(order_df)
                print(f"需要填充{data_count}行数据")
                
                # 如果Word表格中的行数不足，需要添加新行
                rows_needed = data_count - len(data_rows)
                if rows_needed > 0:
                    print(f"需要添加{rows_needed}行到表格中")
                    # 找到最后一行的索引
                    last_row_idx = data_rows[-1] if data_rows else header_row_index
                    
                    # 添加新行
                    for _ in range(rows_needed):
                        # 在表格末尾添加一行
                        new_row = table.add_row()
                        data_rows.append(len(table.rows) - 1)  # 添加新行的索引
                
                # 处理每一行数据
                for i in range(data_count):
                    if i < len(data_rows):
                        row_idx = data_rows[i]
                        row = table.rows[row_idx]
                        
                        # 1. 填写委托单编号
                        if "委托单编号" in column_indices:
                            col_idx = column_indices["委托单编号"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                if cell.paragraphs:
                                    paragraph = cell.paragraphs[0]
                                    paragraph.clear()
                                    run = paragraph.add_run(str(order_number))
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已更新第{row_idx+1}行委托单编号: {order_number}")

                        # 2. 填写检测批号（填入"/"）
                        if "检测批号" in column_indices:
                            col_idx = column_indices["检测批号"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                if cell.paragraphs:
                                    paragraph = cell.paragraphs[0]
                                    paragraph.clear()
                                    run = paragraph.add_run("/")
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已更新第{row_idx+1}行检测批号: /")

                        # 3. 填写单线号（检件编号）
                        if "单线号" in column_indices and i < len(inspection_numbers):
                            col_idx = column_indices["单线号"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                if cell.paragraphs:
                                    paragraph = cell.paragraphs[0]
                                    paragraph.clear()
                                    run = paragraph.add_run(str(inspection_numbers[i]))
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已更新第{row_idx+1}行单线号: {inspection_numbers[i]}")
                        
                        # 4. 填写焊口号
                        if "焊口号" in column_indices and i < len(weld_numbers):
                            col_idx = column_indices["焊口号"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                if cell.paragraphs:
                                    paragraph = cell.paragraphs[0]
                                    paragraph.clear()
                                    run = paragraph.add_run(str(weld_numbers[i]))
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已更新第{row_idx+1}行焊口号: {weld_numbers[i]}")

                        # 5. 填写焊工号
                        if "焊工号" in column_indices and i < len(welder_numbers):
                            col_idx = column_indices["焊工号"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                if cell.paragraphs:
                                    paragraph = cell.paragraphs[0]
                                    paragraph.clear()
                                    run = paragraph.add_run(str(welder_numbers[i]))
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已更新第{row_idx+1}行焊工号: {welder_numbers[i]}")

                        # 6. 填写检测结果（返修补片）
                        if "检测结果" in column_indices and i < len(repair_results):
                            col_idx = column_indices["检测结果"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                repair_result = repair_results[i]
                                if cell.paragraphs:
                                    paragraph = cell.paragraphs[0]
                                    paragraph.clear()
                                    # 检查是否为空或NaN
                                    if pd.isna(repair_result):
                                        run = paragraph.add_run("")
                                    else:
                                        run = paragraph.add_run(str(repair_result))
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已更新第{row_idx+1}行检测结果")

                        # 7. 填写返修张/处数（实际不合格）
                        if "返修张/处数" in column_indices and i < len(failure_counts):
                            col_idx = column_indices["返修张/处数"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                failure_count = failure_counts[i]
                                if cell.paragraphs:
                                    paragraph = cell.paragraphs[0]
                                    paragraph.clear()
                                    # 检查是否为空或NaN
                                    if pd.isna(failure_count):
                                        text_value = "0"  # 为空填写0
                                    else:
                                        # 转换为整数格式，去除小数点
                                        try:
                                            # 如果是数字，转换为整数
                                            if isinstance(failure_count, (int, float)):
                                                text_value = str(int(failure_count))
                                            else:
                                                # 如果是字符串，尝试转换为数字再转整数
                                                numeric_value = float(str(failure_count))
                                                text_value = str(int(numeric_value))
                                        except (ValueError, TypeError):
                                            # 如果转换失败，保持原值
                                            text_value = str(failure_count)

                                    run = paragraph.add_run(text_value)
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已更新第{row_idx+1}行返修张/处数: {text_value}")

                        # 8. 填写备注
                        if "备注" in column_indices and i < len(notes):
                            col_idx = column_indices["备注"]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                note = notes[i]
                                if cell.paragraphs:
                                    paragraph = cell.paragraphs[0]
                                    paragraph.clear()
                                    # 检查是否为空或NaN
                                    if pd.isna(note):
                                        run = paragraph.add_run("")
                                    else:
                                        run = paragraph.add_run(str(note))
                                    # 设置楷体五号字体
                                    run.font.name = "楷体"
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                    run.font.size = Pt(10.5)
                                    print(f"已更新第{row_idx+1}行备注")

                # 在单线号数据内容的下一行添加"以下空白"
                print("\n==== 添加'以下空白'提示 ====")
                if "单线号" in column_indices and data_count > 0:
                    next_empty_row_idx = data_rows[data_count - 1] + 1  # 数据最后一行的下一行
                    if next_empty_row_idx < len(table.rows):
                        next_row = table.rows[next_empty_row_idx]
                        single_line_col_idx = column_indices["单线号"]
                        if single_line_col_idx < len(next_row.cells):
                            cell = next_row.cells[single_line_col_idx]
                            if cell.paragraphs:
                                paragraph = cell.paragraphs[0]
                                paragraph.clear()
                                run = paragraph.add_run("以下空白")
                                # 设置楷体五号字体
                                run.font.name = "楷体"
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                run.font.size = Pt(10.5)
                                set_cell_center_alignment(cell)  # 设置居中
                                print(f"已在第{next_empty_row_idx+1}行单线号列添加'以下空白'并设置居中")
                    else:
                        # 如果没有足够的行，添加新行
                        new_row = table.add_row()
                        single_line_col_idx = column_indices["单线号"]
                        if single_line_col_idx < len(new_row.cells):
                            cell = new_row.cells[single_line_col_idx]
                            if cell.paragraphs:
                                paragraph = cell.paragraphs[0]
                                paragraph.clear()
                                run = paragraph.add_run("以下空白")
                                # 设置楷体五号字体
                                run.font.name = "楷体"
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                run.font.size = Pt(10.5)
                                set_cell_center_alignment(cell)  # 设置居中
                                print(f"已添加新行并在单线号列添加'以下空白'并设置居中")

        # 保存文档
        doc.save(report_output_path)
        print(f"文档已保存至: {report_output_path}")
        success_count += 1
    
    print(f"\n处理完成: 共处理{len(order_numbers)}个委托单编号，成功生成{success_count}份报告")
    return success_count > 0

def main():
    # 创建命令行参数解析器
    parser = argparse.ArgumentParser(description='将Excel数据填入Word文档')
    parser.add_argument('-e', '--excel', default="生成器/Excel/2_生成器结果.xlsx", 
                        help='Excel表格路径 (默认: 生成器/Excel/2_生成器结果.xlsx)')
    parser.add_argument('-w', '--word', default="生成器/word/2_RT结果通知台账_Mode2.docx", 
                        help='Word模板文档路径 (默认: 生成器/word/2_RT结果通知台账_Mode2.docx)')
    parser.add_argument('-o', '--output', 
                        help='输出目录 (可选，默认为"生成器/输出报告"目录)')
    parser.add_argument('-p', '--project', 
                        help='工程名称，用于替换文档中的"工程名称参数值"')
    parser.add_argument('-c', '--client', 
                        help='委托单位，用于替换文档中的"委托单位参数值"')
    parser.add_argument('-m', '--method', 
                        help='检测方法，用于替换文档中的"检测方法参数"')
    
    # 解析命令行参数
    args = parser.parse_args()
    
    # 处理Excel到Word的转换
    success = process_excel_to_word(args.excel, args.word, args.output, args.project, args.client, args.method)
    
    # 返回状态码
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main() 