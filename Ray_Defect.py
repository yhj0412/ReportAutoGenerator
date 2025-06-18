import pandas as pd
import os
import sys
import argparse
from docx import Document
from datetime import datetime
import re

def find_column_with_keyword(df, keyword):
    """查找包含指定关键字的列"""
    matching_cols = [col for col in df.columns if keyword.lower() in col.lower()]
    return matching_cols[0] if matching_cols else None

def get_output_filename(word_template_path, order_number, ray_type):
    """根据Word模板路径、委托单编号和射线类型生成输出文件名
    
    Args:
        word_template_path: Word模板文档路径
        order_number: 委托单编号
        ray_type: 射线类型（"X射线"或"γ射线"）
    
    Returns:
        str: 输出文件名
    """
    # 获取模板文件名（不含路径和扩展名）
    template_name = os.path.splitext(os.path.basename(word_template_path))[0]
    # 射线类型标识
    ray_mark = "γ" if ray_type == "γ射线" else "X"
    # 生成输出文件名
    return f"{template_name}_{order_number}_{ray_mark}_生成结果.docx"

def process_excel_to_word(excel_path, word_template_path, output_path=None):
    """将Excel数据填入Word文档
    
    Args:
        excel_path: Excel表格路径
        word_template_path: Word模板文档路径
        output_path: 输出Word文档路径（如果为None，将自动生成）
    
    Returns:
        bool: 处理是否成功
    """
    # 检查文件是否存在
    if not os.path.exists(excel_path):
        print(f"错误: Excel文件不存在: {excel_path}")
        return False
        
    if not os.path.exists(word_template_path):
        print(f"错误: Word模板文件不存在: {word_template_path}")
        return False
    
    # 创建输出目录
    if output_path is None:
        output_dir = os.path.join("生成器", "输出报告", "4_射线检测记录")
    else:
        output_dir = output_path
        
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            print(f"创建输出目录: {output_dir}")
        except Exception as e:
            print(f"错误: 无法创建输出目录: {e}")
            return False
    
    # 读取Excel数据
    print(f"正在读取Excel文件: {excel_path}")
    try:
        df = pd.read_excel(excel_path)
        print(f"成功读取Excel文件，共有{len(df)}行数据")
    except Exception as e:
        print(f"错误: 无法读取Excel文件: {e}")
        return False
    
    # 打印所有列名，帮助调试
    print(f"Excel表格列名: {list(df.columns)}")
    
    # 定义需要查找的列关键字
    column_keywords = {
        '完成日期': '完成日期',
        '委托单编号': '委托单编号',
        '检件编号': '检件编号',
        '焊口编号': '焊口编号',
        '焊工号': '焊工号',
        '合格级别': '合格级别',
        '检测比例': '检测比例',
        'γ射线': 'γ射线',
        '焊接方法': '焊接方法',
        '检测时机': '检测时机'
    }
    
    # 查找每个关键字对应的实际列名
    column_mapping = {}
    missing_columns = []
    
    for key, keyword in column_keywords.items():
        col_name = find_column_with_keyword(df, keyword)
        if col_name:
            column_mapping[key] = col_name
            print(f"找到列: '{key}' -> '{col_name}'")
        else:
            missing_columns.append(key)
            
    if missing_columns:
        print(f"警告: 未找到以下列: {', '.join(missing_columns)}")
        # 尝试使用列位置
        possible_columns = {
            '完成日期': 'B', 
            '委托单编号': 'C', 
            '检件编号': 'D', 
            '焊口编号': 'E', 
            '焊工号': 'F',
            '合格级别': 'I',
            '检测比例': 'J',
            'γ射线': 'P',
            '焊接方法': 'R',
            '检测时机': 'V'
        }
        
        for key in missing_columns:
            col_letter = possible_columns.get(key)
            if col_letter:
                col_idx = ord(col_letter) - ord('A')
                if col_idx < len(df.columns):
                    column_mapping[key] = df.columns[col_idx]
                    print(f"使用列位置找到: '{key}' -> '{df.columns[col_idx]}'")
    
    # 如果缺少关键列，则无法继续处理
    if '委托单编号' not in column_mapping:
        print("错误: 无法找到委托单编号列")
        return False
    
    if 'γ射线' not in column_mapping:
        print("警告: 无法找到γ射线列，将默认所有记录为X射线")
        # 增加一个全为空的列作为γ射线列
        df['γ射线'] = None
        column_mapping['γ射线'] = 'γ射线'
    
    # 根据委托单编号和射线类型分组数据
    groups = []
    order_numbers = df[column_mapping['委托单编号']].dropna().unique()
    
    for order_number in order_numbers:
        # 获取该委托单编号的所有数据
        order_df = df[df[column_mapping['委托单编号']] == order_number]
        
        # 获取该委托单编号下的所有射线类型
        ray_types = order_df[column_mapping['γ射线']].dropna().unique()
        
        # 如果没有明确的γ射线值，则视为X射线
        if len(ray_types) == 0:
            groups.append({
                'order_number': order_number,
                'ray_type': 'X射线',  # X射线表示为处理逻辑，但射源种类值会设置为空
                'data': order_df
            })
            print(f"委托单编号 {order_number} 没有明确的射线类型，处理为X射线（射源种类值为空）")
        else:
            # 有γ射线值的处理为γ射线
            for ray_type in ray_types:
                ray_df = order_df[order_df[column_mapping['γ射线']] == ray_type]
                groups.append({
                    'order_number': order_number,
                    'ray_type': 'γ射线',
                    'data': ray_df
                })
                print(f"委托单编号 {order_number} 的射线类型 γ射线 有 {len(ray_df)} 条记录")
            
            # 没有γ射线值的处理为X射线
            x_ray_df = order_df[order_df[column_mapping['γ射线']].isna()]
            if len(x_ray_df) > 0:
                groups.append({
                    'order_number': order_number,
                    'ray_type': 'X射线',  # X射线表示为处理逻辑，但射源种类值会设置为空
                    'data': x_ray_df
                })
                print(f"委托单编号 {order_number} 的射线类型 X射线 有 {len(x_ray_df)} 条记录")
    
    print(f"共有 {len(groups)} 个组合需要生成报告")
    
    # 处理每个分组
    success_count = 0
    error_count = 0
    
    for group in groups:
        order_number = group['order_number']
        ray_type = group['ray_type']
        group_df = group['data']
        
        print(f"\n{'='*50}")
        print(f"处理委托单编号: {order_number}, 射线类型: {ray_type}")
        print(f"{'='*50}")
        
        try:
            # 为该分组生成输出文件名
            output_filename = get_output_filename(word_template_path, order_number, ray_type)
            report_output_path = os.path.join(output_dir, output_filename)
            print(f"输出文件路径: {report_output_path}")
            
            # 打开Word文档
            print(f"正在处理Word文档: {word_template_path}")
            
            try:
                # 每次处理新的组合时，重新从模板创建文档对象
                # 这确保了每个组合都会生成一个独立的文档
                doc = Document(word_template_path)
                print(f"成功从模板创建新文档")
            except Exception as e:
                print(f"无法打开Word文档: {e}")
                error_count += 1
                continue
            
            # 填充文档的其余部分将在这里添加...
            
            # 1) 获取该组数据中最晚的完成日期
            date_col = column_mapping.get('完成日期')
            if date_col:
                # 确保日期列是日期类型
                group_df[date_col] = pd.to_datetime(group_df[date_col], errors='coerce')
                latest_date = group_df[date_col].max()
                
                if pd.isna(latest_date):
                    print(f"警告: 委托单编号 {order_number} 没有有效的完成日期")
                    year, month, day = datetime.now().year, datetime.now().month, datetime.now().day
                else:
                    # 将日期转换为年、月、日
                    year = latest_date.year
                    month = latest_date.month
                    day = latest_date.day
                    print(f"找到最晚完成日期: {year}年{month}月{day}日")
            else:
                print("警告: 未找到完成日期列")
                year, month, day = datetime.now().year, datetime.now().month, datetime.now().day
            
            # 获取相关数据
            inspection_numbers = group_df[column_mapping.get('检件编号')].dropna().tolist() if '检件编号' in column_mapping else []
            weld_numbers = group_df[column_mapping.get('焊口编号')].dropna().tolist() if '焊口编号' in column_mapping else []
            welder_numbers = group_df[column_mapping.get('焊工号')].dropna().tolist() if '焊工号' in column_mapping else []
            
            # 获取分组的第一个值
            grade_level = ""
            if '合格级别' in column_mapping:
                grade_levels = group_df[column_mapping['合格级别']].dropna().tolist()
                if grade_levels:
                    grade_level = grade_levels[0]
                    print(f"找到合格级别: {grade_level}")
            
            inspection_ratio = ""
            if '检测比例' in column_mapping:
                ratios = group_df[column_mapping['检测比例']].dropna().tolist()
                if ratios:
                    # 转换为百分数格式
                    try:
                        ratio_value = float(ratios[0])
                        inspection_ratio = f"{ratio_value*100:.0f}%"
                    except (ValueError, TypeError):
                        inspection_ratio = str(ratios[0])
                    print(f"找到检测比例: {inspection_ratio}")
            
            welding_method = ""
            if '焊接方法' in column_mapping:
                methods = group_df[column_mapping['焊接方法']].dropna().tolist()
                if methods:
                    welding_method = methods[0]
                    print(f"找到焊接方法: {welding_method}")
            
            inspection_time = ""
            if '检测时机' in column_mapping:
                times = group_df[column_mapping['检测时机']].dropna().tolist()
                if times:
                    inspection_time = times[0]
                    print(f"找到检测时机: {inspection_time}")
            
            # 根据射线类型设置相关参数
            if ray_type == "γ射线":
                focus_size = "3*3"
                lead_screen = "柯达0.1*2"
                film_grade = "柯达MX125"
                ray_source = "γ射线"
            else:  # X射线
                focus_size = "2.5*2.5"
                lead_screen = "0.03*2"
                film_grade = "锐科R400"
                ray_source = ""  # 当γ射线的值为空时，射源种类值也设置为空
            
            print(f"射线类型: {ray_type}, 焦点尺寸: {focus_size}, 铅增感屏: {lead_screen}, 胶片等级: {film_grade}, 射源种类值: {ray_source}")
            
            # 替换文档中的值
            print("\n==== 开始替换文档中的值 ====")
            
            # 将EPKJ拼接委托单编号代替委托单编号值
            committee_order = f"EPKJ-{order_number}"
            replaced = False
            
            # 遍历所有段落，替换关键词
            for paragraph in doc.paragraphs:
                if "委托单编号值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("委托单编号值", committee_order)
                    print(f"已将段落中的'委托单编号值'替换为'{committee_order}'")
                    replaced = True
                
                if "射源种类值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("射源种类值", ray_source)
                    print(f"已将段落中的'射源种类值'替换为'{ray_source}'")
                    replaced = True
                
                if "合格级别值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("合格级别值", grade_level)
                    print(f"已将段落中的'合格级别值'替换为'{grade_level}'")
                    replaced = True
                
                if "检测比例值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("检测比例值", inspection_ratio)
                    print(f"已将段落中的'检测比例值'替换为'{inspection_ratio}'")
                    replaced = True
                
                if "焊接方法值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("焊接方法值", welding_method)
                    print(f"已将段落中的'焊接方法值'替换为'{welding_method}'")
                    replaced = True
                
                if "检测时机值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("检测时机值", inspection_time)
                    print(f"已将段落中的'检测时机值'替换为'{inspection_time}'")
                    replaced = True
                
                if "焦点尺寸值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("焦点尺寸值", focus_size)
                    print(f"已将段落中的'焦点尺寸值'替换为'{focus_size}'")
                    replaced = True
                
                if "铅增感屏值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("铅增感屏值", lead_screen)
                    print(f"已将段落中的'铅增感屏值'替换为'{lead_screen}'")
                    replaced = True
                
                if "胶片等级值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("胶片等级值", film_grade)
                    print(f"已将段落中的'胶片等级值'替换为'{film_grade}'")
                    replaced = True
            
            # 遍历表格中的单元格，替换关键词
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if "委托单编号值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("委托单编号值", committee_order)
                                print(f"已将表格单元格中的'委托单编号值'替换为'{committee_order}'")
                                replaced = True
                            
                            if "射源种类值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("射源种类值", ray_source)
                                print(f"已将表格单元格中的'射源种类值'替换为'{ray_source}'")
                                replaced = True
                            
                            if "合格级别值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("合格级别值", grade_level)
                                print(f"已将表格单元格中的'合格级别值'替换为'{grade_level}'")
                                replaced = True
                            
                            if "检测比例值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("检测比例值", inspection_ratio)
                                print(f"已将表格单元格中的'检测比例值'替换为'{inspection_ratio}'")
                                replaced = True
                            
                            if "焊接方法值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("焊接方法值", welding_method)
                                print(f"已将表格单元格中的'焊接方法值'替换为'{welding_method}'")
                                replaced = True
                            
                            if "检测时机值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("检测时机值", inspection_time)
                                print(f"已将表格单元格中的'检测时机值'替换为'{inspection_time}'")
                                replaced = True
                            
                            if "焦点尺寸值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("焦点尺寸值", focus_size)
                                print(f"已将表格单元格中的'焦点尺寸值'替换为'{focus_size}'")
                                replaced = True
                            
                            if "铅增感屏值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("铅增感屏值", lead_screen)
                                print(f"已将表格单元格中的'铅增感屏值'替换为'{lead_screen}'")
                                replaced = True
                            
                            if "胶片等级值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("胶片等级值", film_grade)
                                print(f"已将表格单元格中的'胶片等级值'替换为'{film_grade}'")
                                replaced = True
            
            if not replaced:
                print("警告: 未找到需要替换的关键词，可能需要检查Word模板中的占位符命名。")
            
            # 填写日期（洗片人、拍片人、审核人）
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        # 处理日期
                        date_patterns = ["洗片人", "拍片人", "审核人"]
                        for pattern in date_patterns:
                            if pattern in cell.text:
                                print(f"找到{pattern}单元格: 表格行{i+1}, 列{j+1}")
                                
                                # 检查单元格中的所有段落
                                date_found = False
                                for paragraph in cell.paragraphs:
                                    if "年" in paragraph.text and "月" in paragraph.text and "日" in paragraph.text:
                                        print(f"找到日期段落: {paragraph.text}")
                                        
                                        # 创建新的文本，确保只有一个年月日
                                        new_text = paragraph.text
                                        # 确保年月日前没有数字
                                        new_text = re.sub(r'\d*年', '年', new_text)
                                        new_text = re.sub(r'\d*月', '月', new_text)
                                        new_text = re.sub(r'\d*日', '日', new_text)
                                        
                                        # 在年月日前插入正确的数字
                                        new_text = new_text.replace('年', f'{year}年')
                                        new_text = new_text.replace('月', f'{month}月')
                                        new_text = new_text.replace('日', f'{day}日')
                                        
                                        paragraph.text = new_text
                                        date_found = True
                                        print(f"已更新{pattern}日期为 {year}年{month}月{day}日")
                                        break
                                
                                # 如果没有找到日期段落，尝试创建新段落
                                if not date_found:
                                    print(f"未在{pattern}单元格中找到日期段落，尝试添加")
                                    # 添加新段落
                                    p = cell.add_paragraph(f"{year}年{month}月{day}日")
                                    print(f"已添加{pattern}日期: {year}年{month}月{day}日")
            
            # 查找表头行，确定各列的位置
            for table in doc.tables:
                column_indices = {}
                header_row_index = -1
                
                # 查找包含"检件编号"、"焊缝编号"、"焊工号"的行
                for i, row in enumerate(table.rows):
                    header_found = False
                    for j, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        
                        if "检件编号" in cell_text:
                            column_indices["检件编号"] = j
                            header_row_index = i
                            header_found = True
                        elif "焊缝编号" in cell_text or "焊口编号" in cell_text:
                            column_indices["焊口编号"] = j
                            header_found = True
                        elif "焊工号" in cell_text:
                            column_indices["焊工号"] = j
                            header_found = True
                        elif "备注" in cell_text:
                            column_indices["备注"] = j
                            header_found = True
                    
                    if header_found and header_row_index >= 0:
                        print(f"找到表头行: 第{header_row_index+1}行")
                        
                        # 将"焊口编号"的键名更新为"焊缝编号"以保持一致性
                        if "焊口编号" in column_indices:
                            column_indices["焊缝编号"] = column_indices.pop("焊口编号")
                            
                        print(f"列索引: {column_indices}")
                        break
                
                # 如果找到表头行，处理数据填充
                if header_row_index >= 0 and column_indices:
                    # 获取可用于填充数据的行
                    data_rows = []
                    for i in range(header_row_index + 1, len(table.rows)):
                        if i < len(table.rows):
                            # 检查是否是空行或包含特殊标记的行
                            if "以下空白" in table.rows[i].cells[0].text if len(table.rows[i].cells) > 0 else False:
                                print(f"找到'以下空白'行: 第{i+1}行")
                                break
                            # 添加可用于填充数据的行
                            data_rows.append(i)
                    
                    print(f"找到{len(data_rows)}行可用于填充数据")
                    
                    # 确定需要填充的数据行数
                    data_count = len(group_df)
                    print(f"需要填充{data_count}行数据")
                    
                    # 如果Word表格中的行数不足，需要添加新行
                    rows_needed = data_count - len(data_rows)
                    if rows_needed > 0:
                        print(f"需要添加{rows_needed}行到表格中")
                        # 找到最后一行的索引
                        last_row_idx = data_rows[-1] if data_rows else header_row_index
                        
                        # 添加新行
                        for _ in range(rows_needed):
                            # 在表格末尾添加一行
                            new_row = table.add_row()
                            data_rows.append(len(table.rows) - 1)  # 添加新行的索引
                    
                    # 处理每一行数据
                    for i in range(data_count):
                        if i < len(data_rows):
                            row_idx = data_rows[i]
                            row = table.rows[row_idx]
                            
                            # 1. 填写检件编号
                            if "检件编号" in column_indices and i < len(inspection_numbers):
                                col_idx = column_indices["检件编号"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    if cell.paragraphs:
                                        cell.paragraphs[0].text = str(inspection_numbers[i])
                                        print(f"已更新第{row_idx+1}行检件编号: {inspection_numbers[i]}")
                            
                            # 2. 填写焊缝编号
                            if "焊缝编号" in column_indices and i < len(weld_numbers):
                                col_idx = column_indices["焊缝编号"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    if cell.paragraphs:
                                        cell.paragraphs[0].text = str(weld_numbers[i])
                                        print(f"已更新第{row_idx+1}行焊缝编号: {weld_numbers[i]}")
                            
                            # 3. 填写焊工号
                            if "焊工号" in column_indices and i < len(welder_numbers):
                                col_idx = column_indices["焊工号"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    if cell.paragraphs:
                                        cell.paragraphs[0].text = str(welder_numbers[i])
                                        print(f"已更新第{row_idx+1}行焊工号: {welder_numbers[i]}")
                            
                            # 4. 填写备注（填入完成日期）
                            if "备注" in column_indices and date_col and i < len(group_df):
                                col_idx = column_indices["备注"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    # 获取对应行的完成日期
                                    if not pd.isna(group_df[date_col].iloc[i]):
                                        date_value = group_df[date_col].iloc[i]
                                        if isinstance(date_value, pd.Timestamp):
                                            # 如果是日期类型，格式化为字符串
                                            formatted_date = date_value.strftime("%Y年%m月%d日")
                                        else:
                                            # 如果不是日期类型，直接转为字符串
                                            formatted_date = str(date_value)
                                        
                                        if cell.paragraphs:
                                            cell.paragraphs[0].text = formatted_date
                                            print(f"已更新第{row_idx+1}行备注（完成日期）: {formatted_date}")
            
            print("==== 文档填充完成 ====\n")
            
            # 保存文档
            try:
                print(f"\n正在保存文档到: {report_output_path}")
                doc.save(report_output_path)
                print(f"文档已成功保存至: {report_output_path}")
                success_count += 1
            except Exception as e:
                print(f"错误: 无法保存文档: {e}")
                error_count += 1
                
        except Exception as e:
            print(f"错误: 处理委托单编号 {order_number} 和射线类型 {ray_type} 时出错: {e}")
            error_count += 1
    
    print(f"\n处理完成: 共处理{len(groups)}个组合，成功生成{success_count}份报告，失败{error_count}份")
    if error_count > 0:
        print(f"警告: 有{error_count}个组合处理失败，请检查日志")
    
    return success_count > 0

def main():
    # 创建命令行参数解析器
    parser = argparse.ArgumentParser(description='将Excel数据填入Word文档')
    parser.add_argument('-e', '--excel', default="生成器/Excel/4_生成器台账-射线检测记录.xlsx", 
                        help='Excel表格路径 (默认: 生成器/Excel/4_生成器台账-射线检测记录.xlsx)')
    parser.add_argument('-w', '--word', default="生成器/wod/4_射线检测记录.docx", 
                        help='Word模板文档路径 (默认: 生成器/wod/4_射线检测记录.docx)')
    parser.add_argument('-o', '--output', 
                        help='输出目录 (可选，默认为"生成器/输出报告/4_射线检测记录"目录)')
    
    # 解析命令行参数
    args = parser.parse_args()
    
    # 处理Excel到Word的转换
    success = process_excel_to_word(args.excel, args.word, args.output)
    
    # 返回状态码
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
