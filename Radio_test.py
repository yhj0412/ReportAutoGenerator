import pandas as pd
import os
import sys
import argparse
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
import re

def find_column_with_keyword(df, keyword):
    """查找包含指定关键字的列"""
    matching_cols = [col for col in df.columns if keyword.lower() in col.lower()]
    return matching_cols[0] if matching_cols else None

def set_font_style(paragraph, font_name="楷体", font_size=10.5):
    """设置段落字体为楷体五号（10.5磅）

    Args:
        paragraph: Word段落对象
        font_name: 字体名称，默认为"楷体"
        font_size: 字体大小，默认为10.5磅（五号字）
    """
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def normalize_text(text):
    """标准化文本以便更好的匹配"""
    if not text:
        return ""
    # 移除空格、换行符等空白字符
    normalized = re.sub(r'\s+', '', text.strip())
    # 转换为小写
    normalized = normalized.lower()
    return normalized

def find_detection_timing_options(doc):
    """在Word文档中查找检测时机相关的复选框选项"""
    timing_options = []
    processed_cells = set()  # 避免重复处理同一个单元格

    print("开始查找检测时机复选框选项...")

    # 遍历所有表格
    for table_idx, table in enumerate(doc.tables):
        print(f"\n检查表格 {table_idx+1}...")

        for row_idx, row in enumerate(table.rows):
            # 检查每一行是否包含检测时机相关内容
            row_text = ""
            for cell in row.cells:
                row_text += cell.text + " "

            # 如果这一行包含检测时机相关内容，搜索整行的复选框选项
            if "检测时机" in row_text or any(keyword in row_text for keyword in ["焊后", "焊前", "打磨", "热处理"]):
                print(f"找到可能的检测时机行: 表格{table_idx+1}, 行{row_idx+1}")
                print(f"行内容: '{row_text.strip()}'")

                # 查找该行及相邻行中所有包含复选框的单元格
                search_rows = [row_idx]
                if row_idx > 0:
                    search_rows.append(row_idx - 1)  # 上一行
                if row_idx < len(table.rows) - 1:
                    search_rows.append(row_idx + 1)  # 下一行

                for search_row_idx in search_rows:
                    if search_row_idx < 0 or search_row_idx >= len(table.rows):
                        continue

                    search_row = table.rows[search_row_idx]
                    for check_cell_idx, check_cell in enumerate(search_row.cells):
                        cell_key = (table_idx, search_row_idx, check_cell_idx)
                        if cell_key in processed_cells:
                            continue
                        processed_cells.add(cell_key)

                        cell_text = check_cell.text.strip()

                        # 查找包含复选框的选项
                        if '□' in cell_text or '☑' in cell_text or '✓' in cell_text:
                            print(f"检查单元格({search_row_idx+1}, {check_cell_idx+1}): '{cell_text}'")

                            # 分割多个选项（如果在同一个单元格中）
                            lines = cell_text.split('\n')
                            for line in lines:
                                line = line.strip()
                                if ('□' in line or '☑' in line or '✓' in line) and len(line) > 1:
                                    # 提取选项文本（去除复选框符号）
                                    option_text = line.replace('□', '').replace('☑', '').replace('✓', '').strip()
                                    if option_text and len(option_text) > 0:
                                        # 避免重复添加相同的选项
                                        existing_option = None
                                        for existing in timing_options:
                                            if (existing['text'] == option_text and
                                                existing['position'][0] == search_row_idx):
                                                existing_option = existing
                                                break

                                        if not existing_option:
                                            timing_options.append({
                                                'text': option_text,
                                                'original_line': line,
                                                'cell': check_cell,
                                                'position': (search_row_idx, check_cell_idx),
                                                'table_idx': table_idx
                                            })
                                            print(f"找到检测时机选项: '{option_text}' 在位置({search_row_idx+1}, {check_cell_idx+1})")

    # 如果没有找到选项，进行全文档搜索
    if not timing_options:
        print("\n未找到检测时机选项，进行全文档复选框搜索...")
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_key = (table_idx, row_idx, cell_idx)
                    if cell_key in processed_cells:
                        continue

                    cell_text = cell.text.strip()
                    if '□' in cell_text or '☑' in cell_text or '✓' in cell_text:
                        lines = cell_text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if ('□' in line or '☑' in line or '✓' in line) and len(line) > 1:
                                option_text = line.replace('□', '').replace('☑', '').replace('✓', '').strip()
                                if option_text and len(option_text) > 0:
                                    timing_options.append({
                                        'text': option_text,
                                        'original_line': line,
                                        'cell': cell,
                                        'position': (row_idx, cell_idx),
                                        'table_idx': table_idx
                                    })
                                    print(f"全文档搜索找到选项: '{option_text}' 在位置({row_idx+1}, {cell_idx+1})")

    print(f"总共找到 {len(timing_options)} 个检测时机选项")
    if timing_options:
        print("所有检测时机选项:")
        for i, option in enumerate(timing_options):
            print(f"  {i+1}. '{option['text']}' (原文: '{option['original_line']}')")

    return timing_options

def match_timing_option(timing_value, options):
    """将检测时机值与可用选项进行匹配"""
    if not timing_value or not options:
        return None

    normalized_timing = normalize_text(timing_value)
    print(f"尝试匹配检测时机值: '{timing_value}' (标准化: '{normalized_timing}')")

    # 定义检测时机的匹配规则
    timing_patterns = {
        '焊后': ['焊后', '焊接后', '焊完后', '后焊', '焊后检测'],
        '焊前': ['焊前', '焊接前', '前焊', '焊前检测'],
        '打磨': ['打磨', '打磨后', '打磨前'],
        '热处理后': ['热处理后', '热处理', '热处理完成后'],
        '中间': ['中间', '中间检测', '过程中'],
        '最终': ['最终', '最终检测', '终检']
    }

    best_match = None
    best_score = 0

    for option in options:
        option_text = option['text']
        normalized_option = normalize_text(option_text)

        print(f"检查选项: '{option_text}' (标准化: '{normalized_option}')")

        # 1. 完全匹配
        if normalized_timing == normalized_option:
            print(f"找到完全匹配: '{option_text}'")
            return option

        # 2. 使用模式匹配
        for pattern_key, pattern_list in timing_patterns.items():
            for pattern in pattern_list:
                normalized_pattern = normalize_text(pattern)
                if normalized_pattern == normalized_timing:
                    # 检查选项是否包含这个模式
                    if normalized_pattern in normalized_option or pattern_key in normalized_option:
                        score = 1.0  # 模式匹配给最高分
                        if score > best_score:
                            best_score = score
                            best_match = option
                            print(f"找到模式匹配: '{option_text}' 匹配模式 '{pattern}' (得分: {score:.2f})")

        # 3. 包含匹配
        if normalized_timing in normalized_option or normalized_option in normalized_timing:
            score = min(len(normalized_timing), len(normalized_option)) / max(len(normalized_timing), len(normalized_option))
            if score > best_score:
                best_score = score
                best_match = option
                print(f"找到包含匹配: '{option_text}' (得分: {score:.2f})")

        # 4. 关键词匹配
        timing_keywords = ['焊', '后', '前', '打磨', '热处理', '中间', '最终']
        option_keywords = ['焊', '后', '前', '打磨', '热处理', '中间', '最终']

        timing_found_keywords = [kw for kw in timing_keywords if kw in normalized_timing]
        option_found_keywords = [kw for kw in option_keywords if kw in normalized_option]

        if timing_found_keywords and option_found_keywords:
            common_keywords = set(timing_found_keywords) & set(option_found_keywords)
            if common_keywords:
                score = len(common_keywords) / max(len(timing_found_keywords), len(option_found_keywords))
                if score > best_score and score > 0.3:  # 关键词匹配阈值
                    best_score = score
                    best_match = option
                    print(f"找到关键词匹配: '{option_text}' 共同关键词: {common_keywords} (得分: {score:.2f})")

    if best_match and best_score > 0.3:  # 降低最低匹配阈值
        print(f"选择最佳匹配: '{best_match['text']}' (得分: {best_score:.2f})")
        return best_match

    print(f"未找到匹配的检测时机选项")
    return None

def mark_timing_checkbox(option):
    """在匹配的选项前添加勾选标记"""
    try:
        cell = option['cell']
        original_line = option['original_line']
        option_text = option['text']

        print(f"正在标记检测时机选项: '{option_text}'")

        # 遍历单元格中的所有段落
        for paragraph in cell.paragraphs:
            paragraph_text = paragraph.text.strip()

            # 如果段落包含目标选项
            if option_text in paragraph_text and ('□' in paragraph_text or '☑' in paragraph_text or '✓' in paragraph_text):
                # 清空段落并重新构建
                paragraph.clear()

                # 分割段落文本为多行
                lines = paragraph_text.split('\n')
                for i, line in enumerate(lines):
                    line = line.strip()
                    if option_text in line and ('□' in line or '☑' in line or '✓' in line):
                        # 这是目标行，添加勾选标记
                        marked_line = line.replace('□', '☑').replace('✓', '☑')
                        if '☑' not in marked_line:
                            # 如果没有复选框符号，在选项前添加
                            marked_line = f'☑{line}'

                        run = paragraph.add_run(marked_line)
                        run.font.name = "宋体"
                        run.font.size = Pt(9.5)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
                        print(f"已标记选项: '{marked_line}'")
                    else:
                        # 其他行保持原样
                        if line:
                            run = paragraph.add_run(line)
                            run.font.name = "宋体"
                            run.font.size = Pt(9.5)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

                    # 如果不是最后一行，添加换行
                    if i < len(lines) - 1:
                        paragraph.add_run('\n')

                return True

        print(f"警告: 未能在单元格中找到目标选项文本进行标记")
        return False

    except Exception as e:
        print(f"标记检测时机选项时出错: {e}")
        return False

def find_field_options(doc, field_name, field_keywords):
    """通用函数：在Word文档中查找指定字段的复选框选项"""
    field_options = []
    processed_cells = set()

    print(f"开始查找{field_name}复选框选项...")

    # 遍历所有表格
    for table_idx, table in enumerate(doc.tables):
        print(f"\n检查表格 {table_idx+1}...")

        for row_idx, row in enumerate(table.rows):
            # 检查每一行是否包含目标字段相关内容
            row_text = ""
            for cell in row.cells:
                row_text += cell.text + " "

            # 如果这一行包含目标字段相关内容，搜索整行的复选框选项
            if any(keyword in row_text for keyword in field_keywords):
                print(f"找到可能的{field_name}行: 表格{table_idx+1}, 行{row_idx+1}")
                print(f"行内容: '{row_text.strip()}'")

                # 查找该行及相邻行中所有包含复选框的单元格
                search_rows = [row_idx]
                if row_idx > 0:
                    search_rows.append(row_idx - 1)  # 上一行
                if row_idx < len(table.rows) - 1:
                    search_rows.append(row_idx + 1)  # 下一行

                for search_row_idx in search_rows:
                    if search_row_idx < 0 or search_row_idx >= len(table.rows):
                        continue

                    search_row = table.rows[search_row_idx]
                    for check_cell_idx, check_cell in enumerate(search_row.cells):
                        cell_key = (table_idx, search_row_idx, check_cell_idx)
                        if cell_key in processed_cells:
                            continue
                        processed_cells.add(cell_key)

                        cell_text = check_cell.text.strip()

                        # 查找包含复选框的选项
                        if '□' in cell_text or '☑' in cell_text or '✓' in cell_text:
                            print(f"检查单元格({search_row_idx+1}, {check_cell_idx+1}): '{cell_text}'")

                            # 分割多个选项（如果在同一个单元格中）
                            lines = cell_text.split('\n')
                            for line in lines:
                                line = line.strip()
                                if ('□' in line or '☑' in line or '✓' in line) and len(line) > 1:
                                    # 处理单行中的多个复选框选项（如"□GTAW □SMAW"）
                                    checkbox_pattern = r'([□☑✓])([^□☑✓]+?)(?=[□☑✓]|$)'
                                    matches = re.findall(checkbox_pattern, line)

                                    if matches:
                                        for checkbox, option_text in matches:
                                            option_text = option_text.strip()
                                            if option_text and len(option_text) > 0:
                                                # 重构原始行文本
                                                original_line = f"{checkbox}{option_text}"

                                                # 避免重复添加相同的选项
                                                existing_option = None
                                                for existing in field_options:
                                                    if (existing['text'] == option_text and
                                                        existing['position'][0] == search_row_idx):
                                                        existing_option = existing
                                                        break

                                                if not existing_option:
                                                    field_options.append({
                                                        'text': option_text,
                                                        'original_line': original_line,
                                                        'cell': check_cell,
                                                        'position': (search_row_idx, check_cell_idx),
                                                        'table_idx': table_idx
                                                    })
                                                    print(f"找到{field_name}选项: '{option_text}' 在位置({search_row_idx+1}, {check_cell_idx+1})")
                                    else:
                                        # 如果正则匹配失败，使用原来的方法
                                        option_text = line.replace('□', '').replace('☑', '').replace('✓', '').strip()
                                        if option_text and len(option_text) > 0:
                                            # 避免重复添加相同的选项
                                            existing_option = None
                                            for existing in field_options:
                                                if (existing['text'] == option_text and
                                                    existing['position'][0] == search_row_idx):
                                                    existing_option = existing
                                                    break

                                            if not existing_option:
                                                field_options.append({
                                                    'text': option_text,
                                                    'original_line': line,
                                                    'cell': check_cell,
                                                    'position': (search_row_idx, check_cell_idx),
                                                    'table_idx': table_idx
                                                })
                                                print(f"找到{field_name}选项: '{option_text}' 在位置({search_row_idx+1}, {check_cell_idx+1})")

    print(f"总共找到 {len(field_options)} 个{field_name}选项")
    if field_options:
        print(f"所有{field_name}选项:")
        for i, option in enumerate(field_options):
            print(f"  {i+1}. '{option['text']}' (原文: '{option['original_line']}')")

    return field_options

def match_field_option(field_value, options, field_patterns):
    """通用函数：将字段值与可用选项进行匹配"""
    if not field_value or not options:
        return None

    normalized_value = normalize_text(field_value)
    print(f"尝试匹配字段值: '{field_value}' (标准化: '{normalized_value}')")

    best_match = None
    best_score = 0

    for option in options:
        option_text = option['text']
        normalized_option = normalize_text(option_text)

        print(f"检查选项: '{option_text}' (标准化: '{normalized_option}')")

        # 1. 完全匹配
        if normalized_value == normalized_option:
            print(f"找到完全匹配: '{option_text}'")
            return option

        # 2. 使用模式匹配
        for pattern_key, pattern_list in field_patterns.items():
            for pattern in pattern_list:
                normalized_pattern = normalize_text(pattern)
                if normalized_pattern == normalized_value:
                    # 检查选项是否包含这个模式
                    if normalized_pattern in normalized_option or pattern_key in normalized_option:
                        score = 1.0  # 模式匹配给最高分
                        if score > best_score:
                            best_score = score
                            best_match = option
                            print(f"找到模式匹配: '{option_text}' 匹配模式 '{pattern}' (得分: {score:.2f})")

        # 3. 包含匹配
        if normalized_value in normalized_option or normalized_option in normalized_value:
            score = min(len(normalized_value), len(normalized_option)) / max(len(normalized_value), len(normalized_option))
            if score > best_score:
                best_score = score
                best_match = option
                print(f"找到包含匹配: '{option_text}' (得分: {score:.2f})")

        # 4. 关键词匹配
        value_keywords = [kw for kw in ['gtaw', 'smaw', 'saw', '焊', '接', '方法', 'ⅰ', 'ⅱ', 'ⅲ', 'ⅳ', '级', '100%', '50%', '20%', '10%', '5%'] if kw in normalized_value]
        option_keywords = [kw for kw in ['gtaw', 'smaw', 'saw', '焊', '接', '方法', 'ⅰ', 'ⅱ', 'ⅲ', 'ⅳ', '级', '100%', '50%', '20%', '10%', '5%'] if kw in normalized_option]

        if value_keywords and option_keywords:
            common_keywords = set(value_keywords) & set(option_keywords)
            if common_keywords:
                score = len(common_keywords) / max(len(value_keywords), len(option_keywords))
                if score > best_score and score > 0.3:  # 关键词匹配阈值
                    best_score = score
                    best_match = option
                    print(f"找到关键词匹配: '{option_text}' 共同关键词: {common_keywords} (得分: {score:.2f})")

    if best_match and best_score > 0.3:  # 降低最低匹配阈值
        print(f"选择最佳匹配: '{best_match['text']}' (得分: {best_score:.2f})")
        return best_match

    print(f"未找到匹配的选项")
    return None

def mark_field_checkbox(option):
    """通用函数：在匹配的选项前添加勾选标记"""
    try:
        cell = option['cell']
        option_text = option['text']
        original_line = option['original_line']

        print(f"正在标记选项: '{option_text}'")
        print(f"原始行文本: '{original_line}'")

        # 遍历单元格中的所有段落
        for paragraph in cell.paragraphs:
            paragraph_text = paragraph.text.strip()
            print(f"段落文本: '{paragraph_text}'")

            # 检查段落是否包含目标选项
            if option_text in paragraph_text and ('□' in paragraph_text or '☑' in paragraph_text or '✓' in paragraph_text):
                # 清空段落并重新构建
                paragraph.clear()

                # 分割段落文本为多行
                lines = paragraph_text.split('\n')
                for i, line in enumerate(lines):
                    line = line.strip()

                    if line:
                        # 使用精确匹配和替换特定选项
                        marked_line = mark_specific_option_in_line(line, option_text, original_line)

                        # 如果行内容发生了变化（即包含打勾符号），需要分别设置字体
                        if marked_line != line:
                            # 分别处理打勾符号和其他文本的字体
                            add_mixed_font_text(paragraph, marked_line)
                            print(f"已标记选项: '{marked_line}'")
                        else:
                            # 没有变化，使用默认字体
                            run = paragraph.add_run(marked_line)
                            try:
                                run.font.name = "宋体"
                                run.font.size = Pt(9.5)
                                if run._element.rPr is not None:
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
                            except:
                                pass  # 忽略字体设置错误

                    # 如果不是最后一行，添加换行
                    if i < len(lines) - 1:
                        paragraph.add_run('\n')

                return True

        print(f"警告: 未能在单元格中找到目标选项文本进行标记")
        return False

    except Exception as e:
        print(f"标记选项时出错: {e}")
        return False

def add_mixed_font_text(paragraph, text):
    """添加混合字体的文本，打勾符号使用小五号字体(9磅)，其他文本使用五号字体(9.5磅)"""
    try:
        i = 0
        while i < len(text):
            char = text[i]

            # 如果是打勾符号，使用小五号字体
            if char == '☑':
                run = paragraph.add_run(char)
                run.font.name = "楷体"
                run.font.size = Pt(9)  # 小五号字体
                if run._element.rPr is not None:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
                i += 1
            else:
                # 收集连续的非打勾符号字符
                normal_text = ""
                while i < len(text) and text[i] != '☑':
                    normal_text += text[i]
                    i += 1

                # 添加普通文本，使用五号字体
                if normal_text:
                    run = paragraph.add_run(normal_text)
                    run.font.name = "宋体"
                    run.font.size = Pt(9.5)  # 五号字体
                    if run._element.rPr is not None:
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    except Exception as e:
        print(f"设置混合字体时出错: {e}")
        # 如果出错，回退到普通方式
        run = paragraph.add_run(text)
        run.font.name = "宋体"
        run.font.size = Pt(9.5)
        if run._element.rPr is not None:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

def mark_specific_option_in_line(line, option_text, original_line):
    """在一行文本中精确标记特定选项，不影响其他选项"""
    try:
        # 如果行中包含目标选项文本
        if option_text in line:
            # 构建精确匹配的正则表达式模式
            # 匹配 "□选项文本" 或 "☑选项文本" 的模式
            escaped_option = re.escape(option_text)

            # 尝试多种匹配模式
            patterns = [
                f'□{escaped_option}(?=\\s|□|☑|$)',  # □选项文本（后面跟空格、其他复选框或行尾）
                f'☑{escaped_option}(?=\\s|□|☑|$)',  # ☑选项文本（后面跟空格、其他复选框或行尾）
                f'□\\s*{escaped_option}(?=\\s|□|☑|$)',  # □ 选项文本（中间可能有空格）
                f'☑\\s*{escaped_option}(?=\\s|□|☑|$)'   # ☑ 选项文本（中间可能有空格）
            ]

            marked_line = line
            for pattern in patterns:
                if re.search(pattern, line):
                    # 只替换匹配的部分，将□替换为☑
                    marked_line = re.sub(f'□(\\s*{escaped_option})', f'☑\\1', marked_line)
                    print(f"使用模式 '{pattern}' 匹配并标记: '{option_text}'")
                    break

            return marked_line

        return line

    except Exception as e:
        print(f"标记特定选项时出错: {e}")
        return line

def process_detection_timing_checkboxes(doc, timing_value):
    """处理检测时机复选框匹配和标记"""
    try:
        print(f"\n==== 开始处理检测时机复选框匹配 ====")
        print(f"检测时机值: '{timing_value}'")

        # 定义检测时机的匹配规则
        timing_patterns = {
            '焊后': ['焊后', '焊接后', '焊完后', '后焊', '焊后检测'],
            '焊前': ['焊前', '焊接前', '前焊', '焊前检测'],
            '打磨': ['打磨', '打磨后', '打磨前'],
            '热处理后': ['热处理后', '热处理', '热处理完成后'],
            '中间': ['中间', '中间检测', '过程中'],
            '最终': ['最终', '最终检测', '终检']
        }

        # 查找所有检测时机选项
        timing_options = find_field_options(doc, "检测时机", ["检测时机", "焊后", "焊前", "打磨", "热处理"])

        if not timing_options:
            print("警告: 未找到检测时机复选框选项，跳过复选框匹配")
            return False

        # 匹配检测时机值与选项
        matched_option = match_field_option(timing_value, timing_options, timing_patterns)

        if matched_option:
            # 标记匹配的选项
            success = mark_field_checkbox(matched_option)
            if success:
                print(f"成功标记检测时机选项: '{matched_option['text']}'")
                return True
            else:
                print(f"标记检测时机选项失败")
                return False
        else:
            print(f"未找到匹配的检测时机选项，可用选项:")
            for option in timing_options:
                print(f"  - {option['text']}")
            return False

    except Exception as e:
        print(f"处理检测时机复选框时出错: {e}")
        return False

def process_welding_method_checkboxes(doc, welding_method):
    """处理焊接方法复选框匹配和标记"""
    try:
        print(f"\n==== 开始处理焊接方法复选框匹配 ====")
        print(f"焊接方法值: '{welding_method}'")

        # 定义焊接方法的匹配规则
        welding_patterns = {
            'GTAW': ['GTAW', 'gtaw', 'TIG', 'tig', '氩弧焊'],
            'SMAW': ['SMAW', 'smaw', '手工电弧焊', '手弧焊'],
            'SAW': ['SAW', 'saw', '埋弧焊'],
            'GTAW+SMAW': ['GTAW+SMAW', 'gtaw+smaw', 'GTAW＋SMAW', 'TIG+SMAW'],
            'GTAW+SAW': ['GTAW+SAW', 'gtaw+saw', 'GTAW＋SAW', 'TIG+SAW']
        }

        # 查找所有焊接方法选项
        welding_options = find_field_options(doc, "焊接方法", ["焊接方法", "GTAW", "SMAW", "SAW"])

        if not welding_options:
            print("警告: 未找到焊接方法复选框选项，跳过复选框匹配")
            return False

        # 匹配焊接方法值与选项
        matched_option = match_field_option(welding_method, welding_options, welding_patterns)

        if matched_option:
            # 标记匹配的选项
            success = mark_field_checkbox(matched_option)
            if success:
                print(f"成功标记焊接方法选项: '{matched_option['text']}'")
                return True
            else:
                print(f"标记焊接方法选项失败")
                return False
        else:
            print(f"未找到匹配的焊接方法选项，可用选项:")
            for option in welding_options:
                print(f"  - {option['text']}")
            return False

    except Exception as e:
        print(f"处理焊接方法复选框时出错: {e}")
        return False

def process_quality_level_checkboxes(doc, quality_level):
    """处理合格级别复选框匹配和标记"""
    try:
        print(f"\n==== 开始处理合格级别复选框匹配 ====")
        print(f"合格级别值: '{quality_level}'")

        # 定义合格级别的匹配规则
        quality_patterns = {
            'Ⅰ': ['Ⅰ', 'I', '1', '一级', '一'],
            'Ⅱ': ['Ⅱ', 'II', '2', '二级', '二'],
            'Ⅲ': ['Ⅲ', 'III', '3', '三级', '三'],
            'Ⅳ': ['Ⅳ', 'IV', '4', '四级', '四']
        }

        # 查找所有合格级别选项
        quality_options = find_field_options(doc, "合格级别", ["合格级别", "Ⅰ", "Ⅱ", "Ⅲ", "Ⅳ", "级别"])

        if not quality_options:
            print("警告: 未找到合格级别复选框选项，跳过复选框匹配")
            return False

        # 匹配合格级别值与选项
        matched_option = match_field_option(quality_level, quality_options, quality_patterns)

        if matched_option:
            # 标记匹配的选项
            success = mark_field_checkbox(matched_option)
            if success:
                print(f"成功标记合格级别选项: '{matched_option['text']}'")
                return True
            else:
                print(f"标记合格级别选项失败")
                return False
        else:
            print(f"未找到匹配的合格级别选项，可用选项:")
            for option in quality_options:
                print(f"  - {option['text']}")
            return False

    except Exception as e:
        print(f"处理合格级别复选框时出错: {e}")
        return False

def process_detection_ratio_checkboxes(doc, detection_ratio):
    """处理检测比例复选框匹配和标记"""
    try:
        print(f"\n==== 开始处理检测比例复选框匹配 ====")
        print(f"检测比例值: '{detection_ratio}'")

        # 定义检测比例的匹配规则
        ratio_patterns = {
            '100%': ['100%', '100', '全部', '百分之百'],
            '50%': ['50%', '50', '百分之五十'],
            '20%': ['20%', '20', '百分之二十'],
            '10%': ['10%', '10', '百分之十'],
            '5%': ['5%', '5', '百分之五'],
            '1%': ['1%', '1', '百分之一']
        }

        # 查找所有检测比例选项
        ratio_options = find_field_options(doc, "检测比例", ["检测比例", "100%", "50%", "20%", "10%", "5%", "1%", "比例"])

        if not ratio_options:
            print("警告: 未找到检测比例复选框选项，跳过复选框匹配")
            return False

        # 匹配检测比例值与选项
        matched_option = match_field_option(detection_ratio, ratio_options, ratio_patterns)

        if matched_option:
            # 标记匹配的选项
            success = mark_field_checkbox(matched_option)
            if success:
                print(f"成功标记检测比例选项: '{matched_option['text']}'")
                return True
            else:
                print(f"标记检测比例选项失败")
                return False
        else:
            print(f"未找到匹配的检测比例选项，可用选项:")
            for option in ratio_options:
                print(f"  - {option['text']}")
            return False

    except Exception as e:
        print(f"处理检测比例复选框时出错: {e}")
        return False

def get_output_filename(word_template_path, order_number, ray_type):
    """根据Word模板路径、委托单编号和射线类型生成输出文件名
    
    Args:
        word_template_path: Word模板文档路径
        order_number: 委托单编号
        ray_type: 射线类型（"X射线"或"γ射线"）
    
    Returns:
        str: 输出文件名
    """
    # 获取模板文件名（不含路径和扩展名）
    template_name = os.path.splitext(os.path.basename(word_template_path))[0]
    # 射线类型标识
    ray_mark = "γ" if ray_type == "γ射线" else "X"
    # 生成输出文件名
    return f"{template_name}_{order_number}_{ray_mark}_生成结果.docx"

def process_excel_to_word(excel_path, word_template_path, output_path=None, 
                       project_name=None, entrusting_unit=None, 
                       operation_guide_number=None, contracting_unit=None, 
                       equipment_model=None):
    """将Excel数据填入Word文档
    
    Args:
        excel_path: Excel表格路径
        word_template_path: Word模板文档路径
        output_path: 输出Word文档路径（如果为None，将自动生成）
        project_name: 工程名称，用于替换文档中的"工程名称值"
        entrusting_unit: 委托单位，用于替换文档中的"委托单位值"
        operation_guide_number: 操作指导书编号，用于替换文档中的"操作指导书编号值"
        contracting_unit: 承包单位，用于替换文档中的"承包单位值"
        equipment_model: 设备型号，用于替换文档中的"设备型号值"
    
    Returns:
        bool: 处理是否成功
    """
    # 检查文件是否存在
    if not os.path.exists(excel_path):
        print(f"错误: Excel文件不存在: {excel_path}")
        return False
        
    if not os.path.exists(word_template_path):
        print(f"错误: Word模板文件不存在: {word_template_path}")
        return False
    
    # 创建输出目录
    if output_path is None:
        output_dir = os.path.join("生成器", "输出报告", "4_射线检测记录")
    else:
        output_dir = output_path
        
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            print(f"创建输出目录: {output_dir}")
        except Exception as e:
            print(f"错误: 无法创建输出目录: {e}")
            return False
    
    # 读取Excel数据
    print(f"正在读取Excel文件: {excel_path}")
    try:
        df = pd.read_excel(excel_path)
        print(f"成功读取Excel文件，共有{len(df)}行数据")
    except Exception as e:
        print(f"错误: 无法读取Excel文件: {e}")
        return False
    
    # X射线参数表
    xray_params_path = os.path.join("生成器", "Excel", "4_生成器X射线指导书模版.xlsx")
    try:
        if os.path.exists(xray_params_path):
            xray_params_df = pd.read_excel(xray_params_path, sheet_name="曝光参数")
            print(f"成功读取X射线参数表，共有{len(xray_params_df)}行数据")
            # 打印X射线参数表的列名
            print(f"X射线参数表列名: {list(xray_params_df.columns)}")
        else:
            print(f"警告: X射线参数表不存在: {xray_params_path}，将不进行X射线参数匹配")
            xray_params_df = None
    except Exception as e:
        print(f"错误: 无法读取X射线参数表: {e}")
        xray_params_df = None
    
    # γ射线参数表
    gamma_params_path = os.path.join("生成器", "Excel", "4_生成器γ射线指导书模版.xlsx")
    try:
        if os.path.exists(gamma_params_path):
            gamma_params_df = pd.read_excel(gamma_params_path, sheet_name="曝光参数")
            print(f"成功读取γ射线参数表，共有{len(gamma_params_df)}行数据")
            # 打印γ射线参数表的列名
            print(f"γ射线参数表列名: {list(gamma_params_df.columns)}")
        else:
            print(f"警告: γ射线参数表不存在: {gamma_params_path}，将不进行γ射线参数匹配")
            gamma_params_df = None
    except Exception as e:
        print(f"错误: 无法读取γ射线参数表: {e}")
        gamma_params_df = None
    
    # 打印所有列名，帮助调试
    print(f"Excel表格列名: {list(df.columns)}")
    
    # 定义需要查找的列关键字
    column_keywords = {
        '完成日期': '完成日期',
        '委托单编号': '委托单编号',
        '检件编号': '检件编号',
        '焊口编号': '焊口编号',
        '焊工号': '焊工号',
        '合格级别': '合格级别',
        '检测比例': '检测比例',
        'γ射线': 'γ射线',
        '焊接方法': '焊接方法',
        '检测时机': '检测时机',
        '规格': '规格'
    }
    
    # 查找每个关键字对应的实际列名
    column_mapping = {}
    missing_columns = []
    
    for key, keyword in column_keywords.items():
        col_name = find_column_with_keyword(df, keyword)
        if col_name:
            column_mapping[key] = col_name
            print(f"找到列: '{key}' -> '{col_name}'")
        else:
            missing_columns.append(key)
            
    if missing_columns:
        print(f"警告: 未找到以下列: {', '.join(missing_columns)}")
        # 尝试使用列位置
        possible_columns = {
            '完成日期': 'B', 
            '委托单编号': 'C', 
            '检件编号': 'D', 
            '焊口编号': 'E', 
            '焊工号': 'F',
            '合格级别': 'I',
            '检测比例': 'J',
            'γ射线': 'P',
            '焊接方法': 'R',
            '检测时机': 'V',
            '规格': 'G'
        }
        
        for key in missing_columns:
            col_letter = possible_columns.get(key)
            if col_letter:
                col_idx = ord(col_letter) - ord('A')
                if col_idx < len(df.columns):
                    column_mapping[key] = df.columns[col_idx]
                    print(f"使用列位置找到: '{key}' -> '{df.columns[col_idx]}'")
    
    # 如果缺少关键列，则无法继续处理
    if '委托单编号' not in column_mapping:
        print("错误: 无法找到委托单编号列")
        return False
    
    if 'γ射线' not in column_mapping:
        print("警告: 无法找到γ射线列，将默认所有记录为X射线")
        # 增加一个全为空的列作为γ射线列
        df['γ射线'] = None
        column_mapping['γ射线'] = 'γ射线'
    
    # 根据委托单编号和射线类型分组数据
    groups = []
    order_numbers = df[column_mapping['委托单编号']].dropna().unique()
    
    for order_number in order_numbers:
        # 获取该委托单编号的所有数据
        order_df = df[df[column_mapping['委托单编号']] == order_number]
        
        # 获取该委托单编号下的所有射线类型
        ray_types = order_df[column_mapping['γ射线']].dropna().unique()
        
        # 如果没有明确的γ射线值，则视为X射线
        if len(ray_types) == 0:
            groups.append({
                'order_number': order_number,
                'ray_type': 'X射线',  # X射线表示为处理逻辑，但射源种类值会设置为空
                'data': order_df
            })
            print(f"委托单编号 {order_number} 没有明确的射线类型，处理为X射线（射源种类值为空）")
        else:
            # 有γ射线值的处理为γ射线
            for ray_type in ray_types:
                ray_df = order_df[order_df[column_mapping['γ射线']] == ray_type]
                groups.append({
                    'order_number': order_number,
                    'ray_type': 'γ射线',
                    'data': ray_df
                })
                print(f"委托单编号 {order_number} 的射线类型 γ射线 有 {len(ray_df)} 条记录")
            
            # 没有γ射线值的处理为X射线
            x_ray_df = order_df[order_df[column_mapping['γ射线']].isna()]
            if len(x_ray_df) > 0:
                groups.append({
                    'order_number': order_number,
                    'ray_type': 'X射线',  # X射线表示为处理逻辑，但射源种类值会设置为空
                    'data': x_ray_df
                })
                print(f"委托单编号 {order_number} 的射线类型 X射线 有 {len(x_ray_df)} 条记录")
    
    print(f"共有 {len(groups)} 个组合需要生成报告")
    
    # 处理每个分组
    success_count = 0
    error_count = 0
    
    for group in groups:
        order_number = group['order_number']
        ray_type = group['ray_type']
        group_df = group['data']
        
        print(f"\n{'='*50}")
        print(f"处理委托单编号: {order_number}, 射线类型: {ray_type}")
        print(f"{'='*50}")
        
        try:
            # 为该分组生成输出文件名
            output_filename = get_output_filename(word_template_path, order_number, ray_type)
            report_output_path = os.path.join(output_dir, output_filename)
            print(f"输出文件路径: {report_output_path}")
            
            # 打开Word文档
            print(f"正在处理Word文档: {word_template_path}")
            
            try:
                # 每次处理新的组合时，重新从模板创建文档对象
                # 这确保了每个组合都会生成一个独立的文档
                doc = Document(word_template_path)
                print(f"成功从模板创建新文档")
            except Exception as e:
                print(f"无法打开Word文档: {e}")
                error_count += 1
                continue
            
            # 填充文档的其余部分将在这里添加...
            
            # 1) 获取该组数据中最晚的完成日期
            date_col = column_mapping.get('完成日期')
            if date_col:
                # 确保日期列是日期类型
                group_df[date_col] = pd.to_datetime(group_df[date_col], errors='coerce')
                latest_date = group_df[date_col].max()
                
                if pd.isna(latest_date):
                    print(f"警告: 委托单编号 {order_number} 没有有效的完成日期")
                    year, month, day = datetime.now().year, datetime.now().month, datetime.now().day
                else:
                    # 将日期转换为年、月、日
                    year = latest_date.year
                    month = latest_date.month
                    day = latest_date.day
                    print(f"找到最晚完成日期: {year}年{month}月{day}日")
            else:
                print("警告: 未找到完成日期列")
                year, month, day = datetime.now().year, datetime.now().month, datetime.now().day
            
            # 获取相关数据
            inspection_numbers = group_df[column_mapping.get('检件编号')].dropna().tolist() if '检件编号' in column_mapping else []
            weld_numbers = group_df[column_mapping.get('焊口编号')].dropna().tolist() if '焊口编号' in column_mapping else []
            welder_numbers = group_df[column_mapping.get('焊工号')].dropna().tolist() if '焊工号' in column_mapping else []
            # 对规格数据进行去重处理
            specifications = list(group_df[column_mapping.get('规格')].dropna().unique()) if '规格' in column_mapping else []
            print(f"规格列去重前数量: {len(group_df[column_mapping.get('规格')].dropna().tolist() if '规格' in column_mapping else [])}")
            print(f"规格列去重后数量: {len(specifications)}")
            print(f"去重后规格列数据: {specifications}")
            
            # 获取分组的第一个值
            grade_level = ""
            if '合格级别' in column_mapping:
                grade_levels = group_df[column_mapping['合格级别']].dropna().tolist()
                if grade_levels:
                    grade_level = grade_levels[0]
                    print(f"找到合格级别: {grade_level}")
            
            inspection_ratio = ""
            if '检测比例' in column_mapping:
                ratios = group_df[column_mapping['检测比例']].dropna().tolist()
                if ratios:
                    # 转换为百分数格式
                    try:
                        ratio_value = float(ratios[0])
                        inspection_ratio = f"{ratio_value*100:.0f}%"
                    except (ValueError, TypeError):
                        inspection_ratio = str(ratios[0])
                    print(f"找到检测比例: {inspection_ratio}")
            
            welding_method = ""
            if '焊接方法' in column_mapping:
                methods = group_df[column_mapping['焊接方法']].dropna().tolist()
                if methods:
                    welding_method = methods[0]
                    print(f"找到焊接方法: {welding_method}")
            
            inspection_time = ""
            if '检测时机' in column_mapping:
                times = group_df[column_mapping['检测时机']].dropna().tolist()
                if times:
                    inspection_time = times[0]
                    print(f"找到检测时机: {inspection_time}")
            
            # 根据射线类型设置相关参数
            if ray_type == "γ射线":
                focus_size = "3*3"
                lead_screen = "柯达0.1*2"
                film_grade = "柯达MX125"
                ray_source = "γ射线"
            else:  # X射线
                focus_size = "2.5*2.5"
                lead_screen = "0.03*2"
                film_grade = "锐科R400"
                ray_source = "X射线"  # 当γ射线的值为空时，射源种类值也设置为X射线
            
            print(f"射线类型: {ray_type}, 焦点尺寸: {focus_size}, 铅增感屏: {lead_screen}, 胶片等级: {film_grade}, 射源种类值: {ray_source}")
            
            # 替换文档中的值
            print("\n==== 开始替换文档中的值 ====")
            
            # 将EPKJ拼接委托单编号代替委托单编号值
            committee_order = f"EPKJ-{order_number}"
            replaced = False
            
            # 遍历所有段落，替换关键词
            for paragraph in doc.paragraphs:
                if "委托单编号值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("委托单编号值", committee_order)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'委托单编号值'替换为'{committee_order}'")
                    replaced = True

                if "射源种类值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("射源种类值", ray_source)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'射源种类值'替换为'{ray_source}'")
                    replaced = True

                if "合格级别值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("合格级别值", grade_level)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'合格级别值'替换为'{grade_level}'")
                    replaced = True

                if "检测比例值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("检测比例值", inspection_ratio)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'检测比例值'替换为'{inspection_ratio}'")
                    replaced = True
                
                if "焊接方法值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("焊接方法值", welding_method)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'焊接方法值'替换为'{welding_method}'")
                    replaced = True

                if "检测时机值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("检测时机值", inspection_time)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'检测时机值'替换为'{inspection_time}'")
                    replaced = True

                if "焦点尺寸值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("焦点尺寸值", focus_size)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'焦点尺寸值'替换为'{focus_size}'")
                    replaced = True

                if "铅增感屏值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("铅增感屏值", lead_screen)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'铅增感屏值'替换为'{lead_screen}'")
                    replaced = True

                if "胶片等级值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("胶片等级值", film_grade)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'胶片等级值'替换为'{film_grade}'")
                    replaced = True
                
                # 新增的5个参数替换
                if "工程名称值" in paragraph.text and project_name:
                    paragraph.text = paragraph.text.replace("工程名称值", project_name)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'工程名称值'替换为'{project_name}'")
                    replaced = True

                if "委托单位值" in paragraph.text and entrusting_unit:
                    paragraph.text = paragraph.text.replace("委托单位值", entrusting_unit)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'委托单位值'替换为'{entrusting_unit}'")
                    replaced = True

                if "操作指导书编号值" in paragraph.text and operation_guide_number:
                    paragraph.text = paragraph.text.replace("操作指导书编号值", operation_guide_number)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'操作指导书编号值'替换为'{operation_guide_number}'")
                    replaced = True

                if "承包单位值" in paragraph.text and contracting_unit:
                    paragraph.text = paragraph.text.replace("承包单位值", contracting_unit)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'承包单位值'替换为'{contracting_unit}'")
                    replaced = True

                if "设备型号值" in paragraph.text and equipment_model:
                    paragraph.text = paragraph.text.replace("设备型号值", equipment_model)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'设备型号值'替换为'{equipment_model}'")
                    replaced = True
            
            # 遍历表格中的单元格，替换关键词
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        for paragraph in cell.paragraphs:
                            if "委托单编号值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("委托单编号值", committee_order)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'委托单编号值'替换为'{committee_order}'")
                                replaced = True

                            if "射源种类值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("射源种类值", ray_source)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'射源种类值'替换为'{ray_source}'")
                                replaced = True

                            if "合格级别值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("合格级别值", grade_level)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'合格级别值'替换为'{grade_level}'")
                                replaced = True

                            if "检测比例值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("检测比例值", inspection_ratio)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'检测比例值'替换为'{inspection_ratio}'")
                                replaced = True
                            
                            if "焊接方法值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("焊接方法值", welding_method)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'焊接方法值'替换为'{welding_method}'")
                                replaced = True

                            if "检测时机值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("检测时机值", inspection_time)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'检测时机值'替换为'{inspection_time}'")
                                replaced = True

                            if "焦点尺寸值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("焦点尺寸值", focus_size)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'焦点尺寸值'替换为'{focus_size}'")
                                replaced = True

                            if "铅增感屏值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("铅增感屏值", lead_screen)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'铅增感屏值'替换为'{lead_screen}'")
                                replaced = True

                            if "胶片等级值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("胶片等级值", film_grade)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'胶片等级值'替换为'{film_grade}'")
                                replaced = True
                            
                            # 新增的5个参数替换（表格单元格）
                            if "工程名称值" in paragraph.text and project_name:
                                paragraph.text = paragraph.text.replace("工程名称值", project_name)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'工程名称值'替换为'{project_name}'")
                                replaced = True

                            if "委托单位值" in paragraph.text and entrusting_unit:
                                paragraph.text = paragraph.text.replace("委托单位值", entrusting_unit)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'委托单位值'替换为'{entrusting_unit}'")
                                replaced = True

                            if "操作指导书编号值" in paragraph.text and operation_guide_number:
                                paragraph.text = paragraph.text.replace("操作指导书编号值", operation_guide_number)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'操作指导书编号值'替换为'{operation_guide_number}'")
                                replaced = True

                            if "承包单位值" in paragraph.text and contracting_unit:
                                paragraph.text = paragraph.text.replace("承包单位值", contracting_unit)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'承包单位值'替换为'{contracting_unit}'")
                                replaced = True

                            if "设备型号值" in paragraph.text and equipment_model:
                                paragraph.text = paragraph.text.replace("设备型号值", equipment_model)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'设备型号值'替换为'{equipment_model}'")
                                replaced = True
            
            if not replaced:
                print("警告: 未找到需要替换的关键词，可能需要检查Word模板中的占位符命名。")
            
            # 填写日期（洗片人、拍片人、审核人）
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        # 处理日期
                        date_patterns = ["洗片人", "拍片人", "审核人"]
                        for pattern in date_patterns:
                            if pattern in cell.text:
                                print(f"找到{pattern}单元格: 表格行{i+1}, 列{j+1}")

                                # 检查单元格中的所有段落
                                date_found = False
                                for paragraph in cell.paragraphs:
                                    if "年" in paragraph.text and "月" in paragraph.text and "日" in paragraph.text:
                                        print(f"找到日期段落: {paragraph.text}")

                                        # 清空段落内容，重新构建带格式的日期
                                        paragraph.clear()

                                        # 添加年份数字（楷体五号）
                                        run_year = paragraph.add_run(str(year))
                                        run_year.font.name = "楷体"
                                        run_year.font.size = Pt(10.5)
                                        run_year._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")

                                        # 添加"年"字（保持原格式）
                                        paragraph.add_run("年")

                                        # 添加月份数字（楷体五号）
                                        run_month = paragraph.add_run(str(month))
                                        run_month.font.name = "楷体"
                                        run_month.font.size = Pt(10.5)
                                        run_month._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")

                                        # 添加"月"字（保持原格式）
                                        paragraph.add_run("月")

                                        # 添加日期数字（楷体五号）
                                        run_day = paragraph.add_run(str(day))
                                        run_day.font.name = "楷体"
                                        run_day.font.size = Pt(10.5)
                                        run_day._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")

                                        # 添加"日"字（保持原格式）
                                        paragraph.add_run("日")

                                        date_found = True
                                        print(f"已更新{pattern}日期为 {year}年{month}月{day}日")
                                        break

                                # 如果没有找到日期段落，尝试创建新段落
                                if not date_found:
                                    print(f"未在{pattern}单元格中找到日期段落，尝试添加")
                                    # 添加新段落并设置格式
                                    p = cell.add_paragraph()

                                    # 添加年份数字（楷体五号）
                                    run_year = p.add_run(str(year))
                                    run_year.font.name = "楷体"
                                    run_year.font.size = Pt(10.5)
                                    run_year._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")

                                    # 添加"年"字（保持原格式）
                                    p.add_run("年")

                                    # 添加月份数字（楷体五号）
                                    run_month = p.add_run(str(month))
                                    run_month.font.name = "楷体"
                                    run_month.font.size = Pt(10.5)
                                    run_month._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")

                                    # 添加"月"字（保持原格式）
                                    p.add_run("月")

                                    # 添加日期数字（楷体五号）
                                    run_day = p.add_run(str(day))
                                    run_day.font.name = "楷体"
                                    run_day.font.size = Pt(10.5)
                                    run_day._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")

                                    # 添加"日"字（保持原格式）
                                    p.add_run("日")

                                    print(f"已添加{pattern}日期: {year}年{month}月{day}日")
            
            # 查找表头行，确定各列的位置
            for table in doc.tables:
                column_indices = {}
                header_row_index = -1
                
                # 添加完整的表格内容打印，帮助调试
                # print("\n调试：打印表格内容以找到规格列")
                for i, row in enumerate(table.rows):
                    row_text = []
                    for j, cell in enumerate(row.cells):
                        row_text.append(f"[{j}]'{cell.text}'")
                    # if len(row_text) > 0:  # 只打印非空行
                    #     print(f"行 {i}: {', '.join(row_text)}")
                
                # 查找规格表头(mm×mm)位于第10行左右，检件编号等位于第18行左右
                spec_column_index = -1
                for i in range(9, 12):  # 在第9-11行范围内查找规格列
                    if i < len(table.rows):
                        for j, cell in enumerate(table.rows[i].cells):
                            cell_text = cell.text.strip()
                            if "检件规格" in cell_text and ("mm×mm" in cell_text or "mm*mm" in cell_text):
                                spec_column_index = j
                                print(f"找到规格列（透照参数表格）：第{i+1}行，第{j+1}列，文本：{cell_text}")
                                break
                
                # 查找包含"检件编号"、"焊缝编号"、"焊工号"的行
                for i, row in enumerate(table.rows):
                    header_found = False
                    for j, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        
                        # 添加更详细的调试信息，输出表格单元格文本内容
                        if "检件规格" in cell_text or "规格" in cell_text:
                            print(f"找到可能的规格列： 行 {i+1}, 列 {j+1}, 文本: '{cell_text}'")
                        
                        if "检件编号" in cell_text:
                            column_indices["检件编号"] = j
                            header_row_index = i
                            header_found = True
                            print(f"找到检件编号列: 行 {i+1}, 列 {j+1}, 文本: '{cell_text}'")
                        elif "焊缝编号" in cell_text or "焊口编号" in cell_text:
                            column_indices["焊口编号"] = j
                            header_found = True
                            print(f"找到焊缝编号列: 行 {i+1}, 列 {j+1}, 文本: '{cell_text}'")
                        elif "焊工号" in cell_text:
                            column_indices["焊工号"] = j
                            header_found = True
                            print(f"找到焊工号列: 行 {i+1}, 列 {j+1}, 文本: '{cell_text}'")
                        elif "备注" in cell_text:
                            column_indices["备注"] = j
                            header_found = True
                            print(f"找到备注列: 行 {i+1}, 列 {j+1}, 文本: '{cell_text}'")
                        elif "检件规格" in cell_text or "规格" in cell_text or "检件规格(mm×mm)" in cell_text or "检件规格(mm*mm)" in cell_text or "检件规格(mm" in cell_text:
                            column_indices["检件规格"] = j
                            header_found = True
                            print(f"找到检件规格列: 行 {i+1}, 列 {j+1}, 文本: '{cell_text}'")
                        elif "透照参数序号" in cell_text:
                            column_indices["透照参数序号"] = j
                            header_found = True
                            print(f"找到透照参数序号列: 行 {i+1}, 列 {j+1}, 文本: '{cell_text}'")
                    
                    if header_found and header_row_index >= 0:
                        print(f"找到表头行: 第{header_row_index+1}行")
                        
                        # 将"焊口编号"的键名更新为"焊缝编号"以保持一致性
                        if "焊口编号" in column_indices:
                            column_indices["焊缝编号"] = column_indices.pop("焊口编号")
                            
                        print(f"列索引: {column_indices}")
                        break
                
                # 如果找到表头行，处理数据填充
                if header_row_index >= 0 and column_indices:
                    # 获取可用于填充数据的行
                    data_rows = []
                    for i in range(header_row_index + 1, len(table.rows)):
                        if i < len(table.rows):
                            # 检查是否是空行或包含特殊标记的行
                            if "以下空白" in table.rows[i].cells[0].text if len(table.rows[i].cells) > 0 else False:
                                print(f"找到'以下空白'行: 第{i+1}行")
                                break
                            # 添加可用于填充数据的行
                            data_rows.append(i)
                    
                    print(f"找到{len(data_rows)}行可用于填充数据")
                    
                    # 确定需要填充的数据行数
                    data_count = len(group_df)
                    print(f"需要填充{data_count}行数据")
                    
                    # 如果Word表格中的行数不足，需要添加新行
                    rows_needed = data_count - len(data_rows)
                    if rows_needed > 0:
                        print(f"需要添加{rows_needed}行到表格中")
                        # 找到最后一行的索引
                        last_row_idx = data_rows[-1] if data_rows else header_row_index
                        
                        # 添加新行
                        for _ in range(rows_needed):
                            # 在表格末尾添加一行
                            new_row = table.add_row()
                            data_rows.append(len(table.rows) - 1)  # 添加新行的索引
                    
                    # 处理每一行数据
                    for i in range(data_count):
                        if i < len(data_rows):
                            row_idx = data_rows[i]
                            row = table.rows[row_idx]
                            
                            # 1. 填写检件编号
                            if "检件编号" in column_indices and i < len(inspection_numbers):
                                col_idx = column_indices["检件编号"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    if cell.paragraphs:
                                        cell.paragraphs[0].text = str(inspection_numbers[i])
                                        set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                        print(f"已更新第{row_idx+1}行检件编号: {inspection_numbers[i]}")

                            # 2. 填写焊缝编号
                            if "焊缝编号" in column_indices and i < len(weld_numbers):
                                col_idx = column_indices["焊缝编号"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    if cell.paragraphs:
                                        cell.paragraphs[0].text = str(weld_numbers[i])
                                        set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                        print(f"已更新第{row_idx+1}行焊缝编号: {weld_numbers[i]}")

                            # 3. 填写焊工号
                            if "焊工号" in column_indices and i < len(welder_numbers):
                                col_idx = column_indices["焊工号"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    if cell.paragraphs:
                                        cell.paragraphs[0].text = str(welder_numbers[i])
                                        set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                        print(f"已更新第{row_idx+1}行焊工号: {welder_numbers[i]}")
                            
                            # 4. 填写备注（填入完成日期）
                            if "备注" in column_indices and date_col and i < len(group_df):
                                col_idx = column_indices["备注"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    # 获取对应行的完成日期
                                    if not pd.isna(group_df[date_col].iloc[i]):
                                        date_value = group_df[date_col].iloc[i]
                                        if isinstance(date_value, pd.Timestamp):
                                            # 如果是日期类型，格式化为字符串
                                            formatted_date = date_value.strftime("%Y年%m月%d日")
                                        else:
                                            # 如果不是日期类型，直接转为字符串
                                            formatted_date = str(date_value)
                                        
                                        if cell.paragraphs:
                                            cell.paragraphs[0].text = formatted_date
                                            set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                            print(f"已更新第{row_idx+1}行备注（完成日期）: {formatted_date}")

                            # 5. 填写透照参数序号（规格数量）
                            if "透照参数序号" in column_indices:
                                col_idx = column_indices["透照参数序号"]
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    if cell.paragraphs:
                                        # 获取规格去重后的数量
                                        spec_count = len(specifications)
                                        # 直接填写规格总数，不再根据行号判断
                                        param_index = spec_count
                                        cell.paragraphs[0].text = str(param_index)
                                        set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                        print(f"已更新第{row_idx+1}行透照参数序号: {param_index}")
            
                    # 如果找到了规格列，在透照参数表中填写规格信息
                    if spec_column_index >= 0 and len(specifications) > 0:
                        # 透照参数表格一般在第10-15行，我们从第10行开始填充规格数据
                        start_row = 10
                        print(f"开始在透照参数表中填充去重后的{len(specifications)}种规格数据")
                        
                        # 清空现有规格数据
                        for i in range(5):  # 最多清空5行
                            if start_row + i < len(table.rows) and spec_column_index < len(table.rows[start_row + i].cells):
                                cell = table.rows[start_row + i].cells[spec_column_index]
                                if cell.paragraphs:
                                    cell.paragraphs[0].text = ""
                        
                        # 填入去重后的规格数据
                        for i in range(min(len(specifications), 5)):  # 最多填充5行
                            if start_row + i < len(table.rows) and spec_column_index < len(table.rows[start_row + i].cells):
                                cell = table.rows[start_row + i].cells[spec_column_index]
                                if cell.paragraphs:
                                    cell.paragraphs[0].text = str(specifications[i])
                                    set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                    print(f"已更新透照参数表第{start_row+i+1}行检件规格(mm×mm): {specifications[i]}")
                                    
                                    # 如果是X射线模式，则查找并填充X射线参数
                                    if ray_type == "X射线" and xray_params_df is not None:
                                        # 查找与规格匹配的X射线参数
                                        xray_params = find_xray_params_by_spec(xray_params_df, specifications[i])
                                        if xray_params:
                                            # 填充各项X射线参数
                                            param_columns = {
                                                '透照方式': 10,  # 透照方式列索引
                                                '焦距': 15,      # 焦距列索引
                                                '管电压源能量': 28, # 管电压源能量列索引
                                                '管电流源活度': 35, # 管电流源活度列索引
                                                '曝光时间': 40,   # 曝光时间列索引
                                                '有效片长': 23    # 有效片长列索引
                                            }
                                            
                                            # 填入对应参数
                                            for param_name, col_idx in param_columns.items():
                                                if param_name in xray_params and col_idx < len(table.rows[start_row + i].cells):
                                                    value = xray_params[param_name]
                                                    cell = table.rows[start_row + i].cells[col_idx]
                                                    if cell.paragraphs:
                                                        cell.paragraphs[0].text = str(value)
                                                        set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                                        print(f"已更新第{start_row+i+1}行{param_name}: {value}")
                                        else:
                                            print(f"未找到规格 {specifications[i]} 的X射线参数")
                                    
                                    # 如果是γ射线模式，则查找并填充γ射线参数
                                    elif ray_type == "γ射线" and gamma_params_df is not None:
                                        # 查找与规格匹配的γ射线参数
                                        gamma_params = find_gamma_params_by_spec(gamma_params_df, specifications[i])
                                        if gamma_params:
                                            # 填充各项γ射线参数
                                            param_columns = {
                                                '透照方式': 10,  # 透照方式列索引
                                                '焦距': 15,      # 焦距列索引
                                                '管电流源活度': 35, # 管电流源活度列索引 (对应Excel中的源强)
                                                '有效片长': 23    # 有效片长列索引 (对应Excel中的一次透照长度)
                                            }
                                            
                                            # 填入对应参数
                                            for param_name, col_idx in param_columns.items():
                                                if param_name in gamma_params and col_idx < len(table.rows[start_row + i].cells):
                                                    value = gamma_params[param_name]
                                                    cell = table.rows[start_row + i].cells[col_idx]
                                                    if cell.paragraphs:
                                                        cell.paragraphs[0].text = str(value)
                                                        set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                                        print(f"已更新第{start_row+i+1}行{param_name}: {value}")
                                        else:
                                            print(f"未找到规格 {specifications[i]} 的γ射线参数")
                    
                    if ray_type == "X射线":
                        print("X射线参数处理完成")
                    else:
                        print("γ射线参数处理完成")

            # 处理检测时机复选框匹配和标记
            checkbox_success = process_detection_timing_checkboxes(doc, inspection_time)
            if checkbox_success:
                print("检测时机复选框处理完成")
            else:
                print("检测时机复选框处理失败，已保留原有文本替换")

            # 处理焊接方法复选框匹配和标记
            welding_checkbox_success = process_welding_method_checkboxes(doc, welding_method)
            if welding_checkbox_success:
                print("焊接方法复选框处理完成")
            else:
                print("焊接方法复选框处理失败，已保留原有文本替换")

            # 处理合格级别复选框匹配和标记
            quality_checkbox_success = process_quality_level_checkboxes(doc, grade_level)
            if quality_checkbox_success:
                print("合格级别复选框处理完成")
            else:
                print("合格级别复选框处理失败，已保留原有文本替换")

            # 处理检测比例复选框匹配和标记
            ratio_checkbox_success = process_detection_ratio_checkboxes(doc, inspection_ratio)
            if ratio_checkbox_success:
                print("检测比例复选框处理完成")
            else:
                print("检测比例复选框处理失败，已保留原有文本替换")

            print("==== 文档填充完成 ====\n")
            
            # 保存文档
            try:
                print(f"\n正在保存文档到: {report_output_path}")
                doc.save(report_output_path)
                print(f"文档已成功保存至: {report_output_path}")
                success_count += 1
            except Exception as e:
                print(f"错误: 无法保存文档: {e}")
                error_count += 1
                
        except Exception as e:
            print(f"错误: 处理委托单编号 {order_number} 和射线类型 {ray_type} 时出错: {e}")
            error_count += 1
    
    print(f"\n处理完成: 共处理{len(groups)}个组合，成功生成{success_count}份报告，失败{error_count}份")
    if error_count > 0:
        print(f"警告: 有{error_count}个组合处理失败，请检查日志")
    
    return success_count > 0

def find_xray_params_by_spec(xray_params_df, spec_value):
    """
    根据检件规格在X射线参数表中查找对应的参数
    
    Args:
        xray_params_df: X射线参数表DataFrame
        spec_value: 检件规格值
    
    Returns:
        dict: 包含匹配到的各项参数值的字典，若未找到匹配则返回空字典
    """
    if xray_params_df is None or len(xray_params_df) == 0:
        print(f"警告: X射线参数表为空")
        return {}
    
    # 打印所有列以便调试
    print(f"X射线参数表所有列: {xray_params_df.columns.tolist()}")
    
    # 检查列名是否符合预期
    required_columns = ['检件规格', '透照方式', '焦 距（mm）', '源强（管电压）', '源活', '曝光时间', '一次透照长度(mm)']
    missing_columns = [col for col in required_columns if col not in xray_params_df.columns]
    if missing_columns:
        print(f"警告: X射线参数表缺少所需列: {missing_columns}")
        return {}
    
    # 清理规格值以便更准确匹配
    cleaned_spec = re.sub(r'[×*\s]', '', str(spec_value)).lower()
    print(f"查找规格值: '{spec_value}' (清理后: '{cleaned_spec}')")
    
    # 在检件规格列中查找匹配的行
    for idx, row in xray_params_df.iterrows():
        row_spec = str(row['检件规格'])
        cleaned_row_spec = re.sub(r'[×*\s]', '', row_spec).lower()
        
        # 如果找到完全匹配或部分匹配
        if cleaned_spec == cleaned_row_spec or cleaned_spec in cleaned_row_spec or cleaned_row_spec in cleaned_spec:
            print(f"在X射线参数表中找到匹配的规格: '{row_spec}'")
            
            # 返回匹配行的所有所需参数
            params = {
                '透照方式': row['透照方式'],
                '焦距': row['焦 距（mm）'],
                '管电压源能量': row['源强（管电压）'],
                '管电流源活度': row['源活'],
                '曝光时间': row['曝光时间'],
                '有效片长': row['一次透照长度(mm)']
            }
            
            print(f"匹配的X射线参数: {params}")
            return params
    
    print(f"警告: 在X射线参数表中未找到匹配的规格: '{spec_value}'")
    return {}

def find_gamma_params_by_spec(gamma_params_df, spec_value):
    """
    根据检件规格在γ射线参数表中查找对应的参数
    
    Args:
        gamma_params_df: γ射线参数表DataFrame
        spec_value: 检件规格值
    
    Returns:
        dict: 包含匹配到的各项参数值的字典，若未找到匹配则返回空字典
    """
    if gamma_params_df is None or len(gamma_params_df) == 0:
        print(f"警告: γ射线参数表为空")
        return {}
    
    # 打印所有列以便调试
    print(f"γ射线参数表所有列: {gamma_params_df.columns.tolist()}")
    
    # 检查列名是否符合预期
    required_columns = ['检件规格', '透照方式', '焦 距（mm）', '源强（管电压）', '一次透照长度(mm)']
    missing_columns = [col for col in required_columns if col not in gamma_params_df.columns]
    if missing_columns:
        print(f"警告: γ射线参数表缺少所需列: {missing_columns}")
        return {}
    
    # 清理规格值以便更准确匹配
    cleaned_spec = re.sub(r'[×*\s]', '', str(spec_value)).lower()
    print(f"查找规格值: '{spec_value}' (清理后: '{cleaned_spec}')")
    
    # 在检件规格列中查找匹配的行
    for idx, row in gamma_params_df.iterrows():
        row_spec = str(row['检件规格'])
        cleaned_row_spec = re.sub(r'[×*\s]', '', row_spec).lower()
        
        # 如果找到完全匹配或部分匹配
        if cleaned_spec == cleaned_row_spec or cleaned_spec in cleaned_row_spec or cleaned_row_spec in cleaned_spec:
            print(f"在γ射线参数表中找到匹配的规格: '{row_spec}'")
            
            # 返回匹配行的所有所需参数
            params = {
                '透照方式': row['透照方式'],
                '焦距': row['焦 距（mm）'],
                '管电流源活度': row['源强（管电压）'],  # 源强（管电压）列实际是源活度值
                '有效片长': row['一次透照长度(mm)']
            }
            
            print(f"匹配的γ射线参数: {params}")
            return params
    
    print(f"警告: 在γ射线参数表中未找到匹配的规格: '{spec_value}'")
    return {}

def main():
    # 创建命令行参数解析器
    parser = argparse.ArgumentParser(description='将Excel数据填入Word文档')
    parser.add_argument('-e', '--excel', default="生成器/Excel/4_生成器台账-射线检测记录.xlsx", 
                        help='Excel表格路径 (默认: 生成器/Excel/4_生成器台账-射线检测记录.xlsx)')
    parser.add_argument('-w', '--word', default="生成器/word/4_射线检测记录.docx", 
                        help='Word模板文档路径 (默认: 生成器/word/4_射线检测记录.docx)')
    parser.add_argument('-o', '--output', 
                        help='输出目录 (可选，默认为"生成器/输出报告/4_射线检测记录"目录)')
    
    # 新增5个命令行参数
    parser.add_argument('-p', '--project', 
                        help='工程名称，用于替换文档中的"工程名称值"')
    parser.add_argument('-u', '--entrusting_unit', 
                        help='委托单位，用于替换文档中的"委托单位值"')
    parser.add_argument('-g', '--guide_number', 
                        help='操作指导书编号，用于替换文档中的"操作指导书编号值"')
    parser.add_argument('-c', '--contracting_unit', 
                        help='承包单位，用于替换文档中的"承包单位值"')
    parser.add_argument('-m', '--equipment_model', 
                        help='设备型号，用于替换文档中的"设备型号值"')
    
    # 解析命令行参数
    args = parser.parse_args()
    
    # 处理Excel到Word的转换
    success = process_excel_to_word(
        args.excel, args.word, args.output,
        args.project, args.entrusting_unit,
        args.guide_number, args.contracting_unit,
        args.equipment_model
    )
    
    # 返回状态码
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
