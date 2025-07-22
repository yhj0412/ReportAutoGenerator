import pandas as pd
import os
import sys
import argparse
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import re
import gc
import logging
from dataclasses import dataclass
from typing import List, Dict, Optional, Tuple

# 动态表格扩展配置
EXPANSION_CONFIG = {
    'max_rows_per_batch': 50,           # 每批最大添加行数
    'enable_cross_page': True,          # 启用跨页支持
    'preserve_formatting': True,        # 保持格式
    'auto_optimize_layout': True,       # 自动优化布局
    'memory_limit_mb': 512,            # 内存限制(MB)
}

# 日志配置
LOGGING_CONFIG = {
    'log_expansion_details': True,      # 记录扩展详情
    'log_performance_metrics': True,    # 记录性能指标
    'log_format_operations': False,     # 记录格式操作
    'log_level': 'INFO',               # 日志级别
}

@dataclass
class TableExpansionRequirement:
    """表格扩展需求模型"""
    table_index: int
    current_rows: int
    required_rows: int
    rows_to_add: int
    insert_position: int
    reference_row_index: int

@dataclass
class InspectionRowRequirement:
    """检件行需求模型"""
    inspection_number: str
    sheet_count: int
    required_rows: int
    start_row_index: int
    film_numbers: List[str]

class DataRowCalculator:
    """数据行计算器 - 计算检件所需的行数"""

    @staticmethod
    def calculate_rows_for_inspection(sheet_count: int) -> int:
        """根据张数计算单个检件所需行数

        Args:
            sheet_count: 检件张数

        Returns:
            所需行数
        """
        if sheet_count == 2 or sheet_count == 3:
            return sheet_count  # 2张需要2行，3张需要3行
        elif sheet_count >= 6:
            return sheet_count  # ≥6张时，每张一行
        else:
            return 1  # 其他情况默认1行

    @staticmethod
    def calculate_total_rows(inspection_numbers: List[str], sheet_counts: List[int]) -> Tuple[int, List[InspectionRowRequirement]]:
        """计算所有检件所需的总行数

        Args:
            inspection_numbers: 检件编号列表
            sheet_counts: 对应的张数列表

        Returns:
            (总行数, 检件行需求列表)
        """
        total_rows = 0
        requirements = []
        current_start_row = 0

        for inspection_num, sheet_count in zip(inspection_numbers, sheet_counts):
            # 确保sheet_count是整数
            sheet_count = int(float(sheet_count)) if sheet_count is not None else 1
            required_rows = DataRowCalculator.calculate_rows_for_inspection(sheet_count)

            # 生成片号列表
            film_numbers = []
            if sheet_count >= 6:
                # 为每张生成片号
                for j in range(sheet_count):
                    film_numbers.append(f"{inspection_num}-{j+1}")
            else:
                # 使用检件编号作为片号
                film_numbers = [str(inspection_num)] * int(required_rows)

            requirement = InspectionRowRequirement(
                inspection_number=inspection_num,
                sheet_count=sheet_count,
                required_rows=required_rows,
                start_row_index=current_start_row,
                film_numbers=film_numbers
            )

            requirements.append(requirement)
            total_rows += required_rows
            current_start_row += required_rows

            if LOGGING_CONFIG['log_expansion_details']:
                print(f"检件 {inspection_num}: 张数={sheet_count}, 需要行数={required_rows}")

        if LOGGING_CONFIG['log_expansion_details']:
            print(f"总计需要行数: {total_rows}")

        return total_rows, requirements

    @staticmethod
    def optimize_row_allocation(inspection_data: List[InspectionRowRequirement]) -> List[InspectionRowRequirement]:
        """优化行分配策略，确保数据连续性

        Args:
            inspection_data: 检件行需求列表

        Returns:
            优化后的检件行需求列表
        """
        # 按检件编号排序，确保连续性
        optimized_data = sorted(inspection_data, key=lambda x: x.inspection_number)

        # 重新计算起始行索引
        current_start_row = 0
        for requirement in optimized_data:
            requirement.start_row_index = current_start_row
            current_start_row += requirement.required_rows

        return optimized_data

def print_template_usage_summary(template_usage_summary):
    """输出模板使用总结"""
    if not template_usage_summary:
        print("\n未找到模板使用记录")
        return

    print("\n" + "="*100)
    print("模板使用总结 - 委托单编号与Word模板对应关系")
    print("="*100)

    # 按模板类型分组统计
    standard_template_count = 0
    continuation_template_count = 0

    # 表头
    print(f"{'序号':<4} {'委托单编号':<25} {'射线类型':<8} {'张数总和':<8} {'数据行数':<8} {'模板类型':<10} {'使用的模板文件'}")
    print("-"*100)

    # 输出每个委托单的模板使用情况
    for i, info in enumerate(template_usage_summary, 1):
        order_number = info['order_number']
        ray_type = info['ray_type']
        total_sheets = info.get('total_sheets', 0)
        data_rows = info.get('data_rows', 0)
        template_type = info['template_type']
        template_file = info['selected_template']

        # 统计
        if template_type == '标准模板':
            standard_template_count += 1
        else:
            continuation_template_count += 1

        # 格式化输出
        print(f"{i:<4} {order_number:<25} {ray_type:<8} {total_sheets:<8} {data_rows:<8} {template_type:<10} {template_file}")

    print("-"*100)

    # 统计汇总
    total_count = len(template_usage_summary)
    print(f"统计汇总:")
    print(f"  总计处理: {total_count} 个委托单编号")
    print(f"  使用标准模板: {standard_template_count} 个 (5_射线检测记录_续_新.docx)")
    print(f"  使用续表模板: {continuation_template_count} 个 (5_射线检测记录_续.docx)")

    # 模板选择规则说明
    print(f"\n模板选择规则:")
    print(f"  • 张数总和 ≤ 21：使用标准模板 (5_射线检测记录_续_新.docx)")
    print(f"  • 张数总和 > 21：使用续表模板 (5_射线检测记录_续.docx)")

    # 按模板类型分组显示
    if standard_template_count > 0:
        print(f"\n使用标准模板的委托单编号:")
        standard_orders = [info['order_number'] for info in template_usage_summary if info['template_type'] == '标准模板']
        for i, order in enumerate(standard_orders, 1):
            if i % 5 == 1:  # 每行显示5个
                print(f"  ", end="")
            print(f"{order:<25}", end="")
            if i % 5 == 0 or i == len(standard_orders):
                print()

    if continuation_template_count > 0:
        print(f"\n使用续表模板的委托单编号:")
        continuation_orders = [info['order_number'] for info in template_usage_summary if info['template_type'] == '续表模板']
        for i, order in enumerate(continuation_orders, 1):
            if i % 5 == 1:  # 每行显示5个
                print(f"  ", end="")
            print(f"{order:<25}", end="")
            if i % 5 == 0 or i == len(continuation_orders):
                print()

    print("="*100)

class TableCapacityAnalyzer:
    """表格容量分析器 - 分析表格容量并计算扩展需求"""

    def __init__(self, doc: Document, column_indices: Dict[str, int], header_row_index: int):
        self.doc = doc
        self.column_indices = column_indices
        self.header_row_index = header_row_index

    def analyze_available_rows(self, table) -> Tuple[int, List[int]]:
        """分析表格中可用于数据填充的行数

        Args:
            table: Word表格对象

        Returns:
            (可用行数, 可用行索引列表)
        """
        available_rows = []

        # 从表头行之后开始查找可用行
        for i in range(self.header_row_index + 1, len(table.rows)):
            row = table.rows[i]

            # 检查行是否为空或包含特殊标记
            is_empty_or_usable = True
            for cell in row.cells:
                cell_text = cell.text.strip()
                # 如果单元格包含特殊标记，则不可用
                if cell_text and any(marker in cell_text for marker in ['以下空白', '合计', '总计', '备注']):
                    is_empty_or_usable = False
                    break

            if is_empty_or_usable:
                available_rows.append(i)

        if LOGGING_CONFIG['log_expansion_details']:
            print(f"表格分析: 表头行={self.header_row_index}, 可用行数={len(available_rows)}")
            print(f"可用行索引: {available_rows}")

        return len(available_rows), available_rows

    def calculate_required_rows(self, inspection_data: List[str], sheet_counts: List[int]) -> Tuple[int, List[InspectionRowRequirement]]:
        """计算填充所有数据所需的总行数

        Args:
            inspection_data: 检件编号列表
            sheet_counts: 对应的张数列表

        Returns:
            (所需总行数, 检件行需求列表)
        """
        return DataRowCalculator.calculate_total_rows(inspection_data, sheet_counts)

    def get_expansion_requirements(self, table, required_rows: int) -> Optional[TableExpansionRequirement]:
        """获取表格扩展需求

        Args:
            table: Word表格对象
            required_rows: 所需总行数

        Returns:
            表格扩展需求对象，如果不需要扩展则返回None
        """
        available_count, available_rows = self.analyze_available_rows(table)

        if required_rows <= available_count:
            if LOGGING_CONFIG['log_expansion_details']:
                print(f"表格容量充足: 需要{required_rows}行, 可用{available_count}行")
            return None

        rows_to_add = required_rows - available_count

        # 找到最佳插入位置（通常在最后一个可用行之后）
        insert_position = available_rows[-1] + 1 if available_rows else self.header_row_index + 1
        reference_row_index = available_rows[-1] if available_rows else self.header_row_index

        requirement = TableExpansionRequirement(
            table_index=0,  # 假设只有一个主表格
            current_rows=available_count,
            required_rows=required_rows,
            rows_to_add=rows_to_add,
            insert_position=insert_position,
            reference_row_index=reference_row_index
        )

        if LOGGING_CONFIG['log_expansion_details']:
            print(f"表格扩展需求: 当前可用{available_count}行, 需要{required_rows}行, 需添加{rows_to_add}行")
            print(f"插入位置: {insert_position}, 参考行: {reference_row_index}")

        return requirement

class FormatPreserver:
    """格式保持器 - 保持表格格式一致性"""

    def __init__(self, table):
        self.table = table

    def preserve_cell_format(self, source_cell, target_cell):
        """保持单元格格式

        Args:
            source_cell: 源单元格
            target_cell: 目标单元格
        """
        try:
            # 复制单元格属性
            if source_cell._element.tcPr is not None:
                target_cell._element.tcPr = source_cell._element.tcPr

            # 复制段落格式
            if len(source_cell.paragraphs) > 0 and len(target_cell.paragraphs) > 0:
                self.copy_paragraph_format(source_cell.paragraphs[0], target_cell.paragraphs[0])

        except Exception as e:
            if LOGGING_CONFIG['log_format_operations']:
                print(f"复制单元格格式时出错: {e}")

    def preserve_row_format(self, source_row, target_row):
        """保持行格式

        Args:
            source_row: 源行
            target_row: 目标行
        """
        try:
            # 复制行高
            if source_row.height:
                target_row.height = source_row.height

            # 复制每个单元格的格式
            min_cells = min(len(source_row.cells), len(target_row.cells))
            for i in range(min_cells):
                self.preserve_cell_format(source_row.cells[i], target_row.cells[i])

        except Exception as e:
            if LOGGING_CONFIG['log_format_operations']:
                print(f"复制行格式时出错: {e}")

    def copy_paragraph_format(self, source_para, target_para):
        """复制段落格式

        Args:
            source_para: 源段落
            target_para: 目标段落
        """
        try:
            # 复制段落属性
            if source_para._element.pPr is not None:
                target_para._element.pPr = source_para._element.pPr

        except Exception as e:
            if LOGGING_CONFIG['log_format_operations']:
                print(f"复制段落格式时出错: {e}")

    def validate_format_consistency(self) -> bool:
        """验证格式一致性

        Returns:
            格式是否一致
        """
        try:
            # 简单的格式一致性检查
            if len(self.table.rows) < 2:
                return True

            # 检查列数是否一致
            first_row_cells = len(self.table.rows[0].cells)
            for row in self.table.rows[1:]:
                if len(row.cells) != first_row_cells:
                    return False

            return True

        except Exception as e:
            if LOGGING_CONFIG['log_format_operations']:
                print(f"验证格式一致性时出错: {e}")
            return False

class DynamicRowExpander:
    """动态行扩展器 - 执行表格行添加操作"""

    def __init__(self, table, reference_row_index: int):
        self.table = table
        self.reference_row_index = reference_row_index
        self.format_preserver = FormatPreserver(table)

    def add_rows(self, count: int, insert_position: Optional[int] = None) -> List[int]:
        """添加指定数量的行

        Args:
            count: 要添加的行数
            insert_position: 插入位置，如果为None则在末尾添加

        Returns:
            新添加行的索引列表
        """
        if count <= 0:
            return []

        new_row_indices = []

        try:
            # 获取参考行用于格式复制
            reference_row = self.table.rows[self.reference_row_index] if self.reference_row_index < len(self.table.rows) else None

            # 批量添加行以提高性能
            batch_size = min(count, EXPANSION_CONFIG['max_rows_per_batch'])

            for batch_start in range(0, count, batch_size):
                batch_count = min(batch_size, count - batch_start)

                for i in range(batch_count):
                    # 添加新行
                    new_row = self.table.add_row()
                    new_row_index = len(self.table.rows) - 1
                    new_row_indices.append(new_row_index)

                    # 复制格式
                    if reference_row and EXPANSION_CONFIG['preserve_formatting']:
                        self.format_preserver.preserve_row_format(reference_row, new_row)

                # 内存管理
                if EXPANSION_CONFIG['memory_limit_mb'] and batch_start % 100 == 0:
                    gc.collect()

            if LOGGING_CONFIG['log_expansion_details']:
                print(f"成功添加 {count} 行，新行索引: {new_row_indices}")

            return new_row_indices

        except Exception as e:
            print(f"添加表格行时出错: {e}")
            return []

    def copy_row_format(self, source_row, target_row):
        """复制行格式

        Args:
            source_row: 源行
            target_row: 目标行
        """
        self.format_preserver.preserve_row_format(source_row, target_row)

    def ensure_cross_page_compatibility(self):
        """确保跨页兼容性"""
        if not EXPANSION_CONFIG['enable_cross_page']:
            return

        try:
            # 设置表格属性以支持跨页
            tbl = self.table._element
            tblPr = tbl.tblPr

            # 允许表格跨页
            if tblPr.tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'autofit')
                tblPr.append(tblLayout)

            # 设置表头重复
            if hasattr(self.table, 'rows') and len(self.table.rows) > 0:
                header_row = self.table.rows[0]
                trPr = header_row._element.trPr
                if trPr is None:
                    trPr = OxmlElement('w:trPr')
                    header_row._element.insert(0, trPr)

                # 添加表头重复属性
                tblHeader = OxmlElement('w:tblHeader')
                trPr.append(tblHeader)

            if LOGGING_CONFIG['log_expansion_details']:
                print("已设置表格跨页兼容性")

        except Exception as e:
            print(f"设置跨页兼容性时出错: {e}")

def find_column_with_keyword(df, keyword):
    """查找包含指定关键字的列"""
    # 首先尝试精确匹配
    exact_matches = [col for col in df.columns if col.strip() == keyword.strip()]
    if exact_matches:
        return exact_matches[0]

    # 如果没有精确匹配，再尝试部分匹配
    partial_matches = [col for col in df.columns if keyword.lower() in col.lower()]
    return partial_matches[0] if partial_matches else None

def set_font_style(paragraph, font_name="楷体", font_size=10.5):
    """设置段落字体为楷体五号（10.5磅）

    Args:
        paragraph: Word段落对象
        font_name: 字体名称，默认为"楷体"
        font_size: 字体大小，默认为10.5磅（五号字）
    """
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def normalize_text(text):
    """标准化文本以便更好的匹配"""
    if not text:
        return ""
    # 移除空格、换行符等空白字符
    normalized = re.sub(r'\s+', '', text.strip())
    # 转换为小写
    normalized = normalized.lower()
    return normalized

def get_output_filename(word_template_path, order_number, ray_type):
    """根据Word模板路径、委托单编号和射线类型生成输出文件名
    
    Args:
        word_template_path: Word模板文档路径
        order_number: 委托单编号
        ray_type: 射线类型（"X射线"或"γ射线"）
    
    Returns:
        str: 输出文件名
    """
    # 获取模板文件名（不含路径和扩展名）
    template_name = os.path.splitext(os.path.basename(word_template_path))[0]
    # 射线类型标识
    ray_mark = "γ" if ray_type == "γ射线" else "X"
    # 生成输出文件名
    return f"{template_name}_{order_number}_{ray_mark}_续表_生成结果.docx"

# def setup_logging():
#     """设置日志配置"""
#     log_filename = f"radio_test_renewal_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
#     log_file = os.path.join("生成器/输出报告/log", log_filename)

#     # 创建日志目录
#     log_dir = os.path.dirname(log_file)
#     if not os.path.exists(log_dir):
#         os.makedirs(log_dir)

#     # 配置日志
#     logging.basicConfig(
#         level=logging.INFO,
#         format='%(asctime)s - %(levelname)s - %(message)s',
#         handlers=[
#             logging.FileHandler(log_file, encoding='utf-8'),
#             logging.StreamHandler(sys.stdout)
#         ]
#     )

#     logging.info(f"日志文件已创建: {log_file}")
#     return log_file

def calculate_total_inspection_count_by_order(df, column_mapping):
    """计算每个委托单编号的检件编号总个数"""
    # 按委托单编号分组，统计每个委托单的检件编号总个数
    order_inspection_counts = {}

    logging.info("开始计算每个委托单编号的检件编号总个数...")

    for order_number in df[column_mapping['委托单编号']].unique():
        order_df = df[df[column_mapping['委托单编号']] == order_number]
        unique_inspections = order_df[column_mapping['检件编号']].dropna().unique()
        inspection_count = len(unique_inspections)

        order_inspection_counts[order_number] = {
            'count': inspection_count,
            'inspections': unique_inspections.tolist()
        }

        logging.info(f"委托单编号 {order_number}: 检件编号总个数 = {inspection_count}")
        logging.info(f"  检件编号列表: {unique_inspections.tolist()}")

    return order_inspection_counts

def setup_logging(output_dir):
    """设置日志配置"""
    log_file = os.path.join(output_dir, f"radio_test_renewal_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")

    # 创建日志目录
    os.makedirs(os.path.dirname(log_file), exist_ok=True)

    # 配置日志
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8'),
            logging.StreamHandler(sys.stdout)
        ]
    )

    logging.info(f"日志文件: {log_file}")
    return log_file

def process_excel_to_word(excel_path, word_template_path, output_path=None, project_name=None, client_name=None, instruction_number=None):
    """将Excel数据填入Word文档

    Args:
        excel_path: Excel表格路径
        word_template_path: Word模板文档路径
        output_path: 输出目录路径（如果为None，将自动生成）
        project_name: 工程名称，用于替换Word文档中的"工程名称值"
        client_name: 委托单位，用于替换Word文档中的"委托单位值"
        instruction_number: 操作指导书编号，用于替换Word文档中的"操作指导书编号值"

    Returns:
        bool: 处理是否成功
    """
    # 创建输出目录
    if output_path is None:
        output_dir = os.path.join("生成器", "输出报告", "5_射线检测记录续")
    else:
        output_dir = output_path

    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            print(f"创建输出目录: {output_dir}")
        except Exception as e:
            print(f"错误: 无法创建输出目录: {e}")
            return False

    # # 设置日志
    # log_file = setup_logging()
    # logging.info("="*80)
    # logging.info("开始处理射线检测记录续表生成")
    # logging.info("="*80)

    # 检查文件是否存在
    if not os.path.exists(excel_path):
        logging.error(f"Excel文件不存在: {excel_path}")
        return False

    if not os.path.exists(word_template_path):
        logging.error(f"Word模板文件不存在: {word_template_path}")
        return False
    
    # 读取Excel数据
    logging.info(f"正在读取Excel文件: {excel_path}")
    try:
        df = pd.read_excel(excel_path)
        logging.info(f"成功读取Excel文件，共有{len(df)}行数据")
    except Exception as e:
        logging.error(f"无法读取Excel文件: {e}")
        return False

    # 打印所有列名，帮助调试
    logging.info(f"Excel表格列名: {list(df.columns)}")
    
    # 定义需要查找的列关键字
    column_keywords = {
        '完成日期': '完成日期',
        '委托单编号': '委托单编号',
        '检件编号': '检件编号',
        '焊口编号': '焊口编号',
        '焊工号': '焊工号',
        '规格': '规格',
        'γ射线': 'γ射线',
        '张数': '张数',  # 确保精确匹配"张数"列
        '合格级别': '合格级别',
        '检测比例': '检测比列'  # Excel中的列名是"检测比列"
    }
    
    # 查找每个关键字对应的实际列名
    column_mapping = {}
    missing_columns = []
    
    for key, keyword in column_keywords.items():
        col_name = find_column_with_keyword(df, keyword)
        if col_name:
            column_mapping[key] = col_name
            print(f"找到列: '{key}' -> '{col_name}'")
        else:
            missing_columns.append(key)
            
    if missing_columns:
        print(f"警告: 未找到以下列: {', '.join(missing_columns)}")
        # 尝试使用列位置
        possible_columns = {
            '完成日期': 'B',
            '委托单编号': 'C',
            '检件编号': 'D',
            '焊口编号': 'E',
            '焊工号': 'F',
            '规格': 'G',
            'γ射线': 'P',
            '张数': 'M',  # 明确指定张数列为M列
            '合格级别': 'I',
            '检测比例': 'J'
        }
        
        for key in missing_columns:
            col_letter = possible_columns.get(key)
            if col_letter:
                col_idx = ord(col_letter) - ord('A')
                if col_idx < len(df.columns):
                    column_mapping[key] = df.columns[col_idx]
                    print(f"使用列位置找到: '{key}' -> '{df.columns[col_idx]}'")
                else:
                    print(f"警告: 列位置 {col_letter} 超出范围，无法找到 '{key}'")
    
    # 如果缺少关键列，则无法继续处理
    if '委托单编号' not in column_mapping:
        print("错误: 无法找到委托单编号列")
        return False
    
    if 'γ射线' not in column_mapping:
        print("警告: 无法找到γ射线列，将默认所有记录为X射线")
        # 增加一个全为空的列作为γ射线列
        df['γ射线'] = None
        column_mapping['γ射线'] = 'γ射线'
    
    # 根据委托单编号和射线类型分组数据
    groups = []
    order_numbers = df[column_mapping['委托单编号']].dropna().unique()
    
    for order_number in order_numbers:
        # 获取该委托单编号的所有数据
        order_df = df[df[column_mapping['委托单编号']] == order_number]
        
        # 获取该委托单编号下的所有射线类型
        ray_types = order_df[column_mapping['γ射线']].dropna().unique()
        
        # 如果没有明确的γ射线值，则视为X射线
        if len(ray_types) == 0:
            groups.append({
                'order_number': order_number,
                'ray_type': 'X射线',  # X射线表示为处理逻辑，但射源种类值会设置为空
                'data': order_df
            })
            print(f"委托单编号 {order_number} 没有明确的射线类型，处理为X射线")
        else:
            # 有γ射线值的处理为γ射线
            for ray_type in ray_types:
                ray_df = order_df[order_df[column_mapping['γ射线']] == ray_type]
                groups.append({
                    'order_number': order_number,
                    'ray_type': 'γ射线',
                    'data': ray_df
                })
                print(f"委托单编号 {order_number} 的射线类型 γ射线 有 {len(ray_df)} 条记录")
            
            # 没有γ射线值的处理为X射线
            x_ray_df = order_df[order_df[column_mapping['γ射线']].isna()]
            if len(x_ray_df) > 0:
                groups.append({
                    'order_number': order_number,
                    'ray_type': 'X射线',  # X射线表示为处理逻辑，但射源种类值会设置为空
                    'data': x_ray_df
                })
                print(f"委托单编号 {order_number} 的射线类型 X射线 有 {len(x_ray_df)} 条记录")
    
    logging.info(f"共有 {len(groups)} 个组合需要生成报告")

    # 在处理之前，先统计每个分组的张数总和
    logging.info("="*80)
    logging.info("分组张数统计分析")
    logging.info("="*80)

    group_sheet_counts = {}
    for group in groups:
        order_number = group['order_number']
        ray_type = group['ray_type']
        group_df = group['data']

        # 统计该分组中张数的总和
        if '张数' in column_mapping and column_mapping['张数'] in group_df.columns:
            # 将张数列转换为数值，无效值设为0
            sheet_numbers = pd.to_numeric(group_df[column_mapping['张数']], errors='coerce').fillna(0)
            total_sheets = int(sheet_numbers.sum())
        else:
            # 如果没有张数列，默认每行1张
            total_sheets = len(group_df)
            logging.warning(f"未找到张数列，默认每行1张，总计: {total_sheets}")

        group_key = f"{order_number}_{ray_type}"
        group_sheet_counts[group_key] = {
            'order_number': order_number,
            'ray_type': ray_type,
            'total_sheets': total_sheets,
            'data_rows': len(group_df)
        }

        logging.info(f"分组: {order_number} + {ray_type}")
        logging.info(f"  数据行数: {len(group_df)}")
        logging.info(f"  张数总和: {total_sheets}")
        logging.info(f"  预期模板: {'续表模板' if total_sheets > 21 else '标准模板'}")
        logging.info("-" * 60)

    logging.info("="*80)
    logging.info("开始处理各个分组")
    logging.info("="*80)

    # 处理每个分组
    success_count = 0
    error_count = 0

    # 记录每个委托单编号使用的模板信息
    template_usage_summary = []
    
    for group in groups:
        order_number = group['order_number']
        ray_type = group['ray_type']
        group_df = group['data']
        
        print(f"\n{'='*50}")
        print(f"处理委托单编号: {order_number}, 射线类型: {ray_type}")
        print(f"{'='*50}")
        
        try:
            # 智能模板选择：使用预先统计的张数总和
            group_key = f"{order_number}_{ray_type}"
            group_stats = group_sheet_counts.get(group_key, {})

            total_sheets = group_stats.get('total_sheets', 0)
            data_rows = group_stats.get('data_rows', 0)

            logging.info(f"处理分组: {order_number} + {ray_type}")
            logging.info(f"张数统计: 数据行数={data_rows}, 张数总和={total_sheets}")

            # 根据张数总和选择模板
            selected_template_path = word_template_path  # 默认使用传入的模板

            if total_sheets > 21:
                # 当张数总和大于21时，使用续表模板
                alternative_template = "生成器/word/5_射线检测记录_续.docx"
                if os.path.exists(alternative_template):
                    selected_template_path = alternative_template
                    logging.info(f"✓ 张数总和({total_sheets})大于21，自动选用续表模板: {alternative_template}")
                    print(f"✓ 张数总和({total_sheets})大于21，自动选用续表模板: {alternative_template}")
                else:
                    logging.warning(f"⚠ 张数总和({total_sheets})大于21，但续表模板不存在: {alternative_template}")
                    logging.info(f"  继续使用原模板: {word_template_path}")
                    print(f"⚠ 张数总和({total_sheets})大于21，但续表模板不存在: {alternative_template}")
                    print(f"  继续使用原模板: {word_template_path}")
            else:
                logging.info(f"○ 张数总和({total_sheets})≤21，使用标准模板: {word_template_path}")
                print(f"○ 张数总和({total_sheets})≤21，使用标准模板: {word_template_path}")

            # 记录模板使用信息
            template_info = {
                'order_number': order_number,
                'ray_type': ray_type,
                'total_sheets': total_sheets,
                'data_rows': data_rows,
                'selected_template': os.path.basename(selected_template_path),
                'template_path': selected_template_path,
                'template_type': '续表模板' if total_sheets > 21 else '标准模板'
            }
            template_usage_summary.append(template_info)

            # 为该分组生成输出文件名（使用选定的模板）
            output_filename = get_output_filename(selected_template_path, order_number, ray_type)
            report_output_path = os.path.join(output_dir, output_filename)
            print(f"输出文件路径: {report_output_path}")

            # 打开Word文档
            print(f"正在处理Word文档: {selected_template_path}")

            try:
                # 每次处理新的组合时，重新从模板创建文档对象
                # 这确保了每个组合都会生成一个独立的文档
                doc = Document(selected_template_path)
                print(f"成功从模板创建新文档")
            except Exception as e:
                print(f"无法打开Word文档: {e}")
                error_count += 1
                continue
            
            # 1) 获取该组数据中最晚的完成日期
            date_col = column_mapping.get('完成日期')
            if date_col:
                # 确保日期列是日期类型
                group_df[date_col] = pd.to_datetime(group_df[date_col], errors='coerce')
                latest_date = group_df[date_col].max()
                
                if pd.isna(latest_date):
                    print(f"警告: 委托单编号 {order_number} 没有有效的完成日期")
                    year, month, day = datetime.now().year, datetime.now().month, datetime.now().day
                else:
                    # 将日期转换为年、月、日
                    year = latest_date.year
                    month = latest_date.month
                    day = latest_date.day
                    print(f"找到最晚完成日期: {year}年{month}月{day}日")
            else:
                print("警告: 未找到完成日期列")
                year, month, day = datetime.now().year, datetime.now().month, datetime.now().day
            
            # 获取相关数据
            inspection_numbers = []
            weld_numbers = []
            welder_numbers = []
            specifications = []
            sheet_counts = []
            
            # 直接从DataFrame中获取数据
            for i in range(len(group_df)):
                if i < len(group_df):
                    # 获取当前行的各列值
                    inspection_number = group_df.iloc[i][column_mapping['检件编号']] if '检件编号' in column_mapping else ""
                    weld_number = group_df.iloc[i][column_mapping['焊口编号']] if '焊口编号' in column_mapping else ""
                    welder_number = group_df.iloc[i][column_mapping['焊工号']] if '焊工号' in column_mapping else ""
                    specification = group_df.iloc[i][column_mapping['规格']] if '规格' in column_mapping else ""
                    
                    # 添加到对应的列表
                    inspection_numbers.append(str(inspection_number))
                    weld_numbers.append(str(weld_number))
                    welder_numbers.append(str(welder_number))
                    specifications.append(str(specification))
                    
                    # 处理张数
                    if '张数' in column_mapping:
                        sheet_count_raw = group_df.iloc[i][column_mapping['张数']]
                        print(f"行 {i+1} 检件编号 {inspection_number} 原始张数值: '{sheet_count_raw}' (类型: {type(sheet_count_raw).__name__})")
                        
                        # 确保张数是数值类型
                        try:
                            if pd.isna(sheet_count_raw):
                                print(f"警告: 行 {i+1} 检件编号 {inspection_number} 的张数值为NaN，默认为1")
                                sheet_count = 1
                            else:
                                # 尝试从"180*80/3张"这种格式中提取数字
                                if isinstance(sheet_count_raw, str) and '张' in sheet_count_raw:
                                    # 使用正则表达式提取斜杠后、张字前的数字
                                    match = re.search(r'/(\d+)张', sheet_count_raw)
                                    if match:
                                        sheet_count = int(match.group(1))
                                        print(f"从字符串 '{sheet_count_raw}' 中提取到张数: {sheet_count}")
                                    else:
                                        # 如果没有找到匹配的模式，默认为1
                                        print(f"无法从 '{sheet_count_raw}' 中提取张数，默认为1")
                                        sheet_count = 1
                                else:
                                    # 尝试直接转换为整数
                                    sheet_count = int(float(sheet_count_raw))
                                    print(f"行 {i+1} 检件编号 {inspection_number} 的张数值转换为: {sheet_count}")
                        except (ValueError, TypeError) as e:
                            print(f"警告: 行 {i+1} 检件编号 {inspection_number} 的张数值 '{sheet_count_raw}' 转换失败: {e}，默认为1")
                            sheet_count = 1
                        sheet_counts.append(sheet_count)
                    else:
                        sheet_counts.append(1)  # 默认为1
            
            if '张数' in column_mapping:
                print(f"\n张数列名: '{column_mapping['张数']}'")
                print(f"张数列所有值: {group_df[column_mapping['张数']].tolist()}")
                print(f"最终获取到的张数数据: {sheet_counts}")
            else:
                print("未找到张数列，默认所有记录的张数为1")

            # 获取分组的第一个值用于复选框处理
            grade_level = ""
            if '合格级别' in column_mapping:
                grade_levels = group_df[column_mapping['合格级别']].dropna().tolist()
                if grade_levels:
                    grade_level = str(grade_levels[0])
                    print(f"找到合格级别: {grade_level}")

            inspection_ratio = ""
            if '检测比例' in column_mapping:
                ratios = group_df[column_mapping['检测比例']].dropna().tolist()
                if ratios:
                    # 转换为百分数格式
                    try:
                        ratio_value = float(ratios[0])
                        if ratio_value <= 1:  # 如果是小数形式（如0.5），转换为百分数
                            inspection_ratio = f"{ratio_value*100:.0f}%"
                        else:  # 如果已经是百分数形式（如50），直接添加%
                            inspection_ratio = f"{ratio_value:.0f}%"
                    except (ValueError, TypeError):
                        inspection_ratio = str(ratios[0])
                        if not inspection_ratio.endswith('%'):
                            inspection_ratio += '%'
                    print(f"找到检测比例: {inspection_ratio}")
            
            # 替换文档中的值
            print("\n==== 开始替换文档中的值 ====")
            
            # 将EPKJ拼接委托单编号代替委托单编号值
            committee_order = f"EPKJ-{order_number}"
            replaced = False
            
            # 遍历所有段落，替换关键词
            for paragraph in doc.paragraphs:
                if "委托单编号值" in paragraph.text:
                    paragraph.text = paragraph.text.replace("委托单编号值", committee_order)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'委托单编号值'替换为'{committee_order}'")
                    replaced = True
                # 替换工程名称
                if "工程名称值" in paragraph.text and project_name:
                    paragraph.text = paragraph.text.replace("工程名称值", project_name)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'工程名称值'替换为'{project_name}'")
                # 替换委托单位
                if "委托单位值" in paragraph.text and client_name:
                    paragraph.text = paragraph.text.replace("委托单位值", client_name)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'委托单位值'替换为'{client_name}'")
                # 替换操作指导书编号
                if "操作指导书编号值" in paragraph.text and instruction_number:
                    paragraph.text = paragraph.text.replace("操作指导书编号值", instruction_number)
                    set_font_style(paragraph)  # 设置楷体五号字体
                    print(f"已将段落中的'操作指导书编号值'替换为'{instruction_number}'")
            
            # 遍历表格中的单元格，替换关键词
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        for paragraph in cell.paragraphs:
                            if "委托单编号值" in paragraph.text:
                                paragraph.text = paragraph.text.replace("委托单编号值", committee_order)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'委托单编号值'替换为'{committee_order}'")
                                replaced = True
                            # 替换工程名称
                            if "工程名称值" in paragraph.text and project_name:
                                paragraph.text = paragraph.text.replace("工程名称值", project_name)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'工程名称值'替换为'{project_name}'")
                            # 替换委托单位
                            if "委托单位值" in paragraph.text and client_name:
                                paragraph.text = paragraph.text.replace("委托单位值", client_name)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'委托单位值'替换为'{client_name}'")
                            # 替换操作指导书编号
                            if "操作指导书编号值" in paragraph.text and instruction_number:
                                paragraph.text = paragraph.text.replace("操作指导书编号值", instruction_number)
                                set_font_style(paragraph)  # 设置楷体五号字体
                                print(f"已将表格单元格中的'操作指导书编号值'替换为'{instruction_number}'")
            
            if not replaced:
                print("警告: 未找到需要替换的关键词，可能需要检查Word模板中的占位符命名。")
            
            # 填写日期（评片人、审核人）
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        # 处理日期
                        date_patterns = ["评片人", "审核人"]
                        for pattern in date_patterns:
                            if pattern in cell.text:
                                print(f"找到{pattern}单元格: 表格行{i+1}, 列{j+1}")
                                
                                # 检查单元格中的所有段落
                                date_found = False
                                for paragraph in cell.paragraphs:
                                    if "年" in paragraph.text and "月" in paragraph.text and "日" in paragraph.text:
                                        print(f"找到日期段落: {paragraph.text}")

                                        # 清空段落并重新构建，设置正确的字体
                                        original_text = paragraph.text
                                        paragraph.clear()

                                        # 分割文本并重新构建，为数字设置楷体五号字体
                                        parts = re.split(r'(\d*年|\d*月|\d*日)', original_text)
                                        for part in parts:
                                            if part:
                                                if '年' in part:
                                                    # 添加年份数字（楷体五号）
                                                    run_year = paragraph.add_run(str(year))
                                                    run_year.font.name = "楷体"
                                                    run_year.font.size = Pt(10.5)
                                                    run_year._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                                    # 添加"年"字（保持原格式）
                                                    paragraph.add_run("年")
                                                elif '月' in part:
                                                    # 添加月份数字（楷体五号）
                                                    run_month = paragraph.add_run(str(month))
                                                    run_month.font.name = "楷体"
                                                    run_month.font.size = Pt(10.5)
                                                    run_month._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                                    # 添加"月"字（保持原格式）
                                                    paragraph.add_run("月")
                                                elif '日' in part:
                                                    # 添加日期数字（楷体五号）
                                                    run_day = paragraph.add_run(str(day))
                                                    run_day.font.name = "楷体"
                                                    run_day.font.size = Pt(10.5)
                                                    run_day._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                                    # 添加"日"字（保持原格式）
                                                    paragraph.add_run("日")
                                                else:
                                                    # 非日期部分，保持原有格式
                                                    paragraph.add_run(part)

                                        date_found = True
                                        print(f"已更新{pattern}日期为 {year}年{month}月{day}日")
                                        break
                                
                                # 如果没有找到日期段落，尝试创建新段落
                                if not date_found:
                                    print(f"未在{pattern}单元格中找到日期段落，尝试添加")
                                    # 添加新段落并设置格式
                                    p = cell.add_paragraph()

                                    # 添加年份数字（楷体五号）
                                    run_year = p.add_run(str(year))
                                    run_year.font.name = "楷体"
                                    run_year.font.size = Pt(10.5)
                                    run_year._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")

                                    # 添加"年"字（保持原格式）
                                    p.add_run("年")

                                    # 添加月份数字（楷体五号）
                                    run_month = p.add_run(str(month))
                                    run_month.font.name = "楷体"
                                    run_month.font.size = Pt(10.5)
                                    run_month._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")

                                    # 添加"月"字（保持原格式）
                                    p.add_run("月")

                                    # 添加日期数字（楷体五号）
                                    run_day = p.add_run(str(day))
                                    run_day.font.name = "楷体"
                                    run_day.font.size = Pt(10.5)
                                    run_day._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")

                                    # 添加"日"字（保持原格式）
                                    p.add_run("日")

                                    print(f"已添加{pattern}日期: {year}年{month}月{day}日")
            
            # 查找表头行，确定各列的位置
            for table in doc.tables:
                column_indices = {}
                header_row_index = -1
                
                # 查找包含"检件编号"、"焊缝编号"、"焊工号"等的行
                for i, row in enumerate(table.rows):
                    header_found = False
                    for j, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        
                        # 打印表格单元格内容，帮助调试
                        print(f"表格单元格[{i},{j}]内容: '{cell_text}'")
                        
                        if "检件编号" in cell_text:
                            column_indices["检件编号"] = j
                            header_row_index = i
                            header_found = True
                            print(f"找到检件编号列: 行 {i+1}, 列 {j+1}")
                        elif "焊缝编号" in cell_text or "焊口编号" in cell_text:
                            column_indices["焊缝编号"] = j
                            header_found = True
                            print(f"找到焊缝编号列: 行 {i+1}, 列 {j+1}")
                        elif "焊工号" in cell_text:
                            column_indices["焊工号"] = j
                            header_found = True
                            print(f"找到焊工号列: 行 {i+1}, 列 {j+1}")
                        elif "规格" in cell_text:
                            column_indices["规格"] = j
                            header_found = True
                            print(f"找到规格列: 行 {i+1}, 列 {j+1}")
                        elif "片号" in cell_text:
                            column_indices["片号"] = j
                            header_found = True
                            print(f"找到片号列: 行 {i+1}, 列 {j+1}")
                    
                    if header_found and header_row_index >= 0:
                        print(f"找到表头行: 第{header_row_index+1}行")
                        print(f"列索引: {column_indices}")
                        break
                
                # 如果没有找到某些列，尝试通过位置确定
                if "片号" not in column_indices and header_row_index >= 0:
                    # 片号通常在焊缝编号和焊工号之间，尝试通过位置确定
                    if "焊缝编号" in column_indices and "焊工号" in column_indices:
                        expected_pos = min(column_indices["焊缝编号"] + 1, column_indices["焊工号"])
                        column_indices["片号"] = expected_pos
                        print(f"通过位置推断片号列: 列 {expected_pos+1}")
                    elif len(table.rows[header_row_index].cells) > 2:
                        # 如果没有找到焊缝编号和焊工号，但表格有足够的列，假设片号在第3列
                        column_indices["片号"] = 2
                        print(f"默认片号列位置: 列 3")
                
                if "焊缝编号" not in column_indices and header_row_index >= 0:
                    # 如果没有找到焊缝编号列，但找到了检件编号列，假设焊缝编号在检件编号后一列
                    if "检件编号" in column_indices and len(table.rows[header_row_index].cells) > 1:
                        column_indices["焊缝编号"] = column_indices["检件编号"] + 1
                        print(f"通过位置推断焊缝编号列: 列 {column_indices['焊缝编号']+1}")
                
                print(f"最终确定的列索引: {column_indices}")
                
                # 如果找到表头行，处理数据填充
                if header_row_index >= 0 and column_indices:
                    print("==== 开始动态表格扩展分析 ====")

                    # 创建表格容量分析器
                    capacity_analyzer = TableCapacityAnalyzer(doc, column_indices, header_row_index)

                    # 分析当前表格容量
                    available_count, data_rows = capacity_analyzer.analyze_available_rows(table)
                    print(f"表格容量分析: 可用行数={available_count}, 行索引={data_rows}")

                    # 计算数据行需求
                    inspection_numbers = group_df[column_mapping['检件编号']].tolist()
                    sheet_counts_data = group_df[column_mapping['张数']].fillna(1).tolist() if '张数' in column_mapping else [1] * len(inspection_numbers)

                    # 使用数据行计算器计算总需求
                    required_rows, inspection_requirements = capacity_analyzer.calculate_required_rows(
                        inspection_numbers, sheet_counts_data
                    )

                    print(f"数据需求分析: 需要总行数={required_rows}")

                    # 检查是否需要扩展表格
                    expansion_requirement = capacity_analyzer.get_expansion_requirements(table, required_rows)

                    if expansion_requirement:
                        print("==== 执行动态表格扩展 ====")

                        # 创建动态行扩展器
                        row_expander = DynamicRowExpander(table, expansion_requirement.reference_row_index)

                        # 添加所需的行
                        new_row_indices = row_expander.add_rows(
                            expansion_requirement.rows_to_add,
                            expansion_requirement.insert_position
                        )

                        # 更新可用行列表
                        data_rows.extend(new_row_indices)

                        # 确保跨页兼容性
                        row_expander.ensure_cross_page_compatibility()

                        print(f"表格扩展完成: 添加了{len(new_row_indices)}行")
                    else:
                        print("表格容量充足，无需扩展")

                    print("==== 动态表格扩展完成 ====\n")

                    # 获取基础数据
                    data_count = len(group_df)
                    print(f"需要填充{data_count}行数据")

                    # 对于张数≥6的情况，我们需要确保有足够的行来填写所有片号
                    # 检查是否需要添加额外的行来显示完整的片号序列
                    extra_rows_needed = 0
                    extra_rows_for_inspection = {}  # 记录每个检件编号需要的额外行数

                    for i in range(data_count):
                        if i < len(sheet_counts):
                            # 修改逻辑：当张数为2或3时，也需要额外行
                            if sheet_counts[i] >= 6 or sheet_counts[i] in [2, 3]:
                                # 每个符合条件的检件编号需要张数个行
                                extra_needed = sheet_counts[i] - 1  # 减1是因为已经计算了一行
                                extra_rows_needed += extra_needed
                                # 记录该检件编号需要的额外行数
                                inspection_number = inspection_numbers[i]
                                extra_rows_for_inspection[inspection_number] = extra_needed
                    
                    if extra_rows_needed > 0:
                        print(f"为了显示完整的片号序列，需要额外添加{extra_rows_needed}行")
                    
                    # 处理每一行数据
                    row_index = 0  # 用于跟踪当前处理的行索引
                    processed_rows = 0  # 已处理的行数
                    
                    # 首先计算每个检件编号需要的行数
                    inspection_rows_needed = {}
                    for i in range(len(inspection_numbers)):
                        current_inspection = inspection_numbers[i]
                        current_sheet_count = sheet_counts[i] if i < len(sheet_counts) else 1
                        
                        # 对于张数≥6的情况，需要生成多行
                        # 修改逻辑：当张数为2或3时，也生成对应数量的行
                        if current_sheet_count >= 6 or current_sheet_count in [2, 3]:
                            rows_to_generate = current_sheet_count
                        else:
                            rows_to_generate = 1
                        
                        if current_inspection in inspection_rows_needed:
                            inspection_rows_needed[current_inspection] += rows_to_generate
                        else:
                            inspection_rows_needed[current_inspection] = rows_to_generate
                    
                    print(f"每个检件编号需要的行数: {inspection_rows_needed}")
                    
                    # 处理每个检件编号
                    for i in range(data_count):
                        if i < len(inspection_numbers):
                            current_inspection = inspection_numbers[i]
                            current_weld = weld_numbers[i] if i < len(weld_numbers) else ""
                            current_welder = welder_numbers[i] if i < len(welder_numbers) else ""
                            current_spec = specifications[i] if i < len(specifications) else ""
                            current_sheet_count = sheet_counts[i] if i < len(sheet_counts) else 1
                            
                            # 对于张数≥6或张数为2、3的情况，需要生成多行
                            # 修改逻辑：当张数为2或3时，也生成对应数量的行
                            if current_sheet_count >= 6 or current_sheet_count in [2, 3]:
                                rows_to_generate = current_sheet_count
                            else:
                                rows_to_generate = 1
                            
                            # 检查是否需要为当前检件编号添加行
                            if row_index + rows_to_generate > len(data_rows):
                                # 计算需要添加的行数
                                rows_needed = row_index + rows_to_generate - len(data_rows)
                                print(f"为检件编号 {current_inspection} 添加 {rows_needed} 行")
                                
                                # 在当前位置插入新行
                                for _ in range(rows_needed):
                                    # 如果当前位置有效，在当前位置之后插入新行
                                    if row_index > 0 and row_index <= len(data_rows):
                                        # 获取当前行的索引
                                        current_row_idx = data_rows[row_index - 1] if row_index - 1 < len(data_rows) else len(table.rows) - 1
                                        
                                        # 在当前行之后插入新行
                                        new_row = table.add_row()
                                        
                                        # 将新行移动到当前行之后
                                        # 注意：python-docx不直接支持在特定位置插入行，所以我们需要记录行索引
                                        new_row_idx = len(table.rows) - 1
                                        data_rows.insert(row_index, new_row_idx)
                                    else:
                                        # 如果是在末尾添加行
                                        new_row = table.add_row()
                                        data_rows.append(len(table.rows) - 1)
                            
                            # 填写当前检件编号的所有行
                            for j in range(rows_to_generate):
                                if row_index < len(data_rows):
                                    row_idx = data_rows[row_index]
                                    row = table.rows[row_idx]
                                    
                                    # 1. 填写检件编号
                                    if "检件编号" in column_indices:
                                        col_idx = column_indices["检件编号"]
                                        if col_idx < len(row.cells):
                                            cell = row.cells[col_idx]
                                            if cell.paragraphs:
                                                cell.paragraphs[0].text = str(current_inspection)
                                                set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                                print(f"已更新第{row_idx+1}行检件编号: {current_inspection}")

                                    # 2. 填写焊缝编号
                                    if "焊缝编号" in column_indices:
                                        col_idx = column_indices["焊缝编号"]
                                        if col_idx < len(row.cells):
                                            cell = row.cells[col_idx]
                                            if cell.paragraphs:
                                                # 确保单元格内容被完全替换
                                                if cell.paragraphs[0].text:
                                                    cell.paragraphs[0].text = ""
                                                cell.paragraphs[0].text = str(current_weld)
                                                set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                                print(f"已更新第{row_idx+1}行焊缝编号: {current_weld}")

                                    # 3. 填写焊工号
                                    if "焊工号" in column_indices:
                                        col_idx = column_indices["焊工号"]
                                        if col_idx < len(row.cells):
                                            cell = row.cells[col_idx]
                                            if cell.paragraphs:
                                                cell.paragraphs[0].text = str(current_welder)
                                                set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                                print(f"已更新第{row_idx+1}行焊工号: {current_welder}")

                                    # 4. 填写规格
                                    if "规格" in column_indices:
                                        col_idx = column_indices["规格"]
                                        if col_idx < len(row.cells):
                                            cell = row.cells[col_idx]
                                            if cell.paragraphs:
                                                cell.paragraphs[0].text = str(current_spec)
                                                set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                                print(f"已更新第{row_idx+1}行规格: {current_spec}")
                                    
                                    # 5. 填写片号
                                    if "片号" in column_indices:
                                        col_idx = column_indices["片号"]
                                        if col_idx < len(row.cells):
                                            cell = row.cells[col_idx]
                                            
                                            # 确定当前行是该检件编号的第几个实例
                                            current_index_in_group = j
                                            
                                            # 根据张数规则确定片号
                                            film_number = ""
                                            
                                            if current_sheet_count in [1, 4, 5]:
                                                # 不填写片号
                                                film_number = ""
                                                print(f"张数为 {current_sheet_count}，片号保持为空")
                                            elif current_sheet_count == 2:
                                                # 依次填写1，2
                                                if current_index_in_group < 2:
                                                    film_number = str(current_index_in_group + 1)
                                                    print(f"张数为 2，当前是第 {current_index_in_group + 1} 个实例，片号为: {film_number}")
                                            elif current_sheet_count == 3:
                                                # 依次填写1，2，3
                                                if current_index_in_group < 3:
                                                    film_number = str(current_index_in_group + 1)
                                                    print(f"张数为 3，当前是第 {current_index_in_group + 1} 个实例，片号为: {film_number}")
                                            elif current_sheet_count >= 6:
                                                # 依次填写1-2，2-3，3-4，...，(N-1)-N，N-1
                                                if current_index_in_group < current_sheet_count:
                                                    if current_index_in_group < current_sheet_count - 1:
                                                        # 对于前N-1个实例，填写 i-(i+1)
                                                        film_number = f"{current_index_in_group + 1}-{current_index_in_group + 2}"
                                                    else:
                                                        # 对于第N个实例，填写 N-1
                                                        film_number = f"{current_sheet_count}-1"
                                                    print(f"张数为 {current_sheet_count}，当前是第 {current_index_in_group + 1} 个实例，片号为: {film_number}")
                                            
                                            # 打印当前单元格状态
                                            print(f"片号单元格当前内容: '{cell.text}'")
                                            
                                            # 确保单元格内容被完全替换
                                            try:
                                                # 先清空单元格的所有内容
                                                for p in cell.paragraphs:
                                                    p.clear()
                                                
                                                # 如果没有段落，添加一个新段落
                                                if len(cell.paragraphs) == 0:
                                                    p = cell.add_paragraph()
                                                
                                                # 设置片号文本
                                                run = cell.paragraphs[0].add_run(film_number)
                                                # 设置楷体五号字体
                                                run.font.name = "楷体"
                                                run.font.size = Pt(10.5)
                                                run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")

                                                if film_number:
                                                    print(f"已更新第{row_idx+1}行片号: '{film_number}'")
                                                else:
                                                    print(f"第{row_idx+1}行片号保留为空")
                                            except Exception as e:
                                                print(f"设置片号时出错: {e}")
                                                # 尝试另一种方式
                                                try:
                                                    if len(cell.paragraphs) > 0:
                                                        cell.paragraphs[0].text = film_number
                                                        set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                                    else:
                                                        cell.text = film_number
                                                    print(f"使用备用方法设置片号: '{film_number}'")
                                                except Exception as e2:
                                                    print(f"备用方法也失败: {e2}")
                                    
                                    # 6. 填写像质计灵敏度
                                    # 查找表格中的"像质计灵敏度"列
                                    sensitivity_col_idx = -1
                                    for j, cell in enumerate(table.rows[header_row_index].cells):
                                        if "像质计" in cell.text and "灵敏度" in cell.text:
                                            sensitivity_col_idx = j
                                            print(f"找到像质计灵敏度列: 行 {header_row_index+1}, 列 {j+1}")
                                            break
                                    
                                    if sensitivity_col_idx >= 0 and sensitivity_col_idx < len(row.cells):
                                        # 查找对应规格的像质计灵敏度值
                                        sensitivity_value = find_sensitivity_value(current_spec, ray_type)
                                        
                                        if sensitivity_value:
                                            # 填写像质计灵敏度值
                                            cell = row.cells[sensitivity_col_idx]
                                            
                                            try:
                                                # 先清空单元格的所有内容
                                                for p in cell.paragraphs:
                                                    p.clear()
                                                
                                                # 如果没有段落，添加一个新段落
                                                if len(cell.paragraphs) == 0:
                                                    p = cell.add_paragraph()
                                                
                                                # 设置像质计灵敏度文本
                                                run = cell.paragraphs[0].add_run(sensitivity_value)
                                                # 设置楷体五号字体
                                                run.font.name = "楷体"
                                                run.font.size = Pt(10.5)
                                                run._element.rPr.rFonts.set(qn('w:eastAsia'), "楷体")
                                                print(f"已更新第{row_idx+1}行像质计灵敏度: '{sensitivity_value}'")
                                            except Exception as e:
                                                print(f"设置像质计灵敏度时出错: {e}")
                                                # 尝试另一种方式
                                                try:
                                                    if len(cell.paragraphs) > 0:
                                                        cell.paragraphs[0].text = sensitivity_value
                                                        set_font_style(cell.paragraphs[0])  # 设置楷体五号字体
                                                    else:
                                                        cell.text = sensitivity_value
                                                    print(f"使用备用方法设置像质计灵敏度: '{sensitivity_value}'")
                                                except Exception as e2:
                                                    print(f"备用方法也失败: {e2}")
                                    
                                    row_index += 1
                                    processed_rows += 1
            
            print("==== 文档填充完成 ====\n")

            # 处理复选框匹配和标记
            print("==== 开始处理复选框匹配 ====")

            # 处理检测比例复选框匹配和标记
            if inspection_ratio:
                ratio_checkbox_success = process_detection_ratio_checkboxes(doc, inspection_ratio)
                if ratio_checkbox_success:
                    print("检测比例复选框处理完成")
                else:
                    print("检测比例复选框处理失败或未找到匹配选项")

            # 处理合格级别复选框匹配和标记
            if grade_level:
                quality_checkbox_success = process_quality_level_checkboxes(doc, grade_level)
                if quality_checkbox_success:
                    print("合格级别复选框处理完成")
                else:
                    print("合格级别复选框处理失败或未找到匹配选项")

            print("==== 复选框处理完成 ====\n")

            # 保存文档
            try:
                print(f"\n正在保存文档到: {report_output_path}")
                doc.save(report_output_path)
                print(f"文档已成功保存至: {report_output_path}")
                success_count += 1
            except Exception as e:
                print(f"错误: 无法保存文档: {e}")
                error_count += 1
                
        except Exception as e:
            print(f"错误: 处理委托单编号 {order_number} 和射线类型 {ray_type} 时出错: {e}")
            error_count += 1
    
    print(f"\n处理完成: 共处理{len(groups)}个组合，成功生成{success_count}份报告，失败{error_count}份")
    if error_count > 0:
        print(f"警告: 有{error_count}个组合处理失败，请检查日志")

    # 输出模板使用总结
    print_template_usage_summary(template_usage_summary)

    return success_count > 0

# 新增函数：查找像质计灵敏度
def find_sensitivity_value(specification, ray_type):
    """
    根据规格和射线类型查找对应的像质计灵敏度值
    
    Args:
        specification: 规格值
        ray_type: 射线类型 ("X射线" 或 "γ射线")
    
    Returns:
        str: 像质计灵敏度值，如果未找到则返回空字符串
    """
    try:
        # 根据射线类型选择不同的Excel文件
        if ray_type == "X射线":
            excel_path = "生成器/Excel/4_生成器X射线指导书模版.xlsx"
            print(f"查找X射线像质计灵敏度，使用文件: {excel_path}")
        else:  # γ射线
            excel_path = "生成器/Excel/4_生成器γ射线指导书模版.xlsx"
            print(f"查找γ射线像质计灵敏度，使用文件: {excel_path}")
        
        # 检查文件是否存在
        if not os.path.exists(excel_path):
            print(f"错误: 像质计灵敏度查询文件不存在: {excel_path}")
            return ""
        
        # 读取Excel文件
        df = pd.read_excel(excel_path)
        
        # 打印列名，帮助调试
        print(f"文件 {excel_path} 的列名: {list(df.columns)}")
        
        # 检查A列是否存在
        if len(df.columns) == 0:
            print(f"错误: Excel文件 {excel_path} 没有任何列")
            return ""
        
        # 使用A列作为规格列（第一列）
        spec_column = df.columns[0]
        print(f"使用A列 '{spec_column}' 作为规格列")
        
        # 使用I列作为像质计灵敏度列（第9列，因为索引从0开始）
        if len(df.columns) <= 8:
            print(f"错误: Excel文件 {excel_path} 没有足够的列数来使用I列，当前列数: {len(df.columns)}")
            return ""
        
        sensitivity_column = df.columns[8]  # I列
        print(f"使用I列 '{sensitivity_column}' 作为像质计灵敏度列")
        
        # 清理规格字符串，便于匹配
        clean_spec = specification.strip()
        print(f"开始查找规格值: '{clean_spec}'")
        
        # 在规格列中查找匹配项
        for idx, row_spec in enumerate(df[spec_column]):
            if pd.isna(row_spec):
                continue
                
            row_spec_str = str(row_spec).strip()
            
            # 检查是否精确匹配
            if clean_spec == row_spec_str:
                # 找到匹配项，获取对应的像质计灵敏度值
                sensitivity = df.iloc[idx][sensitivity_column]
                if pd.isna(sensitivity):
                    print(f"警告: 规格 '{clean_spec}' 对应的像质计灵敏度值为空")
                    return ""
                
                print(f"找到规格 '{clean_spec}' 对应的像质计灵敏度值: {sensitivity}")
                return str(sensitivity)
        
        # 如果没有找到精确匹配，尝试部分匹配
        for idx, row_spec in enumerate(df[spec_column]):
            if pd.isna(row_spec):
                continue
                
            row_spec_str = str(row_spec).strip()
            
            # 提取规格中的数字部分
            spec_numbers = re.findall(r'\d+\.?\d*', clean_spec)
            row_spec_numbers = re.findall(r'\d+\.?\d*', row_spec_str)
            
            # 检查是否部分匹配（检查规格的数字部分是否匹配）
            if spec_numbers and row_spec_numbers and spec_numbers == row_spec_numbers:
                # 找到部分匹配项，获取对应的像质计灵敏度值
                sensitivity = df.iloc[idx][sensitivity_column]
                if pd.isna(sensitivity):
                    print(f"警告: 规格 '{row_spec_str}' (部分匹配 '{clean_spec}') 对应的像质计灵敏度值为空")
                    return ""
                
                print(f"找到规格 '{row_spec_str}' (部分匹配 '{clean_spec}') 对应的像质计灵敏度值: {sensitivity}")
                return str(sensitivity)
            
            # 尝试更宽松的匹配：只要数字部分有重叠即可
            if spec_numbers and row_spec_numbers:
                # 检查是否有共同的数字
                common_numbers = set(spec_numbers).intersection(set(row_spec_numbers))
                if common_numbers:
                    sensitivity = df.iloc[idx][sensitivity_column]
                    if pd.isna(sensitivity):
                        continue
                    
                    print(f"找到规格 '{row_spec_str}' (宽松匹配 '{clean_spec}') 对应的像质计灵敏度值: {sensitivity}")
                    return str(sensitivity)
        
        print(f"警告: 未找到规格 '{clean_spec}' 对应的像质计灵敏度值")
        return ""
    
    except Exception as e:
        print(f"查找像质计灵敏度时出错: {e}")
        return ""

def find_field_options(doc, field_name, field_keywords):
    """通用函数：在Word文档中查找指定字段的复选框选项"""
    field_options = []
    processed_cells = set()

    # 遍历所有表格
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            # 检查每一行是否包含目标字段相关内容
            row_text = ""
            for cell in row.cells:
                row_text += cell.text + " "

            # 如果这一行包含目标字段相关内容，搜索整行的复选框选项
            if any(keyword in row_text for keyword in field_keywords):
                # 查找该行及相邻行中所有包含复选框的单元格
                search_rows = [row_idx]
                if row_idx > 0:
                    search_rows.append(row_idx - 1)  # 上一行
                if row_idx < len(table.rows) - 1:
                    search_rows.append(row_idx + 1)  # 下一行

                for search_row_idx in search_rows:
                    if search_row_idx < 0 or search_row_idx >= len(table.rows):
                        continue

                    search_row = table.rows[search_row_idx]
                    for check_cell_idx, check_cell in enumerate(search_row.cells):
                        cell_key = (table_idx, search_row_idx, check_cell_idx)
                        if cell_key in processed_cells:
                            continue
                        processed_cells.add(cell_key)

                        cell_text = check_cell.text.strip()

                        # 查找包含复选框的选项
                        if '□' in cell_text or '☑' in cell_text or '✓' in cell_text:
                            # 分割多个选项（如果在同一个单元格中）
                            lines = cell_text.split('\n')
                            for line in lines:
                                line = line.strip()
                                if ('□' in line or '☑' in line or '✓' in line) and len(line) > 1:
                                    # 处理单行中的多个复选框选项（如"□100% □50%"）
                                    checkbox_pattern = r'([□☑✓])([^□☑✓]+?)(?=[□☑✓]|$)'
                                    matches = re.findall(checkbox_pattern, line)

                                    if matches:
                                        for checkbox, option_text in matches:
                                            option_text = option_text.strip()
                                            if option_text and len(option_text) > 0:
                                                # 重构原始行文本
                                                original_line = f"{checkbox}{option_text}"

                                                # 避免重复添加相同的选项
                                                existing_option = None
                                                for existing in field_options:
                                                    if (existing['text'] == option_text and
                                                        existing['position'][0] == search_row_idx):
                                                        existing_option = existing
                                                        break

                                                if not existing_option:
                                                    field_options.append({
                                                        'text': option_text,
                                                        'original_line': original_line,
                                                        'cell': check_cell,
                                                        'position': (search_row_idx, check_cell_idx),
                                                        'table_idx': table_idx
                                                    })
                                    else:
                                        # 如果正则匹配失败，使用原来的方法
                                        option_text = line.replace('□', '').replace('☑', '').replace('✓', '').strip()
                                        if option_text and len(option_text) > 0:
                                            # 避免重复添加相同的选项
                                            existing_option = None
                                            for existing in field_options:
                                                if (existing['text'] == option_text and
                                                    existing['position'][0] == search_row_idx):
                                                    existing_option = existing
                                                    break

                                            if not existing_option:
                                                field_options.append({
                                                    'text': option_text,
                                                    'original_line': line,
                                                    'cell': check_cell,
                                                    'position': (search_row_idx, check_cell_idx),
                                                    'table_idx': table_idx
                                                })

    return field_options

def match_field_option(field_value, options, field_patterns):
    """通用函数：将字段值与可用选项进行匹配"""
    if not field_value or not options:
        return None

    normalized_value = normalize_text(field_value)

    best_match = None
    best_score = 0

    for option in options:
        option_text = option['text']
        normalized_option = normalize_text(option_text)

        # 1. 完全匹配
        if normalized_value == normalized_option:
            return option

        # 2. 使用模式匹配
        for pattern_key, pattern_list in field_patterns.items():
            for pattern in pattern_list:
                normalized_pattern = normalize_text(pattern)
                if normalized_pattern == normalized_value:
                    # 检查选项是否包含这个模式
                    if normalized_pattern in normalized_option or pattern_key in normalized_option:
                        score = 1.0  # 模式匹配给最高分
                        if score > best_score:
                            best_score = score
                            best_match = option

        # 3. 包含匹配
        if normalized_value in normalized_option or normalized_option in normalized_value:
            score = min(len(normalized_value), len(normalized_option)) / max(len(normalized_value), len(normalized_option))
            if score > best_score:
                best_score = score
                best_match = option

        # 4. 关键词匹配
        value_keywords = [kw for kw in ['100%', '50%', '20%', '10%', '5%', '1%', 'ⅰ', 'ⅱ', 'ⅲ', 'ⅳ', '级'] if kw in normalized_value]
        option_keywords = [kw for kw in ['100%', '50%', '20%', '10%', '5%', '1%', 'ⅰ', 'ⅱ', 'ⅲ', 'ⅳ', '级'] if kw in normalized_option]

        if value_keywords and option_keywords:
            common_keywords = set(value_keywords) & set(option_keywords)
            if common_keywords:
                score = len(common_keywords) / max(len(value_keywords), len(option_keywords))
                if score > best_score and score > 0.3:  # 关键词匹配阈值
                    best_score = score
                    best_match = option

    if best_match and best_score > 0.3:  # 降低最低匹配阈值
        return best_match

    return None

def mark_field_checkbox(option):
    """通用函数：在匹配的选项前添加勾选标记"""
    try:
        cell = option['cell']
        option_text = option['text']
        original_line = option['original_line']

        # 遍历单元格中的所有段落
        for paragraph in cell.paragraphs:
            paragraph_text = paragraph.text.strip()

            # 检查段落是否包含目标选项
            if option_text in paragraph_text and ('□' in paragraph_text or '☑' in paragraph_text or '✓' in paragraph_text):
                # 清空段落并重新构建
                paragraph.clear()

                # 分割段落文本为多行
                lines = paragraph_text.split('\n')
                for i, line in enumerate(lines):
                    line = line.strip()

                    if line:
                        # 使用精确匹配和替换特定选项
                        marked_line = mark_specific_option_in_line(line, option_text, original_line)

                        # 如果行内容发生了变化（即包含打勾符号），需要分别设置字体
                        if marked_line != line:
                            # 分别处理打勾符号和其他文本的字体
                            add_mixed_font_text(paragraph, marked_line)
                        else:
                            # 没有变化，使用默认字体
                            run = paragraph.add_run(marked_line)
                            try:
                                run.font.name = "宋体"
                                run.font.size = Pt(9.5)
                                if run._element.rPr is not None:
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
                            except:
                                pass  # 忽略字体设置错误

                    # 如果不是最后一行，添加换行
                    if i < len(lines) - 1:
                        paragraph.add_run('\n')

                return True

        return False

    except Exception as e:
        return False

def add_mixed_font_text(paragraph, text):
    """添加混合字体的文本，打勾符号使用小五号字体(9磅)，其他文本使用五号字体(9.5磅)"""
    try:
        i = 0
        while i < len(text):
            char = text[i]

            # 如果是打勾符号，使用小五号字体
            if char == '☑':
                run = paragraph.add_run(char)
                run.font.name = "楷体"
                run.font.size = Pt(9)  # 小五号字体
                if run._element.rPr is not None:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
                i += 1
            else:
                # 收集连续的非打勾符号字符
                normal_text = ""
                while i < len(text) and text[i] != '☑':
                    normal_text += text[i]
                    i += 1

                # 添加普通文本，使用五号字体
                if normal_text:
                    run = paragraph.add_run(normal_text)
                    run.font.name = "宋体"
                    run.font.size = Pt(9.5)  # 五号字体
                    if run._element.rPr is not None:
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    except Exception as e:
        print(f"设置混合字体时出错: {e}")
        # 如果出错，回退到普通方式
        run = paragraph.add_run(text)
        run.font.name = "宋体"
        run.font.size = Pt(9.5)
        if run._element.rPr is not None:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

def mark_specific_option_in_line(line, option_text, original_line):
    """在一行文本中精确标记特定选项，不影响其他选项"""
    try:
        # 如果行中包含目标选项文本
        if option_text in line:
            # 构建精确匹配的正则表达式模式
            # 匹配 "□选项文本" 或 "☑选项文本" 的模式
            escaped_option = re.escape(option_text)

            # 尝试多种匹配模式
            patterns = [
                f'□{escaped_option}(?=\\s|□|☑|$)',  # □选项文本（后面跟空格、其他复选框或行尾）
                f'☑{escaped_option}(?=\\s|□|☑|$)',  # ☑选项文本（后面跟空格、其他复选框或行尾）
                f'□\\s*{escaped_option}(?=\\s|□|☑|$)',  # □ 选项文本（中间可能有空格）
                f'☑\\s*{escaped_option}(?=\\s|□|☑|$)'   # ☑ 选项文本（中间可能有空格）
            ]

            marked_line = line
            for pattern in patterns:
                if re.search(pattern, line):
                    # 只替换匹配的部分，将□替换为☑
                    marked_line = re.sub(f'□(\\s*{escaped_option})', f'☑\\1', marked_line)
                    break

            return marked_line

        return line

    except Exception as e:
        return line

def process_detection_ratio_checkboxes(doc, detection_ratio):
    """处理检测比例复选框匹配和标记"""
    try:
        # 定义检测比例的匹配规则
        ratio_patterns = {
            '100%': ['100%', '100', '全部', '百分之百'],
            '50%': ['50%', '50', '百分之五十'],
            '20%': ['20%', '20', '百分之二十'],
            '10%': ['10%', '10', '百分之十'],
            '5%': ['5%', '5', '百分之五'],
            '1%': ['1%', '1', '百分之一']
        }

        # 查找所有检测比例选项
        ratio_options = find_field_options(doc, "检测比例", ["检测比例", "100%", "50%", "20%", "10%", "5%", "1%", "比例"])

        if not ratio_options:
            return False

        # 匹配检测比例值与选项
        matched_option = match_field_option(detection_ratio, ratio_options, ratio_patterns)

        if matched_option:
            # 标记匹配的选项
            success = mark_field_checkbox(matched_option)
            if success:
                print(f"成功标记检测比例选项: '{matched_option['text']}'")
                return True
            else:
                print(f"标记检测比例选项失败")
                return False
        else:
            print(f"未找到匹配的检测比例选项")
            return False

    except Exception as e:
        print(f"处理检测比例复选框时出错: {e}")
        return False

def process_quality_level_checkboxes(doc, quality_level):
    """处理合格级别复选框匹配和标记"""
    try:
        # 定义合格级别的匹配规则
        quality_patterns = {
            'Ⅰ': ['Ⅰ', 'I', '1', '一级', '一'],
            'Ⅱ': ['Ⅱ', 'II', '2', '二级', '二'],
            'Ⅲ': ['Ⅲ', 'III', '3', '三级', '三'],
            'Ⅳ': ['Ⅳ', 'IV', '4', '四级', '四']
        }

        # 查找所有合格级别选项
        quality_options = find_field_options(doc, "合格级别", ["合格级别", "Ⅰ", "Ⅱ", "Ⅲ", "Ⅳ", "级别"])

        if not quality_options:
            return False

        # 匹配合格级别值与选项
        matched_option = match_field_option(quality_level, quality_options, quality_patterns)

        if matched_option:
            # 标记匹配的选项
            success = mark_field_checkbox(matched_option)
            if success:
                print(f"成功标记合格级别选项: '{matched_option['text']}'")
                return True
            else:
                print(f"标记合格级别选项失败")
                return False
        else:
            print(f"未找到匹配的合格级别选项")
            return False

    except Exception as e:
        print(f"处理合格级别复选框时出错: {e}")
        return False

def main():
    # 创建命令行参数解析器
    parser = argparse.ArgumentParser(description='将Excel数据填入Word文档')
    parser.add_argument('-e', '--excel', default="生成器/Excel/5_生成器评片记录续表模版.xlsx", 
                        help='Excel表格路径 (默认: 生成器/Excel/5_生成器评片记录续表模版.xlsx)')
    parser.add_argument('-w', '--word', default="生成器/word/5_射线检测记录_续_新.docx", 
                        help='Word模板文档路径 (默认: 生成器/word/5_射线检测记录_续_新.docx)')
    parser.add_argument('-o', '--output', 
                        help='输出目录 (可选，默认为"生成器/输出报告/5_射线检测记录续"目录)')
    parser.add_argument('-p', '--project', 
                        help='工程名称 (用于替换Word文档中的"工程名称值")')
    parser.add_argument('-c', '--client', 
                        help='委托单位 (用于替换Word文档中的"委托单位值")')
    parser.add_argument('-i', '--instruction', 
                        help='操作指导书编号 (用于替换Word文档中的"操作指导书编号值")')
    
    # 解析命令行参数
    args = parser.parse_args()
    
    # 处理Excel到Word的转换
    success = process_excel_to_word(args.excel, args.word, args.output, 
                                   args.project, args.client, args.instruction)
    
    # 返回状态码
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
