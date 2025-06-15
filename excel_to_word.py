import pandas as pd
import os
import sys
import argparse
from docx import Document
from datetime import datetime
import re

def find_column_with_keyword(df, keyword):
    """查找包含指定关键字的列"""
    matching_cols = [col for col in df.columns if keyword.lower() in col.lower()]
    return matching_cols[0] if matching_cols else None

def process_excel_to_word(excel_path, word_template_path, output_path):
    """将Excel数据填入Word文档
    
    Args:
        excel_path: Excel表格路径
        word_template_path: Word模板文档路径
        output_path: 输出Word文档路径
    
    Returns:
        bool: 处理是否成功
    """
    # 检查文件是否存在
    if not os.path.exists(excel_path):
        print(f"错误: Excel文件不存在: {excel_path}")
        return False
        
    if not os.path.exists(word_template_path):
        print(f"错误: Word模板文件不存在: {word_template_path}")
        return False
    
    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 读取Excel数据
    print(f"正在读取Excel文件: {excel_path}")
    df = pd.read_excel(excel_path)
    
    # 打印所有列名，帮助调试
    print(f"Excel表格列名: {list(df.columns)}")
    
    # 1) 获取A列(委托日期)的最晚日期
    # 确保日期列是日期类型
    date_col = find_column_with_keyword(df, '委托日期')
    if not date_col:
        print("警告: 未找到'委托日期'列")
        return False
        
    df[date_col] = pd.to_datetime(df[date_col], errors='coerce')
    latest_date = df[date_col].max()
    
    if pd.isna(latest_date):
        print("警告: 未找到有效的委托日期")
        return False
    
    # 将日期转换为年、月、日
    year = latest_date.year
    month = latest_date.month
    day = latest_date.day
    
    print(f"找到最晚委托日期: {year}年{month}月{day}日")
    
    # 2) 获取相关列的数据，使用模糊匹配
    # 定义需要查找的列关键字
    column_keywords = {
        '检件编号': '检件编号',
        '焊口编号': '焊口编号',
        '焊工号': '焊工号',
        '规格': '规格',
        '材质': '材质'
    }
    
    # 查找每个关键字对应的实际列名
    column_mapping = {}
    missing_columns = []
    
    for key, keyword in column_keywords.items():
        col_name = find_column_with_keyword(df, keyword)
        if col_name:
            column_mapping[key] = col_name
            print(f"找到列: '{key}' -> '{col_name}'")
        else:
            missing_columns.append(key)
    
    if missing_columns:
        print(f"警告: 未找到以下列: {', '.join(missing_columns)}")
        return False
    
    # 获取所有数据列的值
    # 获取D列(检件编号)的值 - 保留为列表，不去重
    pipe_codes = df[column_mapping['检件编号']].dropna().tolist()
    print(f"找到{len(pipe_codes)}个检件编号: {pipe_codes}")
    
    # 获取E列(焊口编号)的值
    weld_numbers = df[column_mapping['焊口编号']].dropna().tolist()
    print(f"找到{len(weld_numbers)}个焊口编号")
    
    # 获取F列(焊工号)的值
    welder_numbers = df[column_mapping['焊工号']].dropna().tolist()
    print(f"找到{len(welder_numbers)}个焊工号")
    
    # 获取G列(规格)的值
    specifications = df[column_mapping['规格']].dropna().tolist()
    print(f"找到{len(specifications)}个规格")
    
    # 获取H列(材质)的值
    materials = df[column_mapping['材质']].dropna().tolist()
    print(f"找到{len(materials)}个材质")
    
    # 打开Word文档
    print(f"正在处理Word文档: {word_template_path}")
    doc = Document(word_template_path)
    
    # 处理表格
    for table in doc.tables:
        # 查找委托人单元格
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                # 1) 处理委托人日期
                if "委托人" in cell.text:
                    print(f"找到委托人单元格: 第{i+1}行, 第{j+1}列")
                    
                    # 先清空单元格中可能包含的日期
                    for paragraph in cell.paragraphs:
                        if "年" in paragraph.text and "月" in paragraph.text and "日" in paragraph.text:
                            paragraph.text = paragraph.text.replace(re.search(r'\d+年', paragraph.text).group(0) if re.search(r'\d+年', paragraph.text) else "年", "年")
                            paragraph.text = paragraph.text.replace(re.search(r'\d+月', paragraph.text).group(0) if re.search(r'\d+月', paragraph.text) else "月", "月")
                            paragraph.text = paragraph.text.replace(re.search(r'\d+日', paragraph.text).group(0) if re.search(r'\d+日', paragraph.text) else "日", "日")
                    
                    # 检查单元格中的所有段落
                    date_found = False
                    for paragraph in cell.paragraphs:
                        if "年" in paragraph.text and "月" in paragraph.text and "日" in paragraph.text:
                            print(f"找到日期段落: {paragraph.text}")
                            
                            # 创建新的文本，确保只有一个年月日
                            new_text = paragraph.text
                            # 确保年月日前没有数字
                            new_text = re.sub(r'\d*年', '年', new_text)
                            new_text = re.sub(r'\d*月', '月', new_text)
                            new_text = re.sub(r'\d*日', '日', new_text)
                            
                            # 在年月日前插入正确的数字
                            new_text = new_text.replace('年', f'{year}年')
                            new_text = new_text.replace('月', f'{month}月')
                            new_text = new_text.replace('日', f'{day}日')
                            
                            paragraph.text = new_text
                            date_found = True
                            print("已更新委托人日期")
                            break
                    
                    # 如果没有找到日期段落，尝试创建新段落
                    if not date_found:
                        print("未在委托人单元格中找到日期段落，尝试其他方法...")
                        # 添加新段落
                        p = cell.add_paragraph(f"{year}年{month}月{day}日")
                        print("已添加委托人日期")
        
        # 查找表头行，确定各列的位置
        column_indices = {}
        header_row_index = -1
        
        # 查找包含"管道编号"、"焊口号"、"焊工号"等的行
        for i, row in enumerate(table.rows):
            header_found = False
            for j, cell in enumerate(row.cells):
                if "管道编号" in cell.text:
                    column_indices["管道编号"] = j
                    header_row_index = i
                    header_found = True
                elif "焊口号" in cell.text:
                    column_indices["焊口号"] = j
                    header_found = True
                elif "焊工号" in cell.text:
                    column_indices["焊工号"] = j
                    header_found = True
                elif "焊口规格" in cell.text:
                    column_indices["焊口规格"] = j
                    header_found = True
                elif "焊口材质" in cell.text:
                    column_indices["焊口材质"] = j
                    header_found = True
            
            if header_found and header_row_index >= 0:
                break
        
        print(f"找到表头行: 第{header_row_index+1}行")
        print(f"列索引: {column_indices}")
        
        # 如果找到表头行，处理数据填充
        if header_row_index >= 0 and column_indices:
            # 获取可用于填充数据的行
            data_rows = []
            for i in range(header_row_index + 1, len(table.rows)):
                if i < len(table.rows):
                    # 检查是否是空行或包含特殊标记的行
                    if "以下空白" in table.rows[i].cells[0].text if len(table.rows[i].cells) > 0 else False:
                        print(f"找到'以下空白'行: 第{i+1}行")
                        break
                    # 添加可用于填充数据的行
                    data_rows.append(i)
            
            print(f"找到{len(data_rows)}行可用于填充数据")
            
            # 确定需要填充的数据行数
            data_count = min(len(pipe_codes), len(weld_numbers), len(welder_numbers), 
                            len(specifications), len(materials))
            print(f"需要填充{data_count}行数据")
            
            # 如果Word表格中的行数不足，需要添加新行
            rows_needed = data_count - len(data_rows)
            if rows_needed > 0:
                print(f"需要添加{rows_needed}行到表格中")
                # 找到最后一行的索引
                last_row_idx = data_rows[-1] if data_rows else header_row_index
                
                # 添加新行
                for _ in range(rows_needed):
                    # 在表格末尾添加一行
                    new_row = table.add_row()
                    data_rows.append(len(table.rows) - 1)  # 添加新行的索引
            
            # 处理每一行数据
            for i in range(data_count):
                if i < len(data_rows):
                    row_idx = data_rows[i]
                    row = table.rows[row_idx]
                    
                    # 1. 填写管道编号
                    if "管道编号" in column_indices and i < len(pipe_codes):
                        col_idx = column_indices["管道编号"]
                        if col_idx < len(row.cells):
                            cell = row.cells[col_idx]
                            if cell.paragraphs:
                                cell.paragraphs[0].text = str(pipe_codes[i])
                                print(f"已更新第{row_idx+1}行管道编号: {pipe_codes[i]}")
                    
                    # 2. 填写焊口号
                    if "焊口号" in column_indices and i < len(weld_numbers):
                        col_idx = column_indices["焊口号"]
                        if col_idx < len(row.cells):
                            cell = row.cells[col_idx]
                            if cell.paragraphs:
                                cell.paragraphs[0].text = str(weld_numbers[i])
                                print(f"已更新第{row_idx+1}行焊口号: {weld_numbers[i]}")
                    
                    # 3. 填写焊工号
                    if "焊工号" in column_indices and i < len(welder_numbers):
                        col_idx = column_indices["焊工号"]
                        if col_idx < len(row.cells):
                            cell = row.cells[col_idx]
                            if cell.paragraphs:
                                cell.paragraphs[0].text = str(welder_numbers[i])
                                print(f"已更新第{row_idx+1}行焊工号: {welder_numbers[i]}")
                    
                    # 4. 填写焊口规格
                    if "焊口规格" in column_indices and i < len(specifications):
                        col_idx = column_indices["焊口规格"]
                        if col_idx < len(row.cells):
                            cell = row.cells[col_idx]
                            if cell.paragraphs:
                                cell.paragraphs[0].text = str(specifications[i])
                                print(f"已更新第{row_idx+1}行焊口规格: {specifications[i]}")
                    
                    # 5. 填写焊口材质
                    if "焊口材质" in column_indices and i < len(materials):
                        col_idx = column_indices["焊口材质"]
                        if col_idx < len(row.cells):
                            cell = row.cells[col_idx]
                            if cell.paragraphs:
                                cell.paragraphs[0].text = str(materials[i])
                                print(f"已更新第{row_idx+1}行焊口材质: {materials[i]}")
    
    # 保存文档
    doc.save(output_path)
    print(f"文档已保存至: {output_path}")
    return True

def main():
    # 创建命令行参数解析器
    parser = argparse.ArgumentParser(description='将Excel数据填入Word文档')
    parser.add_argument('-e', '--excel', required=True, help='Excel表格路径')
    parser.add_argument('-w', '--word', required=True, help='Word模板文档路径')
    parser.add_argument('-o', '--output', required=True, help='输出Word文档路径')
    
    # 解析命令行参数
    args = parser.parse_args()
    
    # 处理Excel到Word的转换
    success = process_excel_to_word(args.excel, args.word, args.output)
    
    # 返回状态码
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main() 